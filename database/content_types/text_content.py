"""Text Content Management Module - Professional Text Content Processing System

Module spécialisé pour la gestion, l'analyse et la protection du contenu textuel
dans la plateforme IA Influencer Agent.

Author: Fahed Mlaiel <mlaiel@live.de>
Team: NLP Expert, Text Analysis Specialist, Content Protection Expert
Copyright: Fahed Mlaiel - All rights reserved

⚠️  AVERTISSEMENT LÉGAL ⚠️
Ce code est la propriété intellectuelle exclusive de Fahed Mlaiel.
Toute utilisation, copie, modification ou distribution non autorisée
est strictement interdite et fera l'objet de poursuites judiciaires.
Contact: mlaiel@live.de
"""

from typing import Dict, List, Any, Optional, Union, Tuple
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
import logging
import hashlib
import json
import asyncio
import re
from enum import Enum

import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.sentiment import SentimentIntensityAnalyzer
import textstat
from langdetect import detect, DetectorFactory
from collections import Counter
import docx
import PyPDF2
from bs4 import BeautifulSoup
import markdown

# Set seed for consistent language detection
DetectorFactory.seed = 0

logger = logging.getLogger(__name__)

class TextFormat(Enum):
    """
Supported text formats with processing capabilities"""

    TXT = {"ext": ".txt", "structured": False, "encoding": "utf-8", "metadata": False}
    MD = {"ext": ".md", "structured": True, "encoding": "utf-8", "metadata": True}
    PDF = {"ext": ".pdf", "structured": True, "encoding": "utf-8", "metadata": True}
    DOCX = {"ext": ".docx", "structured": True, "encoding": "utf-8", "metadata": True}
    HTML = {"ext": ".html", "structured": True, "encoding": "utf-8", "metadata": True}
    RTF = {"ext": ".rtf", "structured": True, "encoding": "utf-8", "metadata": True}
    TEX = {"ext": ".tex", "structured": True, "encoding": "utf-8", "metadata": False}
    ODT = {"ext": ".odt", "structured": True, "encoding": "utf-8", "metadata": True}

class TextContentType(Enum):
    """Text content classification types"""

    ARTICLE = "article"
    BLOG_POST = "blog_post"
    NEWS = "news"
    ACADEMIC_PAPER = "academic_paper"
    TECHNICAL_DOCUMENTATION = "technical_documentation"
    CREATIVE_WRITING = "creative_writing"
    POETRY = "poetry"
    SCRIPT = "script"
    LEGAL_DOCUMENT = "legal_document"
    BUSINESS_DOCUMENT = "business_document"
    EMAIL = "email"
    SOCIAL_MEDIA_POST = "social_media_post"
    PRODUCT_DESCRIPTION = "product_description"
    REVIEW = "review"
    TUTORIAL = "tutorial"
    FAQ = "faq"
    TERMS_CONDITIONS = "terms_conditions"
    PRIVACY_POLICY = "privacy_policy"

class TextGenre(Enum):
    """Text genre classifications"""

    INFORMATIVE = "informative"
    PERSUASIVE = "persuasive"
    NARRATIVE = "narrative"
    DESCRIPTIVE = "descriptive"
    EXPOSITORY = "expository"
    ARGUMENTATIVE = "argumentative"
    INSTRUCTIONAL = "instructional"
    ENTERTAINMENT = "entertainment"

@dataclass
class TextMetadata:
    """Comprehensive text metadata structure"""
    # Basic properties
    character_count: int
    word_count: int
    sentence_count: int
    paragraph_count: int
    language: Optional[str] = None
    encoding: Optional[str] = None
    file_size: Optional[int] = None
    
    # Linguistic analysis
    lexical_diversity: Optional[float] = None
    average_word_length: Optional[float] = None
    average_sentence_length: Optional[float] = None
    readability_score: Optional[float] = None
    reading_time_minutes: Optional[float] = None
    
    # Content analysis
    content_type: Optional[TextContentType] = None
    genre: Optional[TextGenre] = None
    topics: List[str] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    entities: List[Dict[str, Any]] = field(default_factory=list)
    
    # Sentiment analysis
    sentiment_polarity: Optional[float] = None
    sentiment_subjectivity: Optional[float] = None
    emotion_scores: Dict[str, float] = field(default_factory=dict)
    
    # Quality metrics
    grammar_score: Optional[float] = None
    spelling_errors: int = 0
    style_score: Optional[float] = None
    clarity_score: Optional[float] = None
    
    # Structure analysis
    headings: List[str] = field(default_factory=list)
    sections: List[Dict[str, Any]] = field(default_factory=list)
    has_tables: bool = False
    has_images: bool = False
    has_links: bool = False
    
    # Descriptive metadata
    title: Optional[str] = None
    author: Optional[str] = None
    description: Optional[str] = None
    subject: Optional[str] = None
    category: Optional[str] = None
    
    # Rights and licensing
    copyright: Optional[str] = None
    license: Optional[str] = None
    usage_rights: Optional[str] = None
    
    # Publication metadata
    publication_date: Optional[datetime] = None
    last_modified: Optional[datetime] = None
    version: Optional[str] = None
    
    # SEO metadata
    meta_description: Optional[str] = None
    meta_keywords: List[str] = field(default_factory=list)
    seo_score: Optional[float] = None
    
    # Timestamps
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    analyzed_at: Optional[datetime] = None
    
    # Additional metadata
    tags: List[str] = field(default_factory=list)
    custom_metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class TextFingerprint:
    """
Text fingerprint for content identification and protection"""
    content_id: str
    primary_hash: str
    semantic_hash: str
    structural_hash: str
    stylistic_hash: str
    n_gram_hash: str
    sentence_hash: str
    word_frequency_hash: str
    linguistic_features: Optional[np.ndarray] = None
    tfidf_features: Optional[np.ndarray] = None
    semantic_embedding: Optional[np.ndarray] = None
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    confidence_score: float = 0.0
    quality_indicators: Dict[str, float] = field(default_factory=dict)

class TextContentManager:
    """
    Professional text content management system with advanced NLP capabilities
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the Text Content Manager
        
        Args:
            config: Configuration dictionary for text processing
        """
        self.config = config or self._get_default_config()
        self.logger = logging.getLogger(f"{__name__}.TextContentManager")
        self.supported_formats = [fmt.value["ext"] for fmt in TextFormat]
        
        # Initialize processing components
        self._init_components()
        
    def _get_default_config(self) -> Dict[str, Any]:
        """Get default configuration for text processing"""
        return {
            "max_file_size_mb": 10,
            "max_text_length": 1000000,  # 1M characters
            "enable_fingerprinting": True,
            "enable_nlp_analysis": True,
            "enable_sentiment_analysis": True,
            "enable_readability_analysis": True,
            "enable_plagiarism_detection": True,
            "languages_supported": ["en", "de", "fr", "es", "it", "pt"],
            "min_word_count": 10,
            "ngram_sizes": [1, 2, 3],
            "keyword_extraction_count": 20,
            "entity_recognition": True,
            "topic_modeling": True
        }
    
    def _init_components(self):
        """Initialize text processing components"""
        self.logger.info("Initializing Text Content Manager components...")
        
        # Download required NLTK data
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('wordnet', quiet=True)
            nltk.download('vader_lexicon', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
        except Exception as e:
            self.logger.warning(f"NLTK data download failed: {e}")
        
        # Initialize NLTK components
        try:
            self.lemmatizer = WordNetLemmatizer()
            self.sentiment_analyzer = SentimentIntensityAnalyzer()
            self.stop_words = set(stopwords.words('english'))
        except Exception as e:
            self.logger.warning(f"NLTK component initialization failed: {e}")
            self.lemmatizer = None
            self.sentiment_analyzer = None
            self.stop_words = set()
        
        # Text processing patterns
        self.patterns = {
            "email": re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            "url": re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'),
            "phone": re.compile(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'),
            "hashtag": re.compile(r'#\w+'),
            "mention": re.compile(r'@\w+')
        }
        
        self.logger.info("Text Content Manager initialized successfully")
    
    async def process_text_file(
        self,
        file_path: Union[str, Path],
        extract_metadata: bool = True,
        generate_fingerprint: bool = True,
        nlp_analysis: bool = True,
        quality_analysis: bool = True
    ) -> Dict[str, Any]:
        """
        Process text file with comprehensive analysis
        
        Args:
            file_path: Path to text file
            extract_metadata: Whether to extract metadata
            generate_fingerprint: Whether to generate fingerprint
            nlp_analysis: Whether to perform NLP analysis
            quality_analysis: Whether to perform quality analysis
            
        Returns:
            Dict containing processed text information
        """
        try:
            file_path = Path(file_path)
            self.logger.info(f"Processing text file: {file_path}")
            
            # Validate file
            if not await self._validate_text_file(file_path):
                raise ValueError(f"Invalid text file: {file_path}")
            
            # Extract text content
            text_content = await self._extract_text_content(file_path)
            
            if not text_content or len(text_content.strip()) < self.config["min_word_count"]:
                raise ValueError("Text content is too short or empty")
            
            results = {
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "processing_timestamp": datetime.now(timezone.utc),
                "text_length": len(text_content),
                "format": file_path.suffix.lower()
            }
            
            # Extract metadata
            if extract_metadata:
                metadata = await self._extract_text_metadata(file_path, text_content)
                results["metadata"] = metadata
            
            # Generate fingerprint
            if generate_fingerprint:
                fingerprint = await self._generate_text_fingerprint(text_content, str(file_path))
                results["fingerprint"] = fingerprint
            
            # NLP analysis
            if nlp_analysis:
                nlp_results = await self._perform_nlp_analysis(text_content)
                results["nlp_analysis"] = nlp_results
            
            # Quality analysis
            if quality_analysis:
                quality_metrics = await self._analyze_text_quality(text_content)
                results["quality_metrics"] = quality_metrics
            
            # Content classification
            content_type = await self._classify_text_content(text_content, metadata if extract_metadata else None)
            results["content_classification"] = content_type
            
            self.logger.info(f"Text processing completed for: {file_path}")
            return results
            
        except Exception as e:
            self.logger.error(f"Failed to process text file {file_path}: {e}")
            raise
    
    async def _validate_text_file(self, file_path: Path) -> bool:
        """Validate text file format and accessibility"""
        try:
            # Check file existence and size
            if not file_path.exists():
                return False
            
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.config["max_file_size_mb"]:
                self.logger.warning(f"File size {file_size_mb:.2f}MB exceeds limit")
                return False
            
            # Check format support
            if file_path.suffix.lower() not in self.supported_formats:
                return False
            
            # Try to read a portion of the file
            try:
                if file_path.suffix.lower() == '.pdf':
                    # For PDF, just check if it can be opened
                    with open(file_path, 'rb') as f:
                        PyPDF2.PdfReader(f)
                elif file_path.suffix.lower() == '.docx':
                    # For DOCX, check if it can be opened
                    docx.Document(file_path)
                else:
                    # For text files, try reading with different encodings
                    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                f.read(1000)  # Read first 1000 characters
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        return False
                return True
            except Exception:
                return False
                
        except Exception as e:
            self.logger.error(f"Text file validation failed: {e}")
            return False
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extract text content from various file formats"""
        try:
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                return await self._extract_docx_text(file_path)
            elif file_ext == '.html':
                return await self._extract_html_text(file_path)
            elif file_ext == '.md':
                return await self._extract_markdown_text(file_path)
            else:
                return await self._extract_plain_text(file_path)
                
        except Exception as e:
            self.logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            return text_content.strip()
        except Exception as e:
            self.logger.error(f"PDF text extraction failed: {e}")
            raise
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content.strip()
        except Exception as e:
            self.logger.error(f"DOCX text extraction failed: {e}")
            raise
    
    async def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text_content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            return text_content
        except Exception as e:
            self.logger.error(f"HTML text extraction failed: {e}")
            raise
    
    async def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML then extract text
            html_content = markdown.markdown(md_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            return text_content.strip()
        except Exception as e:
            self.logger.error(f"Markdown text extraction failed: {e}")
            raise
    
    async def _extract_plain_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors ignored
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"Plain text extraction failed: {e}")
            raise
    
    async def _extract_text_metadata(self, file_path: Path, text_content: str) -> TextMetadata:
        """Extract comprehensive text metadata"""
        try:
            # Basic text statistics
            char_count = len(text_content)
            words = word_tokenize(text_content.lower()) if self.lemmatizer else text_content.split()
            sentences = sent_tokenize(text_content) if self.lemmatizer else text_content.split('.')
            paragraphs = [p for p in text_content.split('\n\n') if p.strip()]
            
            word_count = len(words)
            sentence_count = len(sentences)
            paragraph_count = len(paragraphs)
            
            metadata = TextMetadata(
                character_count=char_count,
                word_count=word_count,
                sentence_count=sentence_count,
                paragraph_count=paragraph_count,
                file_size=file_path.stat().st_size,
                encoding='utf-8'  # Assume UTF-8 for processed text
            )
            
            # Language detection
            try:
                detected_language = detect(text_content[:1000])  # Use first 1000 chars
                metadata.language = detected_language
            except Exception:
                metadata.language = "unknown"
            
            # Linguistic analysis
            if word_count > 0:
                # Lexical diversity (unique words / total words)
                unique_words = set(words)
                metadata.lexical_diversity = len(unique_words) / word_count
                
                # Average word length
                metadata.average_word_length = sum(len(word) for word in words) / word_count
                
                # Average sentence length
                if sentence_count > 0:
                    metadata.average_sentence_length = word_count / sentence_count
                
                # Reading time (average 200 words per minute)
                metadata.reading_time_minutes = word_count / 200
            
            # Readability analysis
            if self.config.get("enable_readability_analysis", True):
                readability_score = await self._calculate_readability(text_content)
                metadata.readability_score = readability_score
            
            # Keyword extraction
            keywords = await self._extract_keywords(text_content)
            metadata.keywords = keywords
            
            # Sentiment analysis
            if self.config.get("enable_sentiment_analysis", True) and self.sentiment_analyzer:
                sentiment_scores = self.sentiment_analyzer.polarity_scores(text_content)
                metadata.sentiment_polarity = sentiment_scores['compound']
                metadata.emotion_scores = {
                    'positive': sentiment_scores['pos'],
                    'negative': sentiment_scores['neg'],
                    'neutral': sentiment_scores['neu']
                }
            
            # Structure analysis
            structure_analysis = await self._analyze_text_structure(text_content)
            metadata.headings = structure_analysis["headings"]
            metadata.has_links = structure_analysis["has_links"]
            metadata.sections = structure_analysis["sections"]
            
            # Entity recognition (simplified)
            entities = await self._extract_entities(text_content)
            metadata.entities = entities
            
            metadata.analyzed_at = datetime.now(timezone.utc)
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Text metadata extraction failed: {e}")
            raise
    
    async def _calculate_readability(self, text: str) -> float:
        """Calculate readability score using various metrics"""
        try:
            # Use textstat library for readability metrics
            scores = []
            
            # Flesch Reading Ease
            flesch_score = textstat.flesch_reading_ease(text)
            if flesch_score >= 0:
                scores.append(flesch_score / 100)  # Normalize to 0-1
            
            # Flesch-Kincaid Grade Level (inverse normalized)
            fk_grade = textstat.flesch_kincaid_grade(text)
            if fk_grade >= 0:
                scores.append(max(0, 1 - (fk_grade / 20)))  # Normalize to 0-1
            
            # Automated Readability Index (inverse normalized)
            ari = textstat.automated_readability_index(text)
            if ari >= 0:
                scores.append(max(0, 1 - (ari / 20)))  # Normalize to 0-1
            
            # Return average score
            return float(sum(scores) / len(scores)) if scores else 0.5
            
        except Exception as e:
            self.logger.error(f"Readability calculation failed: {e}")
            return 0.5
    
    async def _extract_keywords(self, text: str) -> List[str]:
        """Extract keywords from text using frequency analysis"""
        try:
            if not self.lemmatizer:
                # Simple word frequency if NLTK not available
                words = text.lower().split()
                word_freq = Counter(words)
                return [word for word, _ in word_freq.most_common(10)]
            
            # Tokenize and clean text
            words = word_tokenize(text.lower())
            
            # Remove stopwords and non-alphabetic words
            filtered_words = [
                self.lemmatizer.lemmatize(word) 
                for word in words 
                if word.isalpha() and word not in self.stop_words and len(word) > 2
            ]
            
            # Get most common words
            word_freq = Counter(filtered_words)
            keywords = [word for word, _ in word_freq.most_common(self.config["keyword_extraction_count"])]
            
            return keywords
            
        except Exception as e:
            self.logger.error(f"Keyword extraction failed: {e}")
            return []
    
    async def _analyze_text_structure(self, text: str) -> Dict[str, Any]:
        """Analyze text structure (headings, links, etc.)"""
        try:
            structure = {
                "headings": [],
                "has_links": False,
                "sections": []
            }
            
            # Extract headings (lines that start with # in markdown or are all caps)
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('#'):  # Markdown heading
                    structure["headings"].append(line)
                elif line.isupper() and len(line) > 5 and len(line) < 100:  # Potential heading
                    structure["headings"].append(line)
            
            # Check for links
            if self.patterns["url"].search(text):
                structure["has_links"] = True
            
            # Basic section detection (paragraphs as sections)
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            for i, paragraph in enumerate(paragraphs[:10]):  # Limit to first 10 sections
                structure["sections"].append({
                    "section_id": i,
                    "preview": paragraph[:100] + "..." if len(paragraph) > 100 else paragraph,
                    "word_count": len(paragraph.split())
                })
            
            return structure
            
        except Exception as e:
            self.logger.error(f"Text structure analysis failed: {e}")
            return {"headings": [], "has_links": False, "sections": []}
    
    async def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities from text (simplified)"""
        try:
            entities = []
            
            # Extract emails
            emails = self.patterns["email"].findall(text)
            for email in emails:
                entities.append({
                    "type": "email",
                    "value": email,
                    "confidence": 0.9
                })
            
            # Extract URLs
            urls = self.patterns["url"].findall(text)
            for url in urls:
                entities.append({
                    "type": "url",
                    "value": url,
                    "confidence": 0.9
                })
            
            # Extract phone numbers
            phones = self.patterns["phone"].findall(text)
            for phone in phones:
                entities.append({
                    "type": "phone",
                    "value": phone,
                    "confidence": 0.8
                })
            
            # Extract hashtags
            hashtags = self.patterns["hashtag"].findall(text)
            for hashtag in hashtags:
                entities.append({
                    "type": "hashtag",
                    "value": hashtag,
                    "confidence": 0.9
                })
            
            # Extract mentions
            mentions = self.patterns["mention"].findall(text)
            for mention in mentions:
                entities.append({
                    "type": "mention",
                    "value": mention,
                    "confidence": 0.9
                })
            
            return entities[:50]  # Limit to 50 entities
            
        except Exception as e:
            self.logger.error(f"Entity extraction failed: {e}")
            return []
    
    async def _generate_text_fingerprint(self, text_content: str, content_id: str) -> TextFingerprint:
        """Generate comprehensive text fingerprint for content protection"""
        try:
            # Primary hash (raw text)
            primary_hash = hashlib.sha256(text_content.encode('utf-8')).hexdigest()
            
            # Normalize text for semantic analysis
            normalized_text = self._normalize_text(text_content)
            
            # Semantic hash (normalized content)
            semantic_hash = hashlib.sha256(normalized_text.encode('utf-8')).hexdigest()[:32]
            
            # Structural hash (sentence structure)
            structural_hash = await self._generate_structural_hash(text_content)
            
            # Stylistic hash (writing style features)
            stylistic_hash = await self._generate_stylistic_hash(text_content)
            
            # N-gram hash
            ngram_hash = await self._generate_ngram_hash(text_content)
            
            # Sentence hash (sentence-level analysis)
            sentence_hash = await self._generate_sentence_hash(text_content)
            
            # Word frequency hash
            word_freq_hash = await self._generate_word_frequency_hash(text_content)
            
            # Advanced features
            linguistic_features = await self._extract_linguistic_features(text_content)
            
            # Quality indicators
            quality_indicators = {
                "text_length": len(text_content),
                "word_count": len(text_content.split()),
                "sentence_count": len(text_content.split('.')),
                "unique_words": len(set(text_content.lower().split())),
                "avg_word_length": sum(len(word) for word in text_content.split()) / max(len(text_content.split()), 1)
            }
            
            # Confidence score
            confidence_score = min(1.0, (
                min(quality_indicators["text_length"] / 1000, 1.0) * 0.3 +
                min(quality_indicators["word_count"] / 500, 1.0) * 0.3 +
                min(quality_indicators["unique_words"] / 200, 1.0) * 0.4
            ))
            
            fingerprint = TextFingerprint(
                content_id=hashlib.md5(content_id.encode()).hexdigest(),
                primary_hash=primary_hash,
                semantic_hash=semantic_hash,
                structural_hash=structural_hash,
                stylistic_hash=stylistic_hash,
                n_gram_hash=ngram_hash,
                sentence_hash=sentence_hash,
                word_frequency_hash=word_freq_hash,
                linguistic_features=linguistic_features,
                confidence_score=confidence_score,
                quality_indicators=quality_indicators
            )
            
            return fingerprint
            
        except Exception as e:
            self.logger.error(f"Text fingerprint generation failed: {e}")
            raise
    
    def _normalize_text(self, text: str) -> str:
        """Normalize text for semantic comparison"""
        try:
            # Convert to lowercase
            text = text.lower()
            
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove punctuation except sentence endings
            text = re.sub(r'[^\w\s\.\!\?]', '', text)
            
            # Remove URLs, emails, etc.
            text = self.patterns["url"].sub('', text)
            text = self.patterns["email"].sub('', text)
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Text normalization failed: {e}")
            return text
    
    async def _generate_structural_hash(self, text: str) -> str:
        """Generate hash based on text structure"""
        try:
            sentences = sent_tokenize(text) if self.lemmatizer else text.split('.')
            
            # Extract structural features
            structural_features = [
                len(sentences),
                len(text.split('\n\n')),  # Paragraph count
                len([s for s in sentences if len(s.split()) > 20]),  # Long sentences
                len([s for s in sentences if s.strip().endswith('?')]),  # Questions
                len([s for s in sentences if s.strip().endswith('!')]),  # Exclamations
            ]
            
            # Sentence length distribution
            sentence_lengths = [len(s.split()) for s in sentences]
            if sentence_lengths:
                structural_features.extend([
                    int(np.mean(sentence_lengths)),
                    int(np.std(sentence_lengths)) if len(sentence_lengths) > 1 else 0,
                    max(sentence_lengths),
                    min(sentence_lengths)
                ])
            
            structural_str = json.dumps(structural_features, sort_keys=True)
            return hashlib.sha256(structural_str.encode()).hexdigest()[:32]
            
        except Exception as e:
            self.logger.error(f"Structural hash generation failed: {e}")
            return hashlib.sha256(str(np.random.random()).encode()).hexdigest()[:32]
    
    async def _generate_stylistic_hash(self, text: str) -> str:
        """Generate hash based on writing style"""
        try:
            words = text.split()
            
            # Stylistic features
            stylistic_features = [
                len([w for w in words if len(w) > 6]),  # Long words
                len([w for w in words if w.isupper()]),  # Uppercase words
                len(re.findall(r'[.!?]', text)),  # Punctuation count
                len(re.findall(r'[,;:]', text)),  # Comma/semicolon count
                text.count('"'),  # Quote count
                text.count("'"),  # Apostrophe count
            ]
            
            # Average word length
            if words:
                avg_word_length = sum(len(word) for word in words) / len(words)
                stylistic_features.append(int(avg_word_length * 100))  # Scale for int
            
            stylistic_str = json.dumps(stylistic_features, sort_keys=True)
            return hashlib.sha256(stylistic_str.encode()).hexdigest()[:32]
            
        except Exception as e:
            self.logger.error(f"Stylistic hash generation failed: {e}")
            return hashlib.sha256(str(np.random.random()).encode()).hexdigest()[:32]
    
    async def _generate_ngram_hash(self, text: str) -> str:
        """Generate hash based on n-grams"""
        try:
            words = text.lower().split()
            ngram_features = []
            
            for n in self.config["ngram_sizes"]:
                ngrams = [' '.join(words[i:i+n]) for i in range(len(words)-n+1)]
                ngram_freq = Counter(ngrams)
                
                # Get most common n-grams
                top_ngrams = [ngram for ngram, _ in ngram_freq.most_common(20)]
                ngram_features.extend(top_ngrams)
            
            ngram_str = json.dumps(sorted(ngram_features), sort_keys=True)
            return hashlib.sha256(ngram_str.encode()).hexdigest()[:32]
            
        except Exception as e:
            self.logger.error(f"N-gram hash generation failed: {e}")
            return hashlib.sha256(str(np.random.random()).encode()).hexdigest()[:32]
    
    async def _generate_sentence_hash(self, text: str) -> str:
        """Generate hash based on sentence patterns"""
        try:
            sentences = sent_tokenize(text) if self.lemmatizer else text.split('.')
            
            # Extract sentence features
            sentence_features = []
            for sentence in sentences[:50]:  # Limit to first 50 sentences
                sentence = sentence.strip()
                if sentence:
                    # Sentence length and structure
                    words = sentence.split()
                    sentence_features.extend([
                        len(words),
                        1 if sentence.endswith('?') else 0,
                        1 if sentence.endswith('!') else 0,
                        len([w for w in words if w.isupper()]),
                        sentence.count(',')
                    ])
            
            # Limit feature vector size
            sentence_features = sentence_features[:500]
            
            sentence_str = json.dumps(sentence_features, sort_keys=True)
            return hashlib.sha256(sentence_str.encode()).hexdigest()[:32]
            
        except Exception as e:
            self.logger.error(f"Sentence hash generation failed: {e}")
            return hashlib.sha256(str(np.random.random()).encode()).hexdigest()[:32]
    
    async def _generate_word_frequency_hash(self, text: str) -> str:
        """Generate hash based on word frequency distribution"""
        try:
            words = text.lower().split()
            
            # Remove common stopwords
            filtered_words = [w for w in words if w not in self.stop_words and len(w) > 2]
            
            # Get word frequencies
            word_freq = Counter(filtered_words)
            
            # Get top 100 words with frequencies
            top_words = word_freq.most_common(100)
            
            # Create normalized frequency vector
            freq_features = [freq for word, freq in top_words]
            
            freq_str = json.dumps(freq_features, sort_keys=True)
            return hashlib.sha256(freq_str.encode()).hexdigest()[:32]
            
        except Exception as e:
            self.logger.error(f"Word frequency hash generation failed: {e}")
            return hashlib.sha256(str(np.random.random()).encode()).hexdigest()[:32]
    
    async def _extract_linguistic_features(self, text: str) -> Optional[np.ndarray]:
        """Extract linguistic features for advanced analysis"""
        try:
            words = text.split()
            sentences = sent_tokenize(text) if self.lemmatizer else text.split('.')
            
            # Basic linguistic features
            features = [
                len(text),  # Character count
                len(words),  # Word count
                len(sentences),  # Sentence count
                len(set(words)),  # Unique words
                sum(len(word) for word in words) / max(len(words), 1),  # Avg word length
                len(words) / max(len(sentences), 1),  # Avg sentence length
            ]
            
            # Punctuation features
            features.extend([
                text.count('.'),
                text.count(','),
                text.count('!'),
                text.count('?'),
                text.count(';'),
                text.count(':'),
                text.count('"'),
                text.count("'"),
            ])
            
            # Part-of-speech features (simplified)
            if self.lemmatizer:
                try:
                    import nltk
                    pos_tags = nltk.pos_tag(word_tokenize(text[:1000]))  # Limit for performance
                    pos_counts = Counter(tag for word, tag in pos_tags)
                    
                    # Add common POS tag counts
                    common_pos = ['NN', 'VB', 'JJ', 'RB', 'PRP', 'DT', 'IN', 'CC']
                    for pos in common_pos:
                        features.append(pos_counts.get(pos, 0))
                except Exception:
                    features.extend([0] * 8)  # Default values
            else:
                features.extend([0] * 8)
            
            return np.array(features, dtype=np.float32)
            
        except Exception as e:
            self.logger.error(f"Linguistic feature extraction failed: {e}")
            return True
    
    async def _perform_nlp_analysis(self, text: str) -> Dict[str, Any]:
        """Perform comprehensive NLP analysis"""
        try:
            nlp_results = {}
            
            # Topic modeling (simplified keyword-based)
            keywords = await self._extract_keywords(text)
            nlp_results["keywords"] = keywords
            
            # Sentiment analysis
            if self.sentiment_analyzer:
                sentiment_scores = self.sentiment_analyzer.polarity_scores(text)
                nlp_results["sentiment"] = {
                    "compound": sentiment_scores['compound'],
                    "positive": sentiment_scores['pos'],
                    "negative": sentiment_scores['neg'],
                    "neutral": sentiment_scores['neu']
                }
            
            # Language detection confidence
            try:
                from langdetect import detect_langs
                lang_probs = detect_langs(text[:1000])
                nlp_results["language_detection"] = [
                    {"language": lang.lang, "confidence": lang.prob}
                    for lang in lang_probs[:3]
                ]
            except Exception:
                nlp_results["language_detection"] = []
            
            # Text statistics
            nlp_results["statistics"] = {
                "readability_score": await self._calculate_readability(text),
                "complexity_score": self._calculate_complexity(text),
                "formality_score": self._calculate_formality(text)
            }
            
            return nlp_results
            
        except Exception as e:
            self.logger.error(f"NLP analysis failed: {e}")
            return {}
    
    def _calculate_complexity(self, text: str) -> float:
        """Calculate text complexity score"""
        try:
            words = text.split()
            sentences = text.split('.')
            
            if not words or not sentences:
                return 0.0
            
            # Factors contributing to complexity
            avg_word_length = sum(len(word) for word in words) / len(words)
            avg_sentence_length = len(words) / len(sentences)
            unique_word_ratio = len(set(words)) / len(words)
            
            # Normalize and combine
            complexity = (
                min(avg_word_length / 8, 1.0) * 0.4 +
                min(avg_sentence_length / 20, 1.0) * 0.4 +
                unique_word_ratio * 0.2
            )
            
            return float(complexity)
            
        except Exception as e:
            self.logger.error(f"Complexity calculation failed: {e}")
            return 0.5
    
    def _calculate_formality(self, text: str) -> float:
        """Calculate text formality score"""
        try:
            # Simple heuristic based on word choice and structure
            formal_indicators = [
                len(re.findall(r'\b(however|therefore|furthermore|moreover|consequently)\b', text.lower())),
                len(re.findall(r'\b(utilize|demonstrate|facilitate|implement)\b', text.lower())),
                text.count(';'),
                len(re.findall(r'\b[A-Z][a-z]+\b', text)),  # Proper nouns
            ]
            
            informal_indicators = [
                len(re.findall(r'\b(gonna|wanna|yeah|ok|cool)\b', text.lower())),
                text.count('!'),
                len(re.findall(r'\b(very|really|super|totally)\b', text.lower())),
                len(re.findall(r"['\"]", text)),  # Contractions and quotes
            ]
            
            formal_score = sum(formal_indicators)
            informal_score = sum(informal_indicators)
            
            total_score = formal_score + informal_score
            if total_score == 0:
                return 0.5  # Neutral
            
            return formal_score / total_score
            
        except Exception as e:
            self.logger.error(f"Formality calculation failed: {e}")
            return 0.5
    
    async def _analyze_text_quality(self, text: str) -> Dict[str, float]:
        """Analyze text quality metrics"""
        try:
            quality_metrics = {}
            
            words = text.split()
            sentences = text.split('.')
            
            # Grammar and style (simplified heuristics)
            
            # Sentence variety (different sentence lengths)
            if sentences:
                sentence_lengths = [len(s.split()) for s in sentences if s.strip()]
                if sentence_lengths:
                    length_variance = np.var(sentence_lengths)
                    quality_metrics["sentence_variety"] = min(length_variance / 50, 1.0)
                else:
                    quality_metrics["sentence_variety"] = 0.0
            else:
                quality_metrics["sentence_variety"] = 0.0
            
            # Vocabulary richness (lexical diversity)
            if words:
                unique_words = len(set(word.lower() for word in words))
                vocabulary_richness = unique_words / len(words)
                quality_metrics["vocabulary_richness"] = vocabulary_richness
            else:
                quality_metrics["vocabulary_richness"] = 0.0
            
            # Coherence (simplified - based on topic consistency)
            keywords = await self._extract_keywords(text)
            if keywords:
                # Count how often top keywords appear
                top_keywords = keywords[:5]
                keyword_frequency = sum(text.lower().count(keyword) for keyword in top_keywords)
                coherence_score = min(keyword_frequency / len(words), 0.1) * 10  # Normalize
                quality_metrics["coherence"] = coherence_score
            else:
                quality_metrics["coherence"] = 0.0
            
            # Clarity (based on readability and sentence structure)
            readability = await self._calculate_readability(text)
            clarity_score = readability * 0.7 + quality_metrics["sentence_variety"] * 0.3
            quality_metrics["clarity"] = clarity_score
            
            # Overall quality score
            overall_quality = (
                quality_metrics["vocabulary_richness"] * 0.3 +
                quality_metrics["coherence"] * 0.3 +
                quality_metrics["clarity"] * 0.4
            )
            
            quality_metrics["overall_quality"] = max(0.0, min(1.0, overall_quality))
            
            return quality_metrics
            
        except Exception as e:
            self.logger.error(f"Text quality analysis failed: {e}")
            return {"overall_quality": 0.5, "error": str(e)}
    
    async def _classify_text_content(
        self, 
        text: str, 
        metadata: Optional[TextMetadata] = None
    ) -> TextContentType:
        """Classify text content type using linguistic and structural features"""
        try:
            text_lower = text.lower()
            
            # Academic paper indicators
            academic_keywords = ['abstract', 'methodology', 'conclusion', 'references', 'hypothesis']
            if any(keyword in text_lower for keyword in academic_keywords):
                return TextContentType.ACADEMIC_PAPER
            
            # Technical documentation indicators
            tech_keywords = ['api', 'function', 'parameter', 'configuration', 'installation']
            if any(keyword in text_lower for keyword in tech_keywords):
                return TextContentType.TECHNICAL_DOCUMENTATION
            
            # Legal document indicators
            legal_keywords = ['whereas', 'hereby', 'therein', 'shall', 'pursuant']
            if any(keyword in text_lower for keyword in legal_keywords):
                return TextContentType.LEGAL_DOCUMENT
            
            # News article indicators
            news_keywords = ['breaking', 'reported', 'according to', 'spokesperson', 'announced']
            if any(keyword in text_lower for keyword in news_keywords):
                return TextContentType.NEWS
            
            # Blog post indicators (informal tone, personal pronouns)
            blog_indicators = text_lower.count('i ') + text_lower.count('my ') + text_lower.count('we ')
            if blog_indicators > len(text.split()) * 0.02:  # More than 2% personal pronouns
                return TextContentType.BLOG_POST
            
            # Review indicators
            review_keywords = ['rating', 'stars', 'recommend', 'pros', 'cons', 'overall']
            if any(keyword in text_lower for keyword in review_keywords):
                return TextContentType.REVIEW
            
            # Social media post (short, informal)
            if metadata and metadata.word_count < 50:
                hashtag_count = text.count('#')
                mention_count = text.count('@')
                if hashtag_count > 0 or mention_count > 0:
                    return TextContentType.SOCIAL_MEDIA_POST
            
            # Tutorial indicators
            tutorial_keywords = ['step', 'tutorial', 'how to', 'guide', 'instructions']
            if any(keyword in text_lower for keyword in tutorial_keywords):
                return TextContentType.TUTORIAL
            
            # Poetry indicators (line breaks, rhythm)
            if text.count('\n') > len(text.split()) * 0.1:  # Many line breaks
                return TextContentType.POETRY
            
            # Default to article
            return TextContentType.ARTICLE
            
        except Exception as e:
            self.logger.error(f"Text content classification failed: {e}")
            return TextContentType.ARTICLE  # Default fallback
    
    async def store_content(self, text_content: Dict[str, Any]) -> str:
        """Store processed text content in database"""
        try:
            # Generate unique content ID
            content_id = hashlib.sha256(
                f"{text_content['file_path']}{datetime.now().isoformat()}".encode()
            ).hexdigest()
            
            # Here you would implement database storage
            # For now, return the generated ID
            
            self.logger.info(f"Text content stored with ID: {content_id}")
            return content_id
            
        except Exception as e:
            self.logger.error(f"Failed to store text content: {e}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported text formats"""
        return [fmt.value["ext"] for fmt in TextFormat]
    
    def get_format_info(self, format_name: str) -> Optional[Dict[str, Any]]:
        """Get information about a specific text format"""
        for fmt in TextFormat:
            if fmt.value["ext"] == f".{format_name.lower()}" or fmt.name.lower() == format_name.lower():
                return fmt.value
        return True
