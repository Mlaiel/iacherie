"""
 Metadata Processor - IA Influencer Agent Platform Enterprise
================================================================
Module: backend/data_management/processors/metadata_processor.py
Author: Fahed Mlaiel (mlaiel@live.de)
Team: Lead Dev IA + Backend Senior + ML Engineer + DBA + Security + Microservices + Audio + DevOps + IA Prompt Engineer
Type: Industrial Metadata Processing - Enterprise Production-Ready
Responsibility: Extraction et traitement avancé des métadonnées pour tous types de contenu
=======================================================================================

  PROPRIÉTÉ INTELLECTUELLE EXCLUSIVE - FAHED MLAIEL 
© 2025 Fahed Mlaiel. Tous droits réservés.
Usage non autorisé strictement interdit et passible de poursuites judiciaires.
Contact: mlaiel@live.de

LOGIQUE MÉTIER METADATA PROCESSOR:
Content Input → Format Detection → Technical Metadata → Content Metadata → 
AI-Enhanced Metadata → Privacy Analysis → Enrichment → Standardization
"""

import os
import json
import hashlib
from typing import Dict, List, Optional, Any, Union, Tuple
import asyncio
import aiofiles
from concurrent.futures import ThreadPoolExecutor
import logging
from datetime import datetime, timezone
from pathlib import Path
import mimetypes
import magic

# Media metadata libraries
from mutagen import File as MutagenFile
from PIL import Image
from PIL.ExifTags import TAGS
import exifread
import PyPDF2
import docx
import cv2

# AI and ML libraries
import tensorflow as tf
from transformers import pipeline
import torch
from sentence_transformers import SentenceTransformer

# Geolocation and mapping
import geopy
from geopy.geocoders import Nominatim

from .base_processor import BaseProcessor, AsyncBaseProcessor


class MetadataProcessor(BaseProcessor):
    """Processeur avancé de métadonnées - Production Enterprise"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.supported_types = {
            'audio', 'video', 'image', 'document', 'archive', 'data'
        }
        
        # Initialize extractors
        self._init_extractors()
        
        # Privacy settings
        self.privacy_sensitive_fields = {
            'gps_latitude', 'gps_longitude', 'gps_location',
            'camera_serial', 'owner_name', 'author', 'creator',
            'device_id', 'phone_number', 'email'
        }
        
        # Content classification labels
        self.content_labels = [
            'music', 'speech', 'nature', 'urban', 'indoor', 'outdoor',
            'professional', 'personal', 'commercial', 'educational'
        ]
        
        self.logger = logging.getLogger(__name__)
    
    def _init_extractors(self):
        """Initialize metadata extraction tools"""



        try:
            # AI models for content analysis
            self.content_classifier = pipeline(
                "zero-shot-classification",
                model="facebook/bart-large-mnli"
            )
            
            # Sentence transformer for semantic metadata
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Geocoder for location metadata
            self.geocoder = Nominatim(user_agent="ia_influencer_agent")
            
        except Exception as e:
            self.logger.warning(f"Metadata extractors initialization warning: {e}")
            self.content_classifier = None
            self.sentence_model = None
            self.geocoder = None
    
    def validate_input(self, input_data: Any) -> bool:
        """Valide les données d'entrée pour extraction de métadonnées"""
        if isinstance(input_data, str):
            # File path validation
            return Path(input_data).exists()
        elif isinstance(input_data, bytes):
            # Binary data validation
            return len(input_data) > 0
        elif isinstance(input_data, dict):
            # Structured metadata validation
            return 'content' in input_data or 'file_path' in input_data
        elif hasattr(input_data, 'read'):
            # File-like object
            return True
        
        return False
    
    def process(self, input_data: Any) -> Dict[str, Any]:
        """Traite et extrait toutes les métadonnées"""



        try:
            # Determine content type
            content_type = self._detect_content_type(input_data)
            
            # Extract technical metadata
            technical_metadata = self._extract_technical_metadata(input_data, content_type)
            
            # Extract format-specific metadata
            format_metadata = self._extract_format_metadata(input_data, content_type)
            
            # Extract content-based metadata
            content_metadata = self._extract_content_metadata(input_data, content_type)
            
            # AI-enhanced metadata
            ai_metadata = self._extract_ai_metadata(input_data, content_type)
            
            # Privacy analysis
            privacy_analysis = self._analyze_privacy(format_metadata, content_metadata)
            
            # Semantic enrichment
            semantic_metadata = self._enrich_semantic_metadata(
                technical_metadata, format_metadata, content_metadata
            )
            
            # Generate standard metadata schema
            standardized_metadata = self._standardize_metadata(
                technical_metadata, format_metadata, content_metadata, ai_metadata
            )
            
            return {
                "success": True,
                "content_type": content_type,
                "technical_metadata": technical_metadata,
                "format_specific_metadata": format_metadata,
                "content_metadata": content_metadata,
                "ai_enhanced_metadata": ai_metadata,
                "privacy_analysis": privacy_analysis,
                "semantic_metadata": semantic_metadata,
                "standardized_metadata": standardized_metadata,
                "processing_info": {
                    "processor_version": "3.0.0",
                    "processed_at": datetime.now(timezone.utc).isoformat(),
                    "extraction_method": "comprehensive"
                }
            }
            
        except Exception as e:
            self.logger.error(f"Metadata processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__
            }
    
    def _detect_content_type(self, input_data: Any) -> str:
        """Détecte le type de contenu"""
        if isinstance(input_data, str):
            # File path analysis
            path = Path(input_data)
            mime_type, _ = mimetypes.guess_type(input_data)
            
            if mime_type:
                if mime_type.startswith('audio/'):
                    return 'audio'
                elif mime_type.startswith('video/'):
                    return 'video'
                elif mime_type.startswith('image/'):
                    return 'image'
                elif mime_type.startswith('text/') or 'document' in mime_type:
                    return 'document'
            
            # Fallback to extension
            ext = path.suffix.lower()
            if ext in ['.mp3', '.wav', '.flac', '.ogg', '.aac']:
                return 'audio'
            elif ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
                return 'video'
            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                return 'image'
            elif ext in ['.pdf', '.doc', '.docx', '.txt', '.rtf']:
                return 'document'
        
        elif isinstance(input_data, bytes):
            # Magic number detection
            try:
                file_type = magic.from_buffer(input_data[:1024], mime=True)
                if file_type.startswith('audio/'):
                    return 'audio'
                elif file_type.startswith('video/'):
                    return 'video'
                elif file_type.startswith('image/'):
                    return 'image'
                elif file_type.startswith('text/'):
                    return 'document'
            except:
                pass
        
        return 'unknown'
    
    def _extract_technical_metadata(self, input_data: Any, content_type: str) -> Dict[str, Any]:
        """Extrait les métadonnées techniques de base"""
        metadata = {
            "content_type": content_type,
            "extraction_timestamp": datetime.now(timezone.utc).isoformat()
        }
        
        if isinstance(input_data, str):
            # File-based metadata
            path = Path(input_data)
            stat = path.stat()
            
            metadata.update({
                "filename": path.name,
                "file_extension": path.suffix,
                "file_size_bytes": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_at": datetime.fromtimestamp(stat.st_atime).isoformat(),
                "mime_type": mimetypes.guess_type(input_data)[0],
                "file_permissions": oct(stat.st_mode)[-3:]
            })
            
            # File hashes
            metadata["hashes"] = self._calculate_file_hashes(input_data)
            
        elif isinstance(input_data, bytes):
            # Binary data metadata
            metadata.update({
                "data_size_bytes": len(input_data),
                "mime_type": magic.from_buffer(input_data[:1024], mime=True),
                "hashes": self._calculate_data_hashes(input_data)
            })
        
        return metadata
    
    def _extract_format_metadata(self, input_data: Any, content_type: str) -> Dict[str, Any]:
        """Extrait les métadonnées spécifiques au format"""
        if content_type == 'audio':
            return self._extract_audio_metadata(input_data)
        elif content_type == 'video':
            return self._extract_video_metadata(input_data)
        elif content_type == 'image':
            return self._extract_image_metadata(input_data)
        elif content_type == 'document':
            return self._extract_document_metadata(input_data)
        else:
            return {}
    
    def _extract_audio_metadata(self, input_data: Any) -> Dict[str, Any]:
        """Extrait les métadonnées audio avec Mutagen"""
        if not isinstance(input_data, str):
            return {}
        
        try:
            audio_file = MutagenFile(input_data)
            if audio_file is None:
                return {}
            
            metadata = {
                "format": audio_file.mime[0] if audio_file.mime else "unknown",
                "bitrate": getattr(audio_file.info, 'bitrate', None),
                "length_seconds": getattr(audio_file.info, 'length', None),
                "sample_rate": getattr(audio_file.info, 'sample_rate', None),
                "channels": getattr(audio_file.info, 'channels', None),
                "bits_per_sample": getattr(audio_file.info, 'bits_per_sample', None)
            }
            
            # Extract tags
            tags = {}
            if audio_file.tags:
                common_tags = {
                    'TIT2': 'title', 'TPE1': 'artist', 'TALB': 'album',
                    'TDRC': 'year', 'TCON': 'genre', 'TPE2': 'album_artist',
                    'TRCK': 'track_number', 'TPOS': 'disc_number'
                }
                
                for tag_id, tag_name in common_tags.items():
                    if tag_id in audio_file.tags:
                        tags[tag_name] = str(audio_file.tags[tag_id][0])
            
            metadata["tags"] = tags
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Audio metadata extraction failed: {e}")
            return {}
    
    def _extract_video_metadata(self, input_data: Any) -> Dict[str, Any]:
        """Extrait les métadonnées vidéo avec OpenCV"""
        if not isinstance(input_data, str):
            return {}
        
        try:
            cap = cv2.VideoCapture(input_data)
            
            metadata = {
                "width": int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                "height": int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                "fps": cap.get(cv2.CAP_PROP_FPS),
                "frame_count": int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
                "duration_seconds": cap.get(cv2.CAP_PROP_FRAME_COUNT) / cap.get(cv2.CAP_PROP_FPS),
                "codec": int(cap.get(cv2.CAP_PROP_FOURCC))
            }
            
            cap.release()
            
            # Convert codec to string
            fourcc = metadata["codec"]
            metadata["codec_string"] = "".join([chr((fourcc >> 8 * i) & 0xFF) for i in range(4)])
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Video metadata extraction failed: {e}")
            return {}
    
    def _extract_image_metadata(self, input_data: Any) -> Dict[str, Any]:
        """Extrait les métadonnées image avec PIL et EXIF"""
        metadata = {}
        
        try:
            if isinstance(input_data, str):
                image = Image.open(input_data)
            else:
                import io
                image = Image.open(io.BytesIO(input_data))
            
            # Basic image metadata
            metadata.update({
                "width": image.width,
                "height": image.height,
                "mode": image.mode,
                "format": image.format,
                "has_transparency": image.mode in ('RGBA', 'LA') or 'transparency' in image.info
            })
            
            # EXIF data
            exif_data = {}
            if hasattr(image, '_getexif') and image._getexif():
                exif = image._getexif()
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    
                    # Handle special cases
                    if tag == 'DateTime':
                        try:
                            exif_data[tag] = datetime.strptime(str(value), '%Y:%m:%d %H:%M:%S').isoformat()
                        except:
                            exif_data[tag] = str(value)
                    elif tag in ['GPSInfo']:
                        # Parse GPS data
                        gps_data = self._parse_gps_data(value)
                        if gps_data:
                            exif_data.update(gps_data)
                    else:
                        exif_data[tag] = str(value)
            
            metadata["exif"] = exif_data
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Image metadata extraction failed: {e}")
            return {}
    
    def _extract_document_metadata(self, input_data: Any) -> Dict[str, Any]:
        """Extrait les métadonnées de documents"""
        if not isinstance(input_data, str):
            return {}
        
        path = Path(input_data)
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_pdf_metadata(input_data)
            elif extension in ['.docx', '.doc']:
                return self._extract_docx_metadata(input_data)
            else:
                return {}
        except Exception as e:
            self.logger.warning(f"Document metadata extraction failed: {e}")
            return {}
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées PDF"""



        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "is_encrypted": pdf_reader.is_encrypted
                }
                
                # Document info
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata.update({
                        "title": info.get('/Title', ''),
                        "author": info.get('/Author', ''),
                        "subject": info.get('/Subject', ''),
                        "creator": info.get('/Creator', ''),
                        "producer": info.get('/Producer', ''),
                        "creation_date": info.get('/CreationDate', ''),
                        "modification_date": info.get('/ModDate', '')
                    })
                
                return metadata
        except:
            return {}
    
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées DOCX"""



        try:
            doc = docx.Document(file_path)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections)
            }
            
            # Core properties
            props = doc.core_properties
            metadata.update({
                "title": props.title or '',
                "author": props.author or '',
                "subject": props.subject or '',
                "keywords": props.keywords or '',
                "comments": props.comments or '',
                "category": props.category or '',
                "created": props.created.isoformat() if props.created else '',
                "modified": props.modified.isoformat() if props.modified else '',
                "last_modified_by": props.last_modified_by or ''
            })
            
            return metadata
        except:
            return {}
    
    def _extract_content_metadata(self, input_data: Any, content_type: str) -> Dict[str, Any]:
        """Extrait les métadonnées basées sur le contenu"""
        content_metadata = {
            "content_analysis_timestamp": datetime.now(timezone.utc).isoformat()
        }
        
        if content_type == 'audio':
            content_metadata.update(self._analyze_audio_content(input_data))
        elif content_type == 'video':
            content_metadata.update(self._analyze_video_content(input_data))
        elif content_type == 'image':
            content_metadata.update(self._analyze_image_content(input_data))
        elif content_type == 'document':
            content_metadata.update(self._analyze_document_content(input_data))
        
        return content_metadata
    
    def _analyze_audio_content(self, input_data: Any) -> Dict[str, Any]:
        """Analyse le contenu audio"""
        # Placeholder for audio content analysis
        # In production, this would include:
        # - Tempo detection
        # - Key detection
        # - Genre classification
        # - Mood analysis
        # - Vocal/instrumental detection
        
        return {
            "content_type": "audio",
            "estimated_tempo": None,
            "estimated_key": None,
            "estimated_genre": None,
            "has_vocals": None,
            "energy_level": None
        }
    
    def _analyze_video_content(self, input_data: Any) -> Dict[str, Any]:
        """Analyse le contenu vidéo"""
        # Placeholder for video content analysis
        # In production, this would include:
        # - Scene detection
        # - Object detection
        # - Face detection
        # - Motion analysis
        # - Color analysis
        
        return {
            "content_type": "video",
            "scene_count": None,
            "dominant_colors": [],
            "has_faces": None,
            "motion_intensity": None,
            "brightness_level": None
        }
    
    def _analyze_image_content(self, input_data: Any) -> Dict[str, Any]:
        """Analyse le contenu image"""
        # Placeholder for image content analysis
        # In production, this would include:
        # - Object detection
        # - Scene classification
        # - Color analysis
        # - Composition analysis
        
        return {
            "content_type": "image",
            "dominant_colors": [],
            "detected_objects": [],
            "scene_type": None,
            "artistic_style": None,
            "composition_quality": None
        }
    
    def _analyze_document_content(self, input_data: Any) -> Dict[str, Any]:
        """Analyse le contenu document"""
        # Placeholder for document content analysis
        # In production, this would include:
        # - Topic classification
        # - Sentiment analysis
        # - Reading level
        # - Language detection
        
        return {
            "content_type": "document",
            "estimated_language": None,
            "estimated_topics": [],
            "sentiment_score": None,
            "reading_level": None,
            "word_count_estimate": None
        }
    
    def _extract_ai_metadata(self, input_data: Any, content_type: str) -> Dict[str, Any]:
        """Extrait les métadonnées enrichies par IA"""
        ai_metadata = {
            "ai_analysis_timestamp": datetime.now(timezone.utc).isoformat(),
            "ai_models_used": []
        }
        
        # Content classification
        if self.content_classifier and content_type in ['document', 'image']:
            try:
                classification = self._classify_content_with_ai(input_data, content_type)
                ai_metadata["content_classification"] = classification
                ai_metadata["ai_models_used"].append("content_classifier")
            except Exception as e:
                self.logger.warning(f"AI content classification failed: {e}")
        
        # Semantic embeddings
        if self.sentence_model:
            try:
                embeddings = self._generate_semantic_embeddings(input_data, content_type)
                ai_metadata["semantic_embeddings"] = embeddings
                ai_metadata["ai_models_used"].append("sentence_transformer")
            except Exception as e:
                self.logger.warning(f"Semantic embedding generation failed: {e}")
        
        return ai_metadata
    
    def _classify_content_with_ai(self, input_data: Any, content_type: str) -> Dict[str, Any]:
        """Classification du contenu avec IA"""
        if content_type == 'document' and isinstance(input_data, str):
            # Extract text for classification
            try:
                with open(input_data, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()[:1024]  # Limit text length
                
                classification = self.content_classifier(text, self.content_labels)
                return {
                    "predicted_labels": classification['labels'][:3],
                    "confidence_scores": classification['scores'][:3],
                    "top_prediction": classification['labels'][0]
                }
            except:
                return {}
        
        return {}
    
    def _generate_semantic_embeddings(self, input_data: Any, content_type: str) -> List[float]:
        """Génère des embeddings sémantiques"""
        if content_type == 'document' and isinstance(input_data, str):
            try:
                with open(input_data, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()[:512]  # Limit text length
                
                embedding = self.sentence_model.encode(text)
                return embedding.tolist()[:64]  # Reduce dimensionality
            except:
                return []
        
        return []
    
    def _analyze_privacy(self, format_metadata: Dict, content_metadata: Dict) -> Dict[str, Any]:
        """Analyse les risques de confidentialité"""
        privacy_risks = []
        risk_score = 0.0
        
        # Check for sensitive metadata fields
        all_metadata = {**format_metadata, **content_metadata}
        
        for field, value in all_metadata.items():
            if field.lower() in self.privacy_sensitive_fields:
                if value and str(value).strip():
                    privacy_risks.append({
                        "field": field,
                        "risk_type": "sensitive_metadata",
                        "risk_level": "high" if 'gps' in field.lower() else "medium",
                        "recommendation": f"Consider removing {field} before sharing"
                    })
                    risk_score += 0.3 if 'gps' in field.lower() else 0.1
        
        # Check for GPS data specifically
        if 'exif' in format_metadata:
            exif = format_metadata['exif']
            if any(key.startswith('GPS') for key in exif.keys()):
                privacy_risks.append({
                    "field": "gps_location",
                    "risk_type": "location_data",
                    "risk_level": "high",
                    "recommendation": "Remove GPS/location data before sharing"
                })
                risk_score += 0.5
        
        return {
            "privacy_risks": privacy_risks,
            "overall_risk_score": min(risk_score, 1.0),
            "risk_level": self._get_risk_level(risk_score),
            "privacy_recommendations": self._generate_privacy_recommendations(privacy_risks)
        }
    
    def _enrich_semantic_metadata(self, technical: Dict, format_specific: Dict, content: Dict) -> Dict[str, Any]:
        """Enrichit les métadonnées avec des informations sémantiques"""
        semantic_metadata = {
            "enrichment_timestamp": datetime.now(timezone.utc).isoformat()
        }
        
        # Location enrichment
        if format_specific.get('exif', {}).get('gps_latitude') and format_specific.get('exif', {}).get('gps_longitude'):
            location_info = self._enrich_location_data(
                format_specific['exif']['gps_latitude'],
                format_specific['exif']['gps_longitude']
            )
            semantic_metadata["location_info"] = location_info
        
        # Temporal enrichment
        creation_time = self._extract_creation_time(technical, format_specific)
        if creation_time:
            temporal_info = self._enrich_temporal_data(creation_time)
            semantic_metadata["temporal_info"] = temporal_info
        
        # Technical quality assessment
        quality_assessment = self._assess_technical_quality(technical, format_specific)
        semantic_metadata["quality_assessment"] = quality_assessment
        
        return semantic_metadata
    
    def _standardize_metadata(self, technical: Dict, format_specific: Dict, 
                            content: Dict, ai_enhanced: Dict) -> Dict[str, Any]:
        """Standardise les métadonnées selon des schémas communs"""
        # Dublin Core-inspired standardization
        standardized = {
            "schema_version": "1.0",
            "identifier": technical.get("hashes", {}).get("sha256", ""),
            "title": self._extract_title(format_specific),
            "creator": self._extract_creator(format_specific),
            "subject": self._extract_subject(format_specific, ai_enhanced),
            "description": self._extract_description(format_specific, content),
            "date": self._extract_creation_time(technical, format_specific),
            "type": technical.get("content_type", "unknown"),
            "format": technical.get("mime_type", ""),
            "source": technical.get("filename", ""),
            "language": content.get("estimated_language", ""),
            "coverage": self._extract_coverage(format_specific),
            "rights": self._extract_rights(format_specific)
        }
        
        # Technical specifications
        standardized["technical_specs"] = {
            "file_size": technical.get("file_size_bytes", 0),
            "dimensions": self._extract_dimensions(format_specific),
            "duration": self._extract_duration(format_specific),
            "quality_metrics": self._extract_quality_metrics(format_specific)
        }
        
        return standardized
    
    # Utility methods
    def _calculate_file_hashes(self, file_path: str) -> Dict[str, str]:
        """Calcule les hashes du fichier"""
        hashes = {}
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                hashes["md5"] = hashlib.md5(content).hexdigest()
                hashes["sha1"] = hashlib.sha1(content).hexdigest()
                hashes["sha256"] = hashlib.sha256(content).hexdigest()
        except Exception as e:
            self.logger.warning(f"Hash calculation failed: {e}")
        
        return hashes
    
    def _calculate_data_hashes(self, data: bytes) -> Dict[str, str]:
        """Calcule les hashes des données binaires"""



        return {
            "md5": hashlib.md5(data).hexdigest(),
            "sha1": hashlib.sha1(data).hexdigest(),
            "sha256": hashlib.sha256(data).hexdigest()
        }
    
    def _parse_gps_data(self, gps_info: Dict) -> Optional[Dict[str, Any]]:
        """Parse les données GPS depuis EXIF"""



        try:
            def convert_to_degrees(value):
                d, m, s = value
                return d + (m / 60.0) + (s / 3600.0)
            
            gps_data = {}
            
            if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
                lat = convert_to_degrees(gps_info['GPSLatitude'])
                if gps_info['GPSLatitudeRef'] == 'S':
                    lat = -lat
                gps_data['gps_latitude'] = lat
            
            if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
                lon = convert_to_degrees(gps_info['GPSLongitude'])
                if gps_info['GPSLongitudeRef'] == 'W':
                    lon = -lon
                gps_data['gps_longitude'] = lon
            
            if 'GPSAltitude' in gps_info:
                altitude = float(gps_info['GPSAltitude'])
                if gps_info.get('GPSAltitudeRef') == 1:
                    altitude = -altitude
                gps_data['gps_altitude'] = altitude
            
            if 'GPSTimeStamp' in gps_info:
                gps_data['gps_timestamp'] = str(gps_info['GPSTimeStamp'])
            
            return gps_data if gps_data else None
            
        except Exception as e:
            self.logger.warning(f"GPS data parsing failed: {e}")
            return None
    
    def _enrich_location_data(self, latitude: float, longitude: float) -> Dict[str, Any]:
        """Enrichit les données de localisation"""
        location_info = {
            "coordinates": {"latitude": latitude, "longitude": longitude}
        }
        
        if self.geocoder:
            try:
                location = self.geocoder.reverse(f"{latitude}, {longitude}")
                if location:
                    address = location.raw.get('address', {})
                    location_info.update({
                        "country": address.get('country', ''),
                        "country_code": address.get('country_code', ''),
                        "state": address.get('state', ''),
                        "city": address.get('city', ''),
                        "postal_code": address.get('postcode', ''),
                        "formatted_address": location.address
                    })
            except Exception as e:
                self.logger.warning(f"Geocoding failed: {e}")
        
        return location_info
    
    def _extract_creation_time(self, technical: Dict, format_specific: Dict) -> Optional[str]:
        """Extrait la date de création"""
        # Try different sources for creation time
        candidates = [
            format_specific.get('exif', {}).get('DateTime'),
            format_specific.get('creation_date'),
            format_specific.get('created'),
            technical.get('created_at')
        ]
        
        for candidate in candidates:
            if candidate:
                return str(candidate)
        
        return None
    
    def _enrich_temporal_data(self, creation_time: str) -> Dict[str, Any]:
        """Enrichit les données temporelles"""



        try:
            # Parse creation time
            if 'T' in creation_time:
                dt = datetime.fromisoformat(creation_time.replace('Z', '+00:00'))
            else:
                dt = datetime.fromisoformat(creation_time)
            
            return {
                "year": dt.year,
                "month": dt.month,
                "day": dt.day,
                "hour": dt.hour,
                "weekday": dt.strftime('%A'),
                "season": self._get_season(dt.month),
                "time_of_day": self._get_time_of_day(dt.hour),
                "is_weekend": dt.weekday() >= 5
            }
        except Exception:
            return {}
    
    def _assess_technical_quality(self, technical: Dict, format_specific: Dict) -> Dict[str, Any]:
        """Évalue la qualité technique"""
        quality_score = 0.0
        quality_factors = {}
        
        # File size assessment
        file_size = technical.get("file_size_bytes", 0)
        if file_size > 0:
            if file_size < 100 * 1024:  # < 100KB
                quality_factors["file_size"] = "small"
                quality_score += 0.1
            elif file_size < 10 * 1024 * 1024:  # < 10MB
                quality_factors["file_size"] = "optimal"
                quality_score += 0.3
            else:
                quality_factors["file_size"] = "large"
                quality_score += 0.2
        
        # Resolution assessment (for images/videos)
        width = format_specific.get("width", 0)
        height = format_specific.get("height", 0)
        
        if width and height:
            pixels = width * height
            if pixels >= 8000000:  # 8MP+
                quality_factors["resolution"] = "high"
                quality_score += 0.4
            elif pixels >= 2000000:  # 2MP+
                quality_factors["resolution"] = "medium"
                quality_score += 0.3
            else:
                quality_factors["resolution"] = "low"
                quality_score += 0.1
        
        # Bitrate assessment (for audio/video)
        bitrate = format_specific.get("bitrate", 0)
        if bitrate:
            if bitrate >= 320000:  # 320kbps+
                quality_factors["bitrate"] = "high"
                quality_score += 0.3
            elif bitrate >= 128000:  # 128kbps+
                quality_factors["bitrate"] = "medium"
                quality_score += 0.2
            else:
                quality_factors["bitrate"] = "low"
                quality_score += 0.1
        
        return {
            "overall_score": min(quality_score, 1.0),
            "quality_factors": quality_factors,
            "quality_rating": self._get_quality_rating(quality_score)
        }
    
    def _get_risk_level(self, risk_score: float) -> str:
        """Convertit le score de risque en niveau"""
        if risk_score >= 0.7:
            return "high"
        elif risk_score >= 0.4:
            return "medium"
        elif risk_score >= 0.1:
            return "low"
        else:
            return "minimal"
    
    def _generate_privacy_recommendations(self, privacy_risks: List[Dict]) -> List[str]:
        """Génère des recommandations de confidentialité"""
        recommendations = []
        
        for risk in privacy_risks:
            recommendations.append(risk.get("recommendation", "Review privacy settings"))
        
        if any("gps" in risk.get("field", "").lower() for risk in privacy_risks):
            recommendations.append("Consider using tools to strip location data")
        
        if privacy_risks:
            recommendations.append("Review all metadata before sharing publicly")
        
        return list(set(recommendations))  # Remove duplicates
    
    def _extract_title(self, format_specific: Dict) -> str:
        """Extrait le titre"""
        candidates = [
            format_specific.get('tags', {}).get('title'),
            format_specific.get('title'),
            format_specific.get('exif', {}).get('DocumentName')
        ]
        
        for candidate in candidates:
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        
        return ""
    
    def _extract_creator(self, format_specific: Dict) -> str:
        """Extrait le créateur"""
        candidates = [
            format_specific.get('tags', {}).get('artist'),
            format_specific.get('author'),
            format_specific.get('creator'),
            format_specific.get('exif', {}).get('Artist')
        ]
        
        for candidate in candidates:
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        
        return ""
    
    def _extract_subject(self, format_specific: Dict, ai_enhanced: Dict) -> str:
        """Extrait le sujet"""
        # Try traditional metadata first
        candidates = [
            format_specific.get('tags', {}).get('genre'),
            format_specific.get('subject'),
            format_specific.get('keywords')
        ]
        
        for candidate in candidates:
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        
        # Fallback to AI classification
        ai_classification = ai_enhanced.get('content_classification', {})
        if ai_classification.get('top_prediction'):
            return ai_classification['top_prediction']
        
        return ""
    
    def _extract_description(self, format_specific: Dict, content: Dict) -> str:
        """Extrait la description"""
        candidates = [
            format_specific.get('comments'),
            format_specific.get('description'),
            format_specific.get('exif', {}).get('ImageDescription')
        ]
        
        for candidate in candidates:
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        
        return ""
    
    def _extract_coverage(self, format_specific: Dict) -> str:
        """Extrait la couverture géographique/temporelle"""
        gps_data = format_specific.get('exif', {})
        if gps_data.get('gps_latitude') and gps_data.get('gps_longitude'):
            return f"{gps_data['gps_latitude']}, {gps_data['gps_longitude']}"
        
        return ""
    
    def _extract_rights(self, format_specific: Dict) -> str:
        """Extrait les informations de droits"""
        candidates = [
            format_specific.get('exif', {}).get('Copyright'),
            format_specific.get('copyright'),
            format_specific.get('rights')
        ]
        
        for candidate in candidates:
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        
        return ""
    
    def _extract_dimensions(self, format_specific: Dict) -> Dict[str, Any]:
        """Extrait les dimensions"""
        dimensions = {}
        
        if format_specific.get('width'):
            dimensions['width'] = format_specific['width']
        if format_specific.get('height'):
            dimensions['height'] = format_specific['height']
        
        return dimensions
    
    def _extract_duration(self, format_specific: Dict) -> Optional[float]:
        """Extrait la durée"""



        return format_specific.get('length_seconds') or format_specific.get('duration_seconds')
    
    def _extract_quality_metrics(self, format_specific: Dict) -> Dict[str, Any]:
        """Extrait les métriques de qualité"""
        metrics = {}
        
        if format_specific.get('bitrate'):
            metrics['bitrate'] = format_specific['bitrate']
        if format_specific.get('sample_rate'):
            metrics['sample_rate'] = format_specific['sample_rate']
        if format_specific.get('fps'):
            metrics['fps'] = format_specific['fps']
        
        return metrics
    
    def _get_season(self, month: int) -> str:
        """Détermine la saison"""
        if month in [12, 1, 2]:
            return "winter"
        elif month in [3, 4, 5]:
            return "spring"
        elif month in [6, 7, 8]:
            return "summer"
        else:
            return "autumn"
    
    def _get_time_of_day(self, hour: int) -> str:
        """Détermine le moment de la journée"""
        if 5 <= hour < 12:
            return "morning"
        elif 12 <= hour < 17:
            return "afternoon"
        elif 17 <= hour < 21:
            return "evening"
        else:
            return "night"
    
    def _get_quality_rating(self, score: float) -> str:
        """Convertit le score en rating de qualité"""
        if score >= 0.9:
            return "excellent"
        elif score >= 0.7:
            return "good"
        elif score >= 0.5:
            return "acceptable"
        elif score >= 0.3:
            return "poor"
        else:
            return "very_poor"


class AsyncMetadataProcessor(AsyncBaseProcessor):
    """Version asynchrone du processeur de métadonnées"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.sync_processor = MetadataProcessor(config)
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def validate_input(self, input_data: Any) -> bool:
        """Version asynchrone de la validation"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor, 
            self.sync_processor.validate_input, 
            input_data
        )
    
    async def process(self, input_data: Any) -> Dict[str, Any]:
        """Version asynchrone du traitement"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor, 
            self.sync_processor.process, 
            input_data
        )
    
    async def process_batch(self, input_batch: List[Any]) -> List[Dict[str, Any]]:
        """Traitement en lot asynchrone"""
        tasks = [self.process(item) for item in input_batch]
        return await asyncio.gather(*tasks, return_exceptions=True)
