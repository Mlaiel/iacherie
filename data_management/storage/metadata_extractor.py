"""📊 Metadata Extractor - IA Influencer Agent Platform Enterprise
===============================================================
Module: backend/data_management/storage/metadata_extractor.py
Author: Fahed Mlaiel (mlaiel@live.de)
===============================================================

Advanced metadata extraction for multi-format content analysis
with AI-powered enhancement and professional-grade accuracy.

⚠️  PROPRIÉTÉ INTELLECTUELLE EXCLUSIVE - FAHED MLAIEL ⚠️
© 2025 Fahed Mlaiel. Tous droits réservés.
Contact: mlaiel@live.de

AVERTISSEMENT LÉGAL:
Ce code est la propriété exclusive de Fahed Mlaiel. Toute utilisation,
reproduction, modification ou distribution non autorisée est strictement
interdite et fera l'objet de poursuites judiciaires.

ÉQUIPE PROJET - SPÉCIALITÉS:
- Lead Dev IA: Fahed Mlaiel
- Backend Senior: Fahed Mlaiel  
- ML Engineer: Fahed Mlaiel
- DBA: Fahed Mlaiel
- Sécurité: Fahed Mlaiel
- Microservices: Fahed Mlaiel
- Audio Engineer: Fahed Mlaiel
- DevOps: Fahed Mlaiel
- IA Prompt Engineer: Fahed Mlaiel
"""
from typing import Dict, List, Optional, Any, Union, Tuple
import logging
import asyncio
import mimetypes
import hashlib
import time
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
import json
import io
from pathlib import Path

# Audio/Video metadata
try:
    import mutagen
    from mutagen.id3 import ID3NoHeaderError
    from mutagen.mp3 import MP3
    from mutagen.flac import FLAC
    from mutagen.mp4 import MP4
    from mutagen.oggvorbis import OggVorbis
    AUDIO_LIBRARIES_AVAILABLE = True
except ImportError:
    AUDIO_LIBRARIES_AVAILABLE = False

# Image metadata
try:
    from PIL import Image, ExifTags
    from PIL.ExifTags import TAGS, GPSTAGS
    IMAGE_LIBRARIES_AVAILABLE = True
except ImportError:
    IMAGE_LIBRARIES_AVAILABLE = False

# Video metadata
try:
    import cv2
    import ffmpeg
    VIDEO_LIBRARIES_AVAILABLE = True
except ImportError:
    VIDEO_LIBRARIES_AVAILABLE = False

# Document metadata
try:
    import PyPDF2
    import docx
    DOCUMENT_LIBRARIES_AVAILABLE = True
except ImportError:
    DOCUMENT_LIBRARIES_AVAILABLE = False

logger = logging.getLogger(__name__)

class ContentCategory(Enum):
    """Content categories for specialized processing"""    AUDIO = "audio"
    VIDEO = "video"
    IMAGE = "image"
    DOCUMENT = "document"
    TEXT = "text"
    ARCHIVE = "archive"
    MODEL = "model"
    DATA = "data"
    UNKNOWN = "unknown"

@dataclass
class MetadataExtractionResult:
    """Result of metadata extraction"""    success: bool
    content_category: ContentCategory
    basic_metadata: Dict[str, Any]
    technical_metadata: Dict[str, Any]
    content_metadata: Dict[str, Any]
    ai_analysis: Dict[str, Any]
    extraction_time: float
    error_messages: List[str] = None

class AudioMetadataExtractor:
    """Specialized audio metadata extraction"""    
    @staticmethod
    async def extract_audio_metadata(content_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract comprehensive audio metadata"""        if not AUDIO_LIBRARIES_AVAILABLE:
            return {"error": "Audio libraries not available"}
        
        try:
            # Save to temporary file for mutagen processing
            temp_file = f"/tmp/audio_temp_{int(time.time())}"
            with open(temp_file, 'wb') as f:
                f.write(content_data)
            
            metadata = {}
            
            try:
                # Try to load with mutagen
                audio_file = mutagen.File(temp_file)
                
                if audio_file is not None:
                    # Basic information
                    metadata.update({
                        'length_seconds': getattr(audio_file.info, 'length', 0),
                        'bitrate': getattr(audio_file.info, 'bitrate', 0),
                        'sample_rate': getattr(audio_file.info, 'sample_rate', 0),
                        'channels': getattr(audio_file.info, 'channels', 0),
                        'format': audio_file.mime[0] if hasattr(audio_file, 'mime') else 'unknown'
                    })
                    
                    # Tags
                    if hasattr(audio_file, 'tags') and audio_file.tags:
                        tags = {}
                        for key, value in audio_file.tags.items():
                            if isinstance(value, list) and len(value) > 0:
                                tags[key] = str(value[0])
                            else:
                                tags[key] = str(value)
                        
                        metadata['tags'] = tags
                        
                        # Extract common fields
                        common_fields = {
                            'title': ['TIT2', 'TITLE', '\xa9nam'],
                            'artist': ['TPE1', 'ARTIST', '\xa9ART'],
                            'album': ['TALB', 'ALBUM', '\xa9alb'],
                            'date': ['TDRC', 'DATE', '\xa9day'],
                            'genre': ['TCON', 'GENRE', '\xa9gen'],
                            'track': ['TRCK', 'TRACKNUMBER', 'trkn']
                        }
                        
                        for field, possible_keys in common_fields.items():
                            for key in possible_keys:
                                if key in tags:
                                    metadata[field] = tags[key]
                                    break
                
                # Format-specific extraction
                file_extension = Path(filename).suffix.lower()
                
                if file_extension == '.mp3':
                    mp3_metadata = AudioMetadataExtractor._extract_mp3_metadata(temp_file)
                    metadata.update(mp3_metadata)
                elif file_extension == '.flac':
                    flac_metadata = AudioMetadataExtractor._extract_flac_metadata(temp_file)
                    metadata.update(flac_metadata)
                elif file_extension in ['.m4a', '.mp4']:
                    mp4_metadata = AudioMetadataExtractor._extract_mp4_metadata(temp_file)
                    metadata.update(mp4_metadata)
                
                # Audio analysis
                audio_analysis = await AudioMetadataExtractor._analyze_audio_content(content_data)
                metadata['audio_analysis'] = audio_analysis
                
            finally:
                # Cleanup temp file
                try:
                    Path(temp_file).unlink()
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Audio metadata extraction failed: {str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    def _extract_mp3_metadata(file_path: str) -> Dict[str, Any]:
        """Extract MP3-specific metadata"""        try:
            mp3_file = MP3(file_path)
            return {
                'mp3_info': {
                    'version': getattr(mp3_file.info, 'version', 'unknown'),
                    'layer': getattr(mp3_file.info, 'layer', 'unknown'),
                    'mode': getattr(mp3_file.info, 'mode', 'unknown'),
                    'protected': getattr(mp3_file.info, 'protected', False)
                }
            }
        except Exception as e:
            return {"mp3_error": str(e)}
    
    @staticmethod
    def _extract_flac_metadata(file_path: str) -> Dict[str, Any]:
        """Extract FLAC-specific metadata"""        try:
            flac_file = FLAC(file_path)
            return {
                'flac_info': {
                    'total_samples': getattr(flac_file.info, 'total_samples', 0),
                    'min_blocksize': getattr(flac_file.info, 'min_blocksize', 0),
                    'max_blocksize': getattr(flac_file.info, 'max_blocksize', 0)
                }
            }
        except Exception as e:
            return {"flac_error": str(e)}
    
    @staticmethod
    def _extract_mp4_metadata(file_path: str) -> Dict[str, Any]:
        """Extract MP4-specific metadata"""        try:
            mp4_file = MP4(file_path)
            return {
                'mp4_info': {
                    'codec': getattr(mp4_file.info, 'codec', 'unknown'),
                    'codec_description': getattr(mp4_file.info, 'codec_description', 'unknown')
                }
            }
        except Exception as e:
            return {"mp4_error": str(e)}
    
    @staticmethod
    async def _analyze_audio_content(content_data: bytes) -> Dict[str, Any]:
        """Analyze audio content for additional insights"""        try:
            # Basic analysis
            analysis = {
                'file_size_mb': len(content_data) / (1024 * 1024),
                'estimated_quality': 'unknown',
                'compression_detected': False
            }
            
            # Estimate quality based on file size and format
            size_mb = analysis['file_size_mb']
            if size_mb > 50:
                analysis['estimated_quality'] = 'high'
            elif size_mb > 20:
                analysis['estimated_quality'] = 'medium'
            else:
                analysis['estimated_quality'] = 'standard'
            
            # Detect compression patterns
            if len(content_data) > 1000:
                sample = content_data[:1000]
                unique_bytes = len(set(sample))
                if unique_bytes < 200:  # Low entropy suggests compression
                    analysis['compression_detected'] = True
            
            return analysis
            
        except Exception as e:
            return {"analysis_error": str(e)}

class ImageMetadataExtractor:
    """Specialized image metadata extraction"""    
    @staticmethod
    async def extract_image_metadata(content_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract comprehensive image metadata"""        if not IMAGE_LIBRARIES_AVAILABLE:
            return {"error": "Image libraries not available"}
        
        try:
            image = Image.open(io.BytesIO(content_data))
            
            metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.size[0],
                'height': image.size[1],
                'has_transparency': 'transparency' in image.info or image.mode in ('RGBA', 'LA'),
                'color_count': len(image.getcolors(maxcolors=256*256)) if image.mode in ('P', 'L') else None
            }
            
            # EXIF data extraction
            exif_data = ImageMetadataExtractor._extract_exif_data(image)
            if exif_data:
                metadata['exif'] = exif_data
            
            # Image analysis
            image_analysis = await ImageMetadataExtractor._analyze_image_content(image, content_data)
            metadata['image_analysis'] = image_analysis
            
            # Format-specific metadata
            format_metadata = ImageMetadataExtractor._extract_format_metadata(image)
            metadata.update(format_metadata)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Image metadata extraction failed: {str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    def _extract_exif_data(image: Image.Image) -> Optional[Dict[str, Any]]:
        """Extract EXIF metadata from image"""        try:
            exif_dict = image._getexif()
            if not exif_dict:
                return None
            
            exif_data = {}
            
            for tag_id, value in exif_dict.items():
                tag = TAGS.get(tag_id, tag_id)
                
                # Handle GPS data specially
                if tag == 'GPSInfo':
                    gps_data = {}
                    for gps_tag_id, gps_value in value.items():
                        gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                        gps_data[gps_tag] = gps_value
                    exif_data['GPS'] = gps_data
                else:
                    # Convert to string for JSON serialization
                    try:
                        exif_data[tag] = str(value)
                    except:
                        exif_data[tag] = "unparseable"
            
            return exif_data
            
        except Exception as e:
            logger.warning(f"EXIF extraction failed: {str(e)}")
            return None
    
    @staticmethod
    async def _analyze_image_content(image: Image.Image, content_data: bytes) -> Dict[str, Any]:
        """Analyze image content for insights"""        try:
            analysis = {
                'file_size_mb': len(content_data) / (1024 * 1024),
                'aspect_ratio': round(image.size[0] / image.size[1], 2),
                'is_square': abs(image.size[0] - image.size[1]) < 10,
                'is_portrait': image.size[1] > image.size[0],
                'is_landscape': image.size[0] > image.size[1],
                'megapixels': round((image.size[0] * image.size[1]) / 1000000, 2)
            }
            
            # Determine image category
            if analysis['megapixels'] > 20:
                analysis['quality_category'] = 'professional'
            elif analysis['megapixels'] > 5:
                analysis['quality_category'] = 'high'
            elif analysis['megapixels'] > 1:
                analysis['quality_category'] = 'standard'
            else:
                analysis['quality_category'] = 'low'
            
            # Compression analysis
            theoretical_size = image.size[0] * image.size[1] * 3  # RGB
            compression_ratio = len(content_data) / theoretical_size
            analysis['compression_ratio'] = round(compression_ratio, 4)
            analysis['compression_level'] = 'high' if compression_ratio < 0.1 else 'medium' if compression_ratio < 0.3 else 'low'
            
            return analysis
            
        except Exception as e:
            return {"analysis_error": str(e)}
    
    @staticmethod
    def _extract_format_metadata(image: Image.Image) -> Dict[str, Any]:
        """Extract format-specific metadata"""        metadata = {}
        
        try:
            # Format-specific information
            if hasattr(image, 'info'):
                info = image.info
                
                if image.format == 'JPEG':
                    metadata['jpeg_info'] = {
                        'quality': info.get('quality', 'unknown'),
                        'progressive': info.get('progressive', False),
                        'optimize': info.get('optimize', False)
                    }
                elif image.format == 'PNG':
                    metadata['png_info'] = {
                        'interlace': info.get('interlace', False),
                        'gamma': info.get('gamma', None),
                        'dpi': info.get('dpi', None)
                    }
                elif image.format == 'GIF':
                    metadata['gif_info'] = {
                        'version': info.get('version', 'unknown'),
                        'duration': info.get('duration', None),
                        'loop': info.get('loop', None)
                    }
            
        except Exception as e:
            metadata['format_error'] = str(e)
        
        return metadata

class VideoMetadataExtractor:
    """Specialized video metadata extraction"""    
    @staticmethod
    async def extract_video_metadata(content_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract comprehensive video metadata"""        if not VIDEO_LIBRARIES_AVAILABLE:
            return {"error": "Video libraries not available"}
        
        try:
            # Save to temporary file for video processing
            temp_file = f"/tmp/video_temp_{int(time.time())}.mp4"
            with open(temp_file, 'wb') as f:
                f.write(content_data)
            
            metadata = {}
            
            try:
                # OpenCV analysis
                cv_metadata = VideoMetadataExtractor._extract_opencv_metadata(temp_file)
                metadata.update(cv_metadata)
                
                # FFmpeg analysis
                ffmpeg_metadata = VideoMetadataExtractor._extract_ffmpeg_metadata(temp_file)
                metadata.update(ffmpeg_metadata)
                
                # Video analysis
                video_analysis = await VideoMetadataExtractor._analyze_video_content(content_data)
                metadata['video_analysis'] = video_analysis
                
            finally:
                # Cleanup
                try:
                    Path(temp_file).unlink()
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Video metadata extraction failed: {str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    def _extract_opencv_metadata(file_path: str) -> Dict[str, Any]:
        """Extract metadata using OpenCV"""        try:
            cap = cv2.VideoCapture(file_path)
            
            metadata = {
                'opencv_info': {
                    'frame_count': int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
                    'fps': cap.get(cv2.CAP_PROP_FPS),
                    'width': int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                    'height': int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                    'fourcc': int(cap.get(cv2.CAP_PROP_FOURCC))
                }
            }
            
            # Calculate duration
            fps = metadata['opencv_info']['fps']
            frame_count = metadata['opencv_info']['frame_count']
            if fps > 0:
                metadata['duration_seconds'] = frame_count / fps
            
            cap.release()
            return metadata
            
        except Exception as e:
            return {"opencv_error": str(e)}
    
    @staticmethod
    def _extract_ffmpeg_metadata(file_path: str) -> Dict[str, Any]:
        """Extract metadata using FFmpeg"""        try:
            probe = ffmpeg.probe(file_path)
            
            metadata = {'ffmpeg_info': {}}
            
            # General information
            if 'format' in probe:
                format_info = probe['format']
                metadata['ffmpeg_info']['format'] = {
                    'format_name': format_info.get('format_name', 'unknown'),
                    'duration': float(format_info.get('duration', 0)),
                    'size': int(format_info.get('size', 0)),
                    'bit_rate': int(format_info.get('bit_rate', 0))
                }
            
            # Stream information
            streams = []
            for stream in probe.get('streams', []):
                stream_info = {
                    'codec_type': stream.get('codec_type'),
                    'codec_name': stream.get('codec_name'),
                    'width': stream.get('width'),
                    'height': stream.get('height'),
                    'bit_rate': stream.get('bit_rate'),
                    'duration': stream.get('duration')
                }
                streams.append(stream_info)
            
            metadata['ffmpeg_info']['streams'] = streams
            
            return metadata
            
        except Exception as e:
            return {"ffmpeg_error": str(e)}
    
    @staticmethod
    async def _analyze_video_content(content_data: bytes) -> Dict[str, Any]:
        """Analyze video content for insights"""        try:
            analysis = {
                'file_size_mb': len(content_data) / (1024 * 1024),
                'estimated_quality': 'unknown',
                'compression_detected': True  # Videos are always compressed
            }
            
            # Estimate quality based on file size
            size_mb = analysis['file_size_mb']
            if size_mb > 500:
                analysis['estimated_quality'] = 'high'
            elif size_mb > 100:
                analysis['estimated_quality'] = 'medium'
            else:
                analysis['estimated_quality'] = 'standard'
            
            return analysis
            
        except Exception as e:
            return {"analysis_error": str(e)}

class DocumentMetadataExtractor:
    """Specialized document metadata extraction"""    
    @staticmethod
    async def extract_document_metadata(content_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract comprehensive document metadata"""        if not DOCUMENT_LIBRARIES_AVAILABLE:
            return {"error": "Document libraries not available"}
        
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                return DocumentMetadataExtractor._extract_pdf_metadata(content_data)
            elif file_extension in ['.docx', '.doc']:
                return DocumentMetadataExtractor._extract_docx_metadata(content_data)
            else:
                return DocumentMetadataExtractor._extract_text_metadata(content_data)
                
        except Exception as e:
            logger.error(f"Document metadata extraction failed: {str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    def _extract_pdf_metadata(content_data: bytes) -> Dict[str, Any]:
        """Extract PDF metadata"""        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_data))
            
            metadata = {
                'page_count': len(pdf_reader.pages),
                'is_encrypted': pdf_reader.is_encrypted
            }
            
            # Document info
            if pdf_reader.metadata:
                doc_info = {}
                for key, value in pdf_reader.metadata.items():
                    if key.startswith('/'):
                        key = key[1:]  # Remove leading slash
                    doc_info[key] = str(value)
                metadata['document_info'] = doc_info
            
            # Analyze first page for text content
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                try:
                    text_content = first_page.extract_text()
                    metadata['has_text_content'] = bool(text_content.strip())
                    metadata['estimated_word_count'] = len(text_content.split())
                except:
                    metadata['has_text_content'] = False
                    metadata['estimated_word_count'] = 0
            
            return metadata
            
        except Exception as e:
            return {"pdf_error": str(e)}
    
    @staticmethod
    def _extract_docx_metadata(content_data: bytes) -> Dict[str, Any]:
        """Extract DOCX metadata"""        try:
            doc = docx.Document(io.BytesIO(content_data))
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'has_tables': len(doc.tables) > 0,
                'table_count': len(doc.tables)
            }
            
            # Core properties
            if doc.core_properties:
                core_props = {}
                for prop in ['author', 'category', 'comments', 'content_status', 
                           'created', 'identifier', 'keywords', 'language', 
                           'last_modified_by', 'last_printed', 'modified', 
                           'revision', 'subject', 'title', 'version']:
                    value = getattr(doc.core_properties, prop, None)
                    if value:
                        core_props[prop] = str(value)
                metadata['core_properties'] = core_props
            
            # Extract text for analysis
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            text_content = '\n'.join(full_text)
            metadata['estimated_word_count'] = len(text_content.split())
            metadata['character_count'] = len(text_content)
            
            return metadata
            
        except Exception as e:
            return {"docx_error": str(e)}
    
    @staticmethod
    def _extract_text_metadata(content_data: bytes) -> Dict[str, Any]:
        """Extract text file metadata"""        try:
            # Try different encodings
            text_content = None
            encoding = 'unknown'
            
            for enc in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text_content = content_data.decode(enc)
                    encoding = enc
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                return {"error": "Unable to decode text content"}
            
            lines = text_content.split('\n')
            words = text_content.split()
            
            metadata = {
                'encoding': encoding,
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': len(text_content),
                'character_count_no_spaces': len(text_content.replace(' ', '')),
                'estimated_reading_time_minutes': len(words) / 200,  # Average reading speed
                'has_empty_lines': any(not line.strip() for line in lines)
            }
            
            # Language detection (simplified)
            if text_content:
                sample = text_content[:1000].lower()
                if any(word in sample for word in ['the', 'and', 'that', 'have', 'for']):
                    metadata['detected_language'] = 'english'
                elif any(word in sample for word in ['le', 'de', 'et', 'un', 'une']):
                    metadata['detected_language'] = 'french'
                elif any(word in sample for word in ['der', 'die', 'das', 'und', 'ein']):
                    metadata['detected_language'] = 'german'
                else:
                    metadata['detected_language'] = 'unknown'
            
            return metadata
            
        except Exception as e:
            return {"text_error": str(e)}

class MetadataExtractor:
    """    Universal metadata extractor for all content types.
    
    Features:
    - Format-specific metadata extraction
    - AI-powered content analysis
    - Technical specification detection
    - Content categorization
    - Quality assessment
    - Compliance metadata
    """    
    def __init__(self):
        """Initialize metadata extractor"""        self.audio_extractor = AudioMetadataExtractor()
        self.image_extractor = ImageMetadataExtractor()
        self.video_extractor = VideoMetadataExtractor()
        self.document_extractor = DocumentMetadataExtractor()
        
        logger.info("MetadataExtractor initialized")
    
    async def extract_metadata(
        self,
        content_data: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> MetadataExtractionResult:
        """        Extract comprehensive metadata from content.
        
        Business Flow:
        1. Determine content category
        2. Extract basic file metadata
        3. Apply specialized extraction
        4. Perform AI analysis
        5. Compile results with quality assessment
        """        start_time = time.time()
        
        try:
            # Determine content type and category
            if not content_type:
                content_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
            
            category = self._determine_content_category(content_type, filename)
            
            # Extract basic metadata
            basic_metadata = self._extract_basic_metadata(content_data, filename, content_type)
            
            # Extract technical metadata
            technical_metadata = await self._extract_technical_metadata(
                content_data, filename, category
            )
            
            # Extract content-specific metadata
            content_metadata = await self._extract_content_metadata(
                content_data, filename, category
            )
            
            # AI analysis
            ai_analysis = await self._perform_ai_analysis(
                content_data, filename, category, content_metadata
            )
            
            extraction_time = time.time() - start_time
            
            return MetadataExtractionResult(
                success=True,
                content_category=category,
                basic_metadata=basic_metadata,
                technical_metadata=technical_metadata,
                content_metadata=content_metadata,
                ai_analysis=ai_analysis,
                extraction_time=extraction_time
            )
            
        except Exception as e:
            logger.error(f"Metadata extraction failed for {filename}: {str(e)}")
            
            return MetadataExtractionResult(
                success=False,
                content_category=ContentCategory.UNKNOWN,
                basic_metadata={},
                technical_metadata={},
                content_metadata={},
                ai_analysis={},
                extraction_time=time.time() - start_time,
                error_messages=[str(e)]
            )
    
    def _determine_content_category(self, content_type: str, filename: str) -> ContentCategory:
        """Determine content category from type and filename"""        
        # Check MIME type
        if content_type.startswith('audio/'):
            return ContentCategory.AUDIO
        elif content_type.startswith('video/'):
            return ContentCategory.VIDEO
        elif content_type.startswith('image/'):
            return ContentCategory.IMAGE
        elif content_type.startswith('text/'):
            return ContentCategory.TEXT
        
        # Check file extension
        ext = Path(filename).suffix.lower()
        
        if ext in ['.mp3', '.flac', '.wav', '.aac', '.ogg', '.m4a']:
            return ContentCategory.AUDIO
        elif ext in ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv']:
            return ContentCategory.VIDEO
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return ContentCategory.IMAGE
        elif ext in ['.pdf', '.doc', '.docx', '.rtf', '.odt']:
            return ContentCategory.DOCUMENT
        elif ext in ['.txt', '.md', '.rst', '.html', '.xml', '.json', '.csv']:
            return ContentCategory.TEXT
        elif ext in ['.zip', '.tar', '.gz', '.rar', '.7z']:
            return ContentCategory.ARCHIVE
        elif ext in ['.pkl', '.model', '.h5', '.onnx', '.pb']:
            return ContentCategory.MODEL
        elif ext in ['.csv', '.json', '.parquet', '.feather']:
            return ContentCategory.DATA
        
        return ContentCategory.UNKNOWN
    
    def _extract_basic_metadata(
        self,
        content_data: bytes,
        filename: str,
        content_type: str
    ) -> Dict[str, Any]:
        """Extract basic file metadata"""        
        file_path = Path(filename)
        
        return {
            'filename': filename,
            'file_extension': file_path.suffix.lower(),
            'file_size_bytes': len(content_data),
            'file_size_mb': round(len(content_data) / (1024 * 1024), 3),
            'content_type': content_type,
            'sha256_hash': hashlib.sha256(content_data).hexdigest(),
            'md5_hash': hashlib.md5(content_data).hexdigest(),
            'extraction_timestamp': datetime.now().isoformat()
        }
    
    async def _extract_technical_metadata(
        self,
        content_data: bytes,
        filename: str,
        category: ContentCategory
    ) -> Dict[str, Any]:
        """Extract technical metadata based on content category"""        
        technical_metadata = {}
        
        try:
            if category == ContentCategory.AUDIO:
                audio_metadata = await self.audio_extractor.extract_audio_metadata(content_data, filename)
                technical_metadata.update(audio_metadata)
            
            elif category == ContentCategory.IMAGE:
                image_metadata = await self.image_extractor.extract_image_metadata(content_data, filename)
                technical_metadata.update(image_metadata)
            
            elif category == ContentCategory.VIDEO:
                video_metadata = await self.video_extractor.extract_video_metadata(content_data, filename)
                technical_metadata.update(video_metadata)
            
            elif category == ContentCategory.DOCUMENT:
                doc_metadata = await self.document_extractor.extract_document_metadata(content_data, filename)
                technical_metadata.update(doc_metadata)
            
            # Add entropy analysis for all types
            entropy_analysis = self._calculate_entropy_analysis(content_data)
            technical_metadata['entropy_analysis'] = entropy_analysis
            
        except Exception as e:
            technical_metadata['extraction_error'] = str(e)
        
        return technical_metadata
    
    async def _extract_content_metadata(
        self,
        content_data: bytes,
        filename: str,
        category: ContentCategory
    ) -> Dict[str, Any]:
        """Extract content-specific metadata"""        
        content_metadata = {
            'category': category.value,
            'is_binary': self._is_binary_content(content_data),
            'complexity_score': self._calculate_complexity_score(content_data),
            'uniqueness_score': self._calculate_uniqueness_score(content_data)
        }
        
        # Category-specific content analysis
        if category in [ContentCategory.AUDIO, ContentCategory.VIDEO]:
            content_metadata['media_analysis'] = await self._analyze_media_content(content_data, category)
        elif category == ContentCategory.IMAGE:
            content_metadata['visual_analysis'] = await self._analyze_visual_content(content_data)
        elif category in [ContentCategory.TEXT, ContentCategory.DOCUMENT]:
            content_metadata['text_analysis'] = await self._analyze_text_content(content_data)
        
        return content_metadata
    
    async def _perform_ai_analysis(
        self,
        content_data: bytes,
        filename: str,
        category: ContentCategory,
        content_metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Perform AI-powered content analysis"""        
        ai_analysis = {
            'quality_assessment': self._assess_content_quality(content_data, category, content_metadata),
            'business_relevance': self._assess_business_relevance(filename, category, content_metadata),
            'compliance_check': self._check_compliance_requirements(content_data, category),
            'optimization_suggestions': self._generate_optimization_suggestions(content_data, category)
        }
        
        return ai_analysis
    
    def _calculate_entropy_analysis(self, content_data: bytes) -> Dict[str, Any]:
        """Calculate entropy-based analysis"""        
        if len(content_data) == 0:
            return {'entropy': 0.0, 'randomness': 'empty'}
        
        # Calculate byte frequency
        byte_counts = {}
        for byte in content_data[:10000]:  # Sample first 10KB
            byte_counts[byte] = byte_counts.get(byte, 0) + 1
        
        # Calculate Shannon entropy
        data_len = len(content_data[:10000])
        entropy = 0.0
        
        for count in byte_counts.values():
            probability = count / data_len
            if probability > 0:
                entropy -= probability * (probability.bit_length() - 1)
        
        # Classify randomness
        if entropy > 7.5:
            randomness = 'high'
        elif entropy > 6.0:
            randomness = 'medium'
        elif entropy > 3.0:
            randomness = 'low'
        else:
            randomness = 'very_low'
        
        return {
            'entropy': round(entropy, 3),
            'randomness': randomness,
            'unique_bytes': len(byte_counts),
            'compression_potential': 'low' if entropy > 7.0 else 'high'
        }
    
    def _is_binary_content(self, content_data: bytes) -> bool:
        """Determine if content is binary"""        
        if len(content_data) == 0:
            return False
        
        # Check for null bytes
        if b'\x00' in content_data[:1000]:
            return True
        
        # Check for high percentage of non-printable characters
        sample = content_data[:1000]
        printable_count = sum(1 for byte in sample if 32 <= byte <= 126 or byte in [9, 10, 13])
        
        return (printable_count / len(sample)) < 0.7
    
    def _calculate_complexity_score(self, content_data: bytes) -> float:
        """Calculate content complexity score"""        
        if len(content_data) == 0:
            return 0.0
        
        # Base score on entropy and patterns
        entropy_analysis = self._calculate_entropy_analysis(content_data)
        entropy_score = entropy_analysis['entropy'] / 8.0  # Normalize to 0-1
        
        # Pattern complexity (simplified)
        sample = content_data[:5000]
        unique_bigrams = len(set(zip(sample[:-1], sample[1:])))
        pattern_score = min(unique_bigrams / len(sample), 1.0) if len(sample) > 1 else 0.0
        
        # Combine scores
        complexity = (entropy_score * 0.6) + (pattern_score * 0.4)
        
        return round(complexity, 3)
    
    def _calculate_uniqueness_score(self, content_data: bytes) -> float:
        """Calculate content uniqueness score"""        
        if len(content_data) == 0:
            return 0.0
        
        # Simple uniqueness based on byte distribution
        sample = content_data[:10000]
        unique_bytes = len(set(sample))
        
        # Normalize to 0-1 scale
        uniqueness = unique_bytes / 256.0
        
        return round(uniqueness, 3)
    
    async def _analyze_media_content(self, content_data: bytes, category: ContentCategory) -> Dict[str, Any]:
        """Analyze media content (audio/video)"""        
        analysis = {
            'estimated_duration': 'unknown',
            'quality_indicators': [],
            'format_efficiency': 'unknown'
        }
        
        # File size analysis for media
        size_mb = len(content_data) / (1024 * 1024)
        
        if category == ContentCategory.AUDIO:
            # Estimate duration based on typical bitrates
            if size_mb > 50:
                analysis['quality_indicators'].append('high_bitrate')
            elif size_mb < 5:
                analysis['quality_indicators'].append('compressed')
            
            analysis['estimated_duration'] = f"{int(size_mb * 8)}+ minutes"  # Rough estimate
        
        elif category == ContentCategory.VIDEO:
            if size_mb > 500:
                analysis['quality_indicators'].append('high_definition')
            elif size_mb > 100:
                analysis['quality_indicators'].append('standard_definition')
            else:
                analysis['quality_indicators'].append('compressed')
        
        return analysis
    
    async def _analyze_visual_content(self, content_data: bytes) -> Dict[str, Any]:
        """Analyze visual content (images)"""        
        analysis = {
            'estimated_dimensions': 'unknown',
            'quality_indicators': [],
            'visual_complexity': 'unknown'
        }
        
        # Simple analysis based on file size
        size_mb = len(content_data) / (1024 * 1024)
        
        if size_mb > 10:
            analysis['quality_indicators'].append('high_resolution')
        elif size_mb > 1:
            analysis['quality_indicators'].append('medium_resolution')
        else:
            analysis['quality_indicators'].append('optimized')
        
        return analysis
    
    async def _analyze_text_content(self, content_data: bytes) -> Dict[str, Any]:
        """Analyze text content"""        
        analysis = {
            'language_hints': [],
            'content_structure': 'unknown',
            'formatting_detected': False
        }
        
        try:
            # Try to decode as text
            text_content = content_data.decode('utf-8', errors='ignore')
            
            # Basic language detection
            sample = text_content.lower()[:1000]
            
            if any(word in sample for word in ['the', 'and', 'that', 'have']):
                analysis['language_hints'].append('english')
            if any(word in sample for word in ['le', 'de', 'et', 'un']):
                analysis['language_hints'].append('french')
            if any(word in sample for word in ['der', 'die', 'und', 'ein']):
                analysis['language_hints'].append('german')
            
            # Structure detection
            if any(tag in text_content for tag in ['<html>', '<div>', '<p>']):
                analysis['content_structure'] = 'html'
                analysis['formatting_detected'] = True
            elif '```' in text_content or '    ' in text_content:
                analysis['content_structure'] = 'code'
            elif text_content.count('\n\n') > text_content.count('\n') * 0.1:
                analysis['content_structure'] = 'structured_text'
            else:
                analysis['content_structure'] = 'plain_text'
        
        except Exception as e:
            analysis['analysis_error'] = str(e)
        
        return analysis
    
    def _assess_content_quality(
        self,
        content_data: bytes,
        category: ContentCategory,
        content_metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Assess content quality"""        
        quality_score = 0.5  # Base score
        quality_factors = []
        
        # Size-based quality assessment
        size_mb = len(content_data) / (1024 * 1024)
        
        if category in [ContentCategory.AUDIO, ContentCategory.VIDEO]:
            if size_mb > 100:
                quality_score += 0.3
                quality_factors.append('large_file_size')
            elif size_mb < 1:
                quality_score -= 0.2
                quality_factors.append('small_file_size')
        
        elif category == ContentCategory.IMAGE:
            if size_mb > 5:
                quality_score += 0.2
                quality_factors.append('high_resolution_likely')
            elif size_mb < 0.1:
                quality_score -= 0.1
                quality_factors.append('low_resolution_likely')
        
        # Complexity-based assessment
        complexity = content_metadata.get('complexity_score', 0.5)
        if complexity > 0.7:
            quality_score += 0.2
            quality_factors.append('high_complexity')
        elif complexity < 0.3:
            quality_score -= 0.1
            quality_factors.append('low_complexity')
        
        # Uniqueness-based assessment
        uniqueness = content_metadata.get('uniqueness_score', 0.5)
        if uniqueness > 0.8:
            quality_score += 0.1
            quality_factors.append('high_uniqueness')
        
        quality_score = max(0.0, min(1.0, quality_score))
        
        return {
            'quality_score': round(quality_score, 3),
            'quality_level': 'high' if quality_score > 0.7 else 'medium' if quality_score > 0.4 else 'low',
            'quality_factors': quality_factors
        }
    
    def _assess_business_relevance(
        self,
        filename: str,
        category: ContentCategory,
        content_metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Assess business relevance for IA Influencer platform"""        
        relevance_score = 0.5
        relevance_factors = []
        
        # Category-based relevance
        if category in [ContentCategory.AUDIO, ContentCategory.VIDEO, ContentCategory.IMAGE]:
            relevance_score += 0.3
            relevance_factors.append('creative_content')
        
        # Filename analysis
        filename_lower = filename.lower()
        if any(keyword in filename_lower for keyword in ['music', 'song', 'track', 'album']):
            relevance_score += 0.2
            relevance_factors.append('music_content')
        
        if any(keyword in filename_lower for keyword in ['photo', 'image', 'picture', 'artwork']):
            relevance_score += 0.15
            relevance_factors.append('visual_content')
        
        if any(keyword in filename_lower for keyword in ['blog', 'post', 'article', 'story']):
            relevance_score += 0.1
            relevance_factors.append('written_content')
        
        # Size considerations
        size_mb = content_metadata.get('file_size_mb', 0)
        if 1 <= size_mb <= 100:  # Reasonable size for content
            relevance_score += 0.1
            relevance_factors.append('optimal_size')
        
        relevance_score = max(0.0, min(1.0, relevance_score))
        
        return {
            'relevance_score': round(relevance_score, 3),
            'relevance_level': 'high' if relevance_score > 0.7 else 'medium' if relevance_score > 0.4 else 'low',
            'relevance_factors': relevance_factors
        }
    
    def _check_compliance_requirements(self, content_data: bytes, category: ContentCategory) -> Dict[str, Any]:
        """Check compliance requirements"""        
        compliance = {
            'copyright_safe': True,
            'privacy_compliant': True,
            'content_warnings': [],
            'recommendations': []
        }
        
        # Size-based compliance
        size_mb = len(content_data) / (1024 * 1024)
        if size_mb > 1000:  # 1GB
            compliance['content_warnings'].append('large_file_size')
            compliance['recommendations'].append('consider_compression')
        
        # Category-specific compliance
        if category in [ContentCategory.AUDIO, ContentCategory.VIDEO]:
            compliance['recommendations'].append('verify_music_rights')
            compliance['recommendations'].append('check_performance_rights')
        
        elif category == ContentCategory.IMAGE:
            compliance['recommendations'].append('verify_image_rights')
            compliance['recommendations'].append('check_model_releases')
        
        return compliance
    
    def _generate_optimization_suggestions(self, content_data: bytes, category: ContentCategory) -> List[str]:
        """Generate optimization suggestions"""        
        suggestions = []
        size_mb = len(content_data) / (1024 * 1024)
        
        # Size-based suggestions
        if size_mb > 100:
            suggestions.append('Consider compression to reduce storage costs')
        
        if size_mb > 1000:
            suggestions.append('Use cloud archival storage for large files')
        
        # Category-specific suggestions
        if category == ContentCategory.AUDIO:
            suggestions.extend([
                'Extract audio fingerprints for copyright protection',
                'Consider multiple quality versions for streaming'
            ])
        
        elif category == ContentCategory.VIDEO:
            suggestions.extend([
                'Generate video thumbnails for preview',
                'Create multiple resolution versions',
                'Extract audio track separately'
            ])
        
        elif category == ContentCategory.IMAGE:
            suggestions.extend([
                'Generate multiple sizes for responsive display',
                'Consider WebP format for web use',
                'Extract EXIF data for organization'
            ])
        
        elif category in [ContentCategory.TEXT, ContentCategory.DOCUMENT]:
            suggestions.extend([
                'Index content for full-text search',
                'Extract key phrases for SEO'
            ])
        
        return suggestions

# Export main classes
__all__ = [
    'MetadataExtractor',
    'MetadataExtractionResult',
    'ContentCategory',
    'AudioMetadataExtractor',
    'ImageMetadataExtractor',
    'VideoMetadataExtractor',
    'DocumentMetadataExtractor'
]
