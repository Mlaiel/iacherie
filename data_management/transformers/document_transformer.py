"""📄 Document Transformation Engine - IA Influencer Agent Platform Enterprise
========================================================================
Module: backend/data_management/transformers/document_transformer.py
Author: Fahed Mlaiel (mlaiel@live.de)
========================================================================

⚠️  PROPRIÉTÉ INTELLECTUELLE EXCLUSIVE - FAHED MLAIEL ⚠️
© 2025 Fahed Mlaiel. Tous droits réservés.
Contact: mlaiel@live.de

AVERTISSEMENT: Toute tentative de vol, copie ou utilisation non autorisée
de ce code ou de cette technologie est strictement interdite et sera
poursuivie selon les lois allemandes et internationales.

ÉQUIPE PROJET SPÉCIALISÉE:
- Lead Dev IA: Fahed Mlaiel
- Backend Senior: Fahed Mlaiel  
- ML Engineer: Fahed Mlaiel
- NLP Expert: Fahed Mlaiel
- DevOps Engineer: Fahed Mlaiel
- DBA: Fahed Mlaiel
- Sécurité Expert: Fahed Mlaiel
"""
import asyncio
import logging
import time
import tempfile
import subprocess
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass
from enum import Enum
import json
import re

# Document processing
import pypdf
import docx2txt
from docx import Document as DocxDocument
import openpyxl
from openpyxl import Workbook
import markdown
from markdown.extensions import codehilite, tables, toc
import html2text
from bs4 import BeautifulSoup
import pdfplumber
import fitz  # PyMuPDF

# NLP and text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy
from transformers import pipeline, AutoTokenizer, AutoModel
import torch
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
import textstat
from textblob import TextBlob

# SEO and content optimization
import requests
from urllib.parse import urljoin, urlparse
import yake

from ..models.document_models import DocumentMetadata, DocumentQualityMetrics
from ...core.exceptions import DocumentProcessingError, ValidationError
from ...core.config import get_settings
from ...utils.file_manager import FileManager
from ...utils.validation import validate_document_file

settings = get_settings()
logger = logging.getLogger(__name__)

class DocumentFormat(Enum):
    """Formats de document supportés"""    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    RTF = "rtf"
    ODT = "odt"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"

class DocumentType(Enum):
    """Types de documents pour optimisation"""    BLOG_POST = "blog_post"
    ARTICLE = "article"
    SCRIPT = "script"
    BOOK = "book"
    REPORT = "report"
    PRESENTATION = "presentation"
    SOCIAL_MEDIA = "social_media"
    EMAIL = "email"
    MARKETING = "marketing"

class ContentQuality(Enum):
    """Niveaux de qualité de contenu"""    PROFESSIONAL = "professional"
    STANDARD = "standard"
    CASUAL = "casual"
    TECHNICAL = "technical"

@dataclass
class TextAnalysis:
    """Résultats d'analyse textuelle"""    word_count: int
    sentence_count: int
    paragraph_count: int
    reading_level: str
    sentiment_score: float
    language: str
    keywords: List[str]
    topics: List[str]
    readability_score: float
    seo_score: float

@dataclass
class DocumentProcessingResult:
    """Résultat du traitement de document"""    success: bool
    input_file: str
    output_file: Optional[str]
    original_metadata: DocumentMetadata
    processed_metadata: DocumentMetadata
    quality_metrics: DocumentQualityMetrics
    processing_time: float
    operations_performed: List[str]
    warnings: List[str]
    errors: List[str]

class DocumentAnalyzer:
    """Analyseur de document intelligent pour créateurs de contenu"""    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Initialisation des modèles NLP
        try:
            # Modèle spaCy pour analyse syntaxique
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            self.logger.warning("Modèle spaCy non disponible")
            self.nlp = None
        
        # Pipeline Hugging Face pour analyse de sentiment
        try:
            self.sentiment_analyzer = pipeline(
                "sentiment-analysis",
                model="cardiffnlp/twitter-roberta-base-sentiment-latest"
            )
        except Exception as e:
            self.logger.warning(f"Analyseur sentiment non disponible: {e}")
            self.sentiment_analyzer = None
        
        # Pipeline pour résumé automatique
        try:
            self.summarizer = pipeline(
                "summarization",
                model="facebook/bart-large-cnn"
            )
        except Exception as e:
            self.logger.warning(f"Résumeur non disponible: {e}")
            self.summarizer = None
        
        # Lemmatizer pour extraction de mots-clés
        self.lemmatizer = WordNetLemmatizer()
        
        # Téléchargement des ressources NLTK nécessaires
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('wordnet', quiet=True)
            nltk.download('vader_lexicon', quiet=True)
        except Exception:
            pass
        
        # Extracteur de mots-clés YAKE
        self.keyword_extractor = yake.KeywordExtractor(
            lan="en",
            n=3,  # Mots-clés de 3 mots max
            dedupLim=0.7,
            top=20
        )
    
    def analyze_document_file(self, document_path: str) -> DocumentMetadata:
        """Analyse complète d'un fichier document"""        try:
            # Extraction du texte selon le format
            text_content = self._extract_text_from_file(document_path)
            
            if not text_content.strip():
                raise DocumentProcessingError("Aucun texte extractible du document")
            
            # Analyse textuelle approfondie
            text_analysis = self._analyze_text_content(text_content)
            
            # Métadonnées basiques
            file_size = Path(document_path).stat().st_size
            format_name = Path(document_path).suffix.lower().lstrip('.')
            
            # Classification automatique du type de document
            document_type = self._classify_document_type(text_content, text_analysis)
            
            # Analyse SEO pour contenu web
            seo_analysis = self._analyze_seo_potential(text_content, text_analysis)
            
            # Score de qualité global
            quality_score = self._calculate_content_quality_score(text_analysis, seo_analysis)
            
            return DocumentMetadata(
                filename=Path(document_path).name,
                format=format_name,
                file_size=file_size,
                
                # Contenu textuel
                text_content=text_content[:1000],  # Extrait pour aperçu
                full_text_length=len(text_content),
                
                # Analyse linguistique
                word_count=text_analysis.word_count,
                sentence_count=text_analysis.sentence_count,
                paragraph_count=text_analysis.paragraph_count,
                language=text_analysis.language,
                
                # Métriques de lisibilité
                reading_level=text_analysis.reading_level,
                readability_score=text_analysis.readability_score,
                
                # Analyse sémantique
                sentiment_score=text_analysis.sentiment_score,
                keywords=text_analysis.keywords,
                topics=text_analysis.topics,
                
                # Classification
                document_type=document_type,
                
                # SEO
                seo_score=seo_analysis['score'],
                seo_recommendations=seo_analysis['recommendations'],
                
                # Qualité
                content_quality_score=quality_score,
                
                # Métadonnées d'extraction
                extraction_method=self._get_extraction_method(document_path),
                has_images=self._detect_embedded_media(document_path, 'images'),
                has_tables=self._detect_tables(text_content),
                has_links=self._detect_links(text_content)
            )
            
        except Exception as e:
            self.logger.error(f"Erreur analyse document {document_path}: {e}")
            raise DocumentProcessingError(f"Échec analyse document: {str(e)}")
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """Extrait le texte selon le format de fichier"""        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._extract_from_text(file_path)
            elif file_ext == '.html':
                return self._extract_from_html(file_path)
            elif file_ext == '.xlsx':
                return self._extract_from_excel(file_path)
            elif file_ext == '.csv':
                return self._extract_from_csv(file_path)
            elif file_ext == '.json':
                return self._extract_from_json(file_path)
            else:
                # Fallback: tentative de lecture comme texte
                return self._extract_from_text(file_path)
                
        except Exception as e:
            raise DocumentProcessingError(f"Impossible d'extraire le texte: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extraction de texte depuis PDF"""        
        text_content = ""
        
        try:
            # Tentative avec pdfplumber (meilleur pour les tableaux)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                return text_content
        except Exception:
            pass
        
        try:
            # Fallback avec PyMuPDF
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_content += page.get_text() + "\n"
            doc.close()
            
            if text_content.strip():
                return text_content
        except Exception:
            pass
        
        try:
            # Fallback avec pypdf
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            return text_content
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extraction de texte depuis DOCX"""        
        try:
            # Méthode principale avec docx2txt
            text_content = docx2txt.process(file_path)
            
            if text_content.strip():
                return text_content
        except Exception:
            pass
        
        try:
            # Fallback avec python-docx
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extraction des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            return text_content
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction DOCX: {str(e)}")
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extraction depuis fichier texte"""        
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise DocumentProcessingError("Impossible de décoder le fichier texte")
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extraction depuis HTML"""        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Utilisation de BeautifulSoup pour nettoyer le HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Suppression des scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extraction du texte
            text = soup.get_text()
            
            # Nettoyage
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction HTML: {str(e)}")
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extraction depuis Excel"""        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content += f"\n=== {sheet_name} ===\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return text_content
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction Excel: {str(e)}")
    
    def _extract_from_csv(self, file_path: str) -> str:
        """Extraction depuis CSV"""        
        try:
            import csv
            text_content = ""
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text_content += " | ".join(row) + "\n"
            
            return text_content
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction CSV: {str(e)}")
    
    def _extract_from_json(self, file_path: str) -> str:
        """Extraction depuis JSON"""        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Extraction récursive de tous les textes
            def extract_text_from_obj(obj, depth=0):
                if depth > 10:  # Limite de récursion
                    return ""
                
                text = ""
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        text += f"{key}: {extract_text_from_obj(value, depth+1)}\n"
                elif isinstance(obj, list):
                    for item in obj:
                        text += extract_text_from_obj(item, depth+1) + "\n"
                elif isinstance(obj, str):
                    text += obj + " "
                else:
                    text += str(obj) + " "
                
                return text
            
            return extract_text_from_obj(data)
        except Exception as e:
            raise DocumentProcessingError(f"Échec extraction JSON: {str(e)}")
    
    def _analyze_text_content(self, text: str) -> TextAnalysis:
        """Analyse approfondie du contenu textuel"""        
        # Nettoyage du texte
        cleaned_text = self._clean_text(text)
        
        # Métriques basiques
        word_count = len(word_tokenize(cleaned_text))
        sentence_count = len(sent_tokenize(cleaned_text))
        paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
        
        # Détection de langue
        language = self._detect_language(cleaned_text)
        
        # Analyse de lisibilité
        reading_level = self._calculate_reading_level(cleaned_text)
        readability_score = self._calculate_readability_score(cleaned_text)
        
        # Analyse de sentiment
        sentiment_score = self._analyze_sentiment(cleaned_text)
        
        # Extraction de mots-clés
        keywords = self._extract_keywords(cleaned_text)
        
        # Extraction de topics
        topics = self._extract_topics(cleaned_text)
        
        return TextAnalysis(
            word_count=word_count,
            sentence_count=sentence_count,
            paragraph_count=paragraph_count,
            reading_level=reading_level,
            sentiment_score=sentiment_score,
            language=language,
            keywords=keywords,
            topics=topics,
            readability_score=readability_score,
            seo_score=0.0  # Calculé séparément
        )
    
    def _clean_text(self, text: str) -> str:
        """Nettoie le texte pour l'analyse"""        
        # Suppression des caractères de contrôle
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
        
        # Normalisation des espaces
        text = re.sub(r'\s+', ' ', text)
        
        # Suppression des lignes vides multiples
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def _detect_language(self, text: str) -> str:
        """Détecte la langue du texte"""        
        try:
            from textblob import TextBlob
            blob = TextBlob(text[:1000])  # Échantillon pour détection
            return blob.detect_language()
        except Exception:
            return "en"  # Défaut en anglais
    
    def _calculate_reading_level(self, text: str) -> str:
        """Calcule le niveau de lecture"""        
        try:
            flesch_score = textstat.flesch_reading_ease(text)
            
            if flesch_score >= 90:
                return "very_easy"
            elif flesch_score >= 80:
                return "easy"
            elif flesch_score >= 70:
                return "fairly_easy"
            elif flesch_score >= 60:
                return "standard"
            elif flesch_score >= 50:
                return "fairly_difficult"
            elif flesch_score >= 30:
                return "difficult"
            else:
                return "very_difficult"
        except Exception:
            return "unknown"
    
    def _calculate_readability_score(self, text: str) -> float:
        """Calcule un score de lisibilité normalisé"""        
        try:
            # Plusieurs métriques de lisibilité
            flesch = textstat.flesch_reading_ease(text)
            fkgl = textstat.flesch_kincaid_grade(text)
            ari = textstat.automated_readability_index(text)
            
            # Normalisation et moyenne
            flesch_norm = min(100, max(0, flesch)) / 100
            fkgl_norm = max(0, 1 - (fkgl / 20))  # Grade level inversé
            ari_norm = max(0, 1 - (ari / 20))
            
            return (flesch_norm + fkgl_norm + ari_norm) / 3
        except Exception:
            return 0.5  # Score neutre par défaut
    
    def _analyze_sentiment(self, text: str) -> float:
        """Analyse le sentiment du texte"""        
        try:
            if self.sentiment_analyzer:
                # Découpage en chunks pour éviter les limites de tokens
                chunks = [text[i:i+500] for i in range(0, len(text), 500)]
                sentiments = []
                
                for chunk in chunks[:10]:  # Limite à 10 chunks
                    if chunk.strip():
                        result = self.sentiment_analyzer(chunk)[0]
                        score = result['score']
                        if result['label'] == 'NEGATIVE':
                            score = -score
                        sentiments.append(score)
                
                return sum(sentiments) / len(sentiments) if sentiments else 0.0
            else:
                # Fallback avec TextBlob
                blob = TextBlob(text[:1000])
                return blob.sentiment.polarity
        except Exception:
            return 0.0  # Sentiment neutre par défaut
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extrait les mots-clés importants"""        
        keywords = []
        
        try:
            # Utilisation de YAKE
            yake_keywords = self.keyword_extractor.extract_keywords(text)
            keywords.extend([kw[1] for kw in yake_keywords[:10]])
        except Exception:
            pass
        
        try:
            # Fallback avec TF-IDF
            if not keywords:
                tfidf = TfidfVectorizer(
                    max_features=20,
                    stop_words='english',
                    ngram_range=(1, 2)
                )
                tfidf_matrix = tfidf.fit_transform([text])
                feature_names = tfidf.get_feature_names_out()
                scores = tfidf_matrix.toarray()[0]
                
                # Tri par score TF-IDF
                keyword_scores = list(zip(feature_names, scores))
                keyword_scores.sort(key=lambda x: x[1], reverse=True)
                
                keywords = [kw[0] for kw in keyword_scores[:10]]
        except Exception:
            pass
        
        return keywords[:10]
    
    def _extract_topics(self, text: str) -> List[str]:
        """Extrait les sujets principaux"""        
        topics = []
        
        try:
            if self.nlp:
                # Analyse avec spaCy pour entités nommées
                doc = self.nlp(text[:10000])  # Limite pour performance
                
                # Extraction des entités
                entities = [ent.text.lower() for ent in doc.ents 
                           if ent.label_ in ['PERSON', 'ORG', 'GPE', 'EVENT', 'PRODUCT']]
                
                # Comptage et tri
                from collections import Counter
                entity_counts = Counter(entities)
                topics = [entity for entity, count in entity_counts.most_common(10)]
        except Exception:
            pass
        
        # Fallback: extraction de sujets basiques par patterns
        if not topics:
            # Recherche de patterns de sujets communs
            subject_patterns = [
                r'\b(?:about|regarding|concerning)\s+([a-zA-Z\s]+?)(?:\.|,|\n)',
                r'\b(?:topic|subject|theme):\s*([a-zA-Z\s]+?)(?:\.|,|\n)',
                r'\b([A-Z][a-zA-Z\s]{2,20})\s+(?:is|are|was|were)',
            ]
            
            for pattern in subject_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                topics.extend([match.strip().lower() for match in matches[:5]])
        
        return topics[:10]
    
    def _classify_document_type(self, text: str, analysis: TextAnalysis) -> str:
        """Classifie automatiquement le type de document"""        
        text_lower = text.lower()
        word_count = analysis.word_count
        
        # Détection par mots-clés et structure
        if any(keyword in text_lower for keyword in ['title:', 'author:', 'abstract:', 'introduction:', 'conclusion:']):
            return DocumentType.ARTICLE.value
        
        elif any(keyword in text_lower for keyword in ['fade in:', 'fade out:', 'int.', 'ext.', 'character:']):
            return DocumentType.SCRIPT.value
        
        elif any(keyword in text_lower for keyword in ['chapter', 'prologue', 'epilogue']) and word_count > 10000:
            return DocumentType.BOOK.value
        
        elif any(keyword in text_lower for keyword in ['slide', 'presentation', 'agenda:']):
            return DocumentType.PRESENTATION.value
        
        elif any(keyword in text_lower for keyword in ['#hashtag', '@mention', 'follow', 'like', 'share']):
            return DocumentType.SOCIAL_MEDIA.value
        
        elif any(keyword in text_lower for keyword in ['dear', 'sincerely', 'best regards', 'from:', 'to:']):
            return DocumentType.EMAIL.value
        
        elif any(keyword in text_lower for keyword in ['buy now', 'limited offer', 'discount', 'sale', 'promotion']):
            return DocumentType.MARKETING.value
        
        elif word_count > 5000:
            return DocumentType.REPORT.value
        
        elif word_count > 1000:
            return DocumentType.BLOG_POST.value
        
        else:
            return DocumentType.ARTICLE.value
    
    def _analyze_seo_potential(self, text: str, analysis: TextAnalysis) -> Dict[str, Any]:
        """Analyse le potentiel SEO du contenu"""        
        seo_score = 0.0
        recommendations = []
        
        # Longueur du contenu
        if analysis.word_count >= 300:
            seo_score += 0.2
        else:
            recommendations.append("Augmenter la longueur du contenu (minimum 300 mots)")
        
        # Présence de mots-clés
        if len(analysis.keywords) >= 5:
            seo_score += 0.2
        else:
            recommendations.append("Ajouter plus de mots-clés pertinents")
        
        # Structure du document
        if self._has_good_structure(text):
            seo_score += 0.2
        else:
            recommendations.append("Améliorer la structure avec des titres et sous-titres")
        
        # Lisibilité
        if analysis.readability_score >= 0.6:
            seo_score += 0.2
        else:
            recommendations.append("Améliorer la lisibilité du contenu")
        
        # Présence de liens
        if self._detect_links(text):
            seo_score += 0.1
        else:
            recommendations.append("Ajouter des liens internes et externes pertinents")
        
        # Méta-informations détectables
        if self._has_meta_info(text):
            seo_score += 0.1
        else:
            recommendations.append("Ajouter des méta-descriptions et titres optimisés")
        
        return {
            'score': min(1.0, seo_score),
            'recommendations': recommendations
        }
    
    def _has_good_structure(self, text: str) -> bool:
        """Vérifie si le document a une bonne structure"""        
        # Recherche de titres (markdown ou patterns)
        title_patterns = [
            r'^#+ .+$',  # Markdown headers
            r'^[A-Z][A-Za-z\s]+:$',  # Titres avec deux-points
            r'^[A-Z\s]+$',  # Titres en majuscules
        ]
        
        title_count = 0
        for pattern in title_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            title_count += len(matches)
        
        return title_count >= 2
    
    def _detect_links(self, text: str) -> bool:
        """Détecte la présence de liens"""        
        url_pattern = r'https?://[^\s]+'
        markdown_link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        
        return bool(re.search(url_pattern, text) or re.search(markdown_link_pattern, text))
    
    def _has_meta_info(self, text: str) -> bool:
        """Détecte la présence d'informations méta"""        
        meta_patterns = [
            r'title:\s*.+',
            r'description:\s*.+',
            r'keywords:\s*.+',
            r'author:\s*.+',
        ]
        
        for pattern in meta_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _calculate_content_quality_score(
        self,
        analysis: TextAnalysis,
        seo_analysis: Dict[str, Any]
    ) -> float:
        """Calcule un score de qualité global du contenu"""        
        # Facteurs de qualité
        length_score = min(1.0, analysis.word_count / 1000)  # Optimal autour de 1000 mots
        readability_score = analysis.readability_score
        keyword_score = min(1.0, len(analysis.keywords) / 10)
        structure_score = seo_analysis['score']
        
        # Score pondéré
        quality_score = (
            length_score * 0.25 +
            readability_score * 0.3 +
            keyword_score * 0.2 +
            structure_score * 0.25
        )
        
        return round(quality_score, 3)
    
    def _get_extraction_method(self, file_path: str) -> str:
        """Détermine la méthode d'extraction utilisée"""        
        file_ext = Path(file_path).suffix.lower()
        
        method_mapping = {
            '.pdf': 'pdfplumber/pymupdf',
            '.docx': 'docx2txt/python-docx',
            '.txt': 'text_reader',
            '.md': 'text_reader',
            '.html': 'beautifulsoup',
            '.xlsx': 'openpyxl',
            '.csv': 'csv_reader',
            '.json': 'json_parser'
        }
        
        return method_mapping.get(file_ext, 'text_reader')
    
    def _detect_embedded_media(self, file_path: str, media_type: str) -> bool:
        """Détecte la présence de médias embarqués"""        
        file_ext = Path(file_path).suffix.lower()
        
        if media_type == 'images':
            if file_ext == '.pdf':
                try:
                    doc = fitz.open(file_path)
                    for page_num in range(min(5, doc.page_count)):  # Check first 5 pages
                        page = doc[page_num]
                        if page.get_images():
                            doc.close()
                            return True
                    doc.close()
                except Exception:
                    pass
            
            elif file_ext == '.docx':
                try:
                    doc = DocxDocument(file_path)
                    for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
                        if paragraph.runs:
                            for run in paragraph.runs:
                                if hasattr(run, 'element') and run.element.xpath('.//pic:pic'):
                                    return True
                except Exception:
                    pass
        
        return False
    
    def _detect_tables(self, text: str) -> bool:
        """Détecte la présence de tableaux dans le texte"""        
        # Recherche de patterns de tableaux
        table_patterns = [
            r'\|.+\|.+\|',  # Markdown tables
            r'^\s*[^\n]*\t[^\n]*\t[^\n]*$',  # Tab-separated
            r'^\s*[^\n]*\s{3,}[^\n]*\s{3,}[^\n]*$',  # Space-separated columns
        ]
        
        for pattern in table_patterns:
            if re.search(pattern, text, re.MULTILINE):
                return True
        
        return False

class ContentOptimizer:
    """Optimisateur de contenu pour créateurs"""    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Modèle pour amélioration de texte
        try:
            self.text_improver = pipeline(
                "text2text-generation",
                model="t5-base"
            )
        except Exception as e:
            self.logger.warning(f"Modèle d'amélioration non disponible: {e}")
            self.text_improver = None
    
    def optimize_for_seo(self, text: str, target_keywords: List[str]) -> str:
        """Optimise le contenu pour le SEO"""        
        optimized_text = text
        
        # Insertion de mots-clés de manière naturelle
        for keyword in target_keywords[:5]:  # Limite à 5 mots-clés
            if keyword.lower() not in text.lower():
                # Tentative d'insertion naturelle
                sentences = sent_tokenize(optimized_text)
                if sentences:
                    # Insertion dans la première phrase si possible
                    first_sentence = sentences[0]
                    if len(first_sentence.split()) > 5:
                        words = first_sentence.split()
                        insert_pos = len(words) // 2
                        words.insert(insert_pos, keyword)
                        sentences[0] = ' '.join(words)
                        optimized_text = ' '.join(sentences)
        
        return optimized_text
    
    def improve_readability(self, text: str) -> str:
        """Améliore la lisibilité du texte"""        
        # Simplification des phrases longues
        sentences = sent_tokenize(text)
        improved_sentences = []
        
        for sentence in sentences:
            words = word_tokenize(sentence)
            if len(words) > 25:  # Phrase trop longue
                # Tentative de division
                midpoint = len(words) // 2
                part1 = ' '.join(words[:midpoint]) + '.'
                part2 = ' '.join(words[midpoint:])
                improved_sentences.extend([part1, part2])
            else:
                improved_sentences.append(sentence)
        
        return ' '.join(improved_sentences)
    
    def add_structure(self, text: str) -> str:
        """Ajoute une structure au texte"""        
        paragraphs = text.split('\n\n')
        if len(paragraphs) < 2:
            return text
        
        structured_text = ""
        
        # Ajout d'un titre principal si absent
        if not re.match(r'^#+ ', text):
            structured_text += "# Main Title\n\n"
        
        # Ajout de sous-titres entre les paragraphes
        for i, paragraph in enumerate(paragraphs):
            if i > 0 and len(paragraph.strip()) > 100:
                structured_text += f"\n## Section {i}\n\n"
            
            structured_text += paragraph + "\n\n"
        
        return structured_text.strip()

class DocumentTransformer:
    """Transformateur de document principal pour créateurs de contenu"""    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.file_manager = FileManager()
        self.analyzer = DocumentAnalyzer()
        self.optimizer = ContentOptimizer()
        
        # Templates par type de créateur
        self.creator_templates = {
            'blogger': {
                'format': 'md',
                'structure': 'blog_post',
                'seo_optimization': True,
                'readability_target': 'standard'
            },
            'content_writer': {
                'format': 'html',
                'structure': 'article',
                'seo_optimization': True,
                'readability_target': 'easy'
            },
            'author': {
                'format': 'docx',
                'structure': 'book',
                'seo_optimization': False,
                'readability_target': 'standard'
            },
            'marketer': {
                'format': 'html',
                'structure': 'marketing',
                'seo_optimization': True,
                'readability_target': 'easy'
            }
        }
    
    def transform(
        self,
        input_path: str,
        config: 'TransformationConfig',
        output_path: Optional[str] = None
    ) -> 'TransformationResult':
        """Transformation de document selon configuration"""        
        start_time = time.time()
        operations = []
        warnings = []
        errors = []
        
        try:
            # Validation du fichier d'entrée
            if not validate_document_file(input_path):
                raise ValidationError(f"Fichier document invalide: {input_path}")
            
            # Analyse du fichier source
            original_metadata = self.analyzer.analyze_document_file(input_path)
            operations.append("Analyse métadonnées")
            
            # Extraction du texte complet
            full_text = self.analyzer._extract_text_from_file(input_path)
            operations.append("Extraction texte")
            
            # Préparation du chemin de sortie
            if not output_path:
                output_path = self._generate_output_path(input_path, config)
            
            # Application des transformations selon le type
            processed_text = full_text
            
            if config.type.value == 'document_convert':
                processed_text = self._convert_document(processed_text, config.parameters)
                operations.append("Conversion format")
                
            elif config.type.value == 'document_extract_text':
                # Texte déjà extrait
                operations.append("Extraction texte brut")
                
            elif config.type.value == 'document_summarize':
                processed_text = self._summarize_document(processed_text, config.parameters)
                operations.append("Résumé automatique")
            
            # Optimisations supplémentaires
            creator_type = config.parameters.get('creator_type')
            if creator_type and creator_type in self.creator_templates:
                template = self.creator_templates[creator_type]
                
                if template.get('seo_optimization'):
                    target_keywords = config.parameters.get('keywords', [])
                    processed_text = self.optimizer.optimize_for_seo(processed_text, target_keywords)
                    operations.append("Optimisation SEO")
                
                if template.get('readability_target') == 'easy':
                    processed_text = self.optimizer.improve_readability(processed_text)
                    operations.append("Amélioration lisibilité")
            
            # Sauvegarde du document traité
            self._save_document(processed_text, output_path, config)
            operations.append("Sauvegarde")
            
            # Analyse finale
            processed_metadata = self.analyzer.analyze_document_file(output_path)
            
            # Calcul des métriques de qualité
            quality_metrics = self._calculate_quality_metrics(
                original_metadata, processed_metadata
            )
            
            processing_time = time.time() - start_time
            
            from . import TransformationResult, TransformationType
            return TransformationResult(
                success=True,
                input_path=input_path,
                output_path=output_path,
                transformation_type=TransformationType(config.type.value),
                metadata={
                    'original': original_metadata.__dict__,
                    'processed': processed_metadata.__dict__,
                    'quality_metrics': quality_metrics.__dict__
                },
                errors=errors,
                warnings=warnings,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Erreur transformation document {input_path}: {e}")
            processing_time = time.time() - start_time
            
            from . import TransformationResult, TransformationType
            return TransformationResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transformation_type=TransformationType(config.type.value),
                metadata={},
                errors=[str(e)],
                warnings=warnings,
                processing_time=processing_time
            )
    
    def _convert_document(self, text: str, params: Dict[str, Any]) -> str:
        """Convertit le document vers un format spécifié"""        
        target_format = params.get('format', 'txt')
        
        if target_format == 'md':
            return self._convert_to_markdown(text)
        elif target_format == 'html':
            return self._convert_to_html(text)
        elif target_format == 'txt':
            return self._convert_to_text(text)
        else:
            return text
    
    def _convert_to_markdown(self, text: str) -> str:
        """Convertit vers Markdown"""        
        # Détection et conversion des titres
        lines = text.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append('')
                continue
            
            # Détection de titre (ligne en majuscules ou pattern spécifique)
            if (line.isupper() and len(line) > 5) or line.endswith(':'):
                if len(line) < 50:  # Titre probable
                    markdown_lines.append(f"## {line.rstrip(':')}")
                else:
                    markdown_lines.append(line)
            else:
                markdown_lines.append(line)
        
        return '\n'.join(markdown_lines)
    
    def _convert_to_html(self, text: str) -> str:
        """Convertit vers HTML"""        
        # Conversion basique avec structure HTML
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document</title>
</head>
<body>
"""        
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Détection de titre
                if len(paragraph) < 100 and paragraph.count('\n') == 0:
                    html_content += f"    <h2>{paragraph}</h2>\n"
                else:
                    html_content += f"    <p>{paragraph.replace(chr(10), '<br>')}</p>\n"
        
        html_content += """</body>
</html>"""        
        return html_content
    
    def _convert_to_text(self, text: str) -> str:
        """Nettoie pour format texte pur"""        
        # Suppression des formatages spéciaux
        cleaned = re.sub(r'<[^>]+>', '', text)  # HTML tags
        cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned)  # Bold markdown
        cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)  # Italic markdown
        cleaned = re.sub(r'`([^`]+)`', r'\1', cleaned)  # Code markdown
        
        return cleaned
    
    def _summarize_document(self, text: str, params: Dict[str, Any]) -> str:
        """Crée un résumé automatique du document"""        
        max_length = params.get('max_length', 500)
        summary_type = params.get('type', 'extractive')
        
        if summary_type == 'abstractive' and self.analyzer.summarizer:
            try:
                # Résumé abstractif avec modèle de langue
                chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
                summaries = []
                
                for chunk in chunks[:5]:  # Limite à 5 chunks
                    if len(chunk.strip()) > 100:
                        summary = self.analyzer.summarizer(
                            chunk,
                            max_length=min(150, len(chunk)//4),
                            min_length=50,
                            do_sample=False
                        )
                        if summary:
                            summaries.append(summary[0]['summary_text'])
                
                return ' '.join(summaries)[:max_length]
            except Exception as e:
                self.logger.warning(f"Résumé abstractif échoué: {e}")
        
        # Fallback: résumé extractif
        return self._extractive_summarize(text, max_length)
    
    def _extractive_summarize(self, text: str, max_length: int) -> str:
        """Crée un résumé extractif en sélectionnant les meilleures phrases"""        
        sentences = sent_tokenize(text)
        if len(sentences) <= 3:
            return text[:max_length]
        
        # Calcul des scores TF-IDF pour chaque phrase
        try:
            tfidf = TfidfVectorizer(stop_words='english', max_features=100)
            tfidf_matrix = tfidf.fit_transform(sentences)
            
            # Score de chaque phrase (somme des TF-IDF des mots)
            sentence_scores = []
            for i, sentence in enumerate(sentences):
                score = tfidf_matrix[i].sum()
                sentence_scores.append((score, sentence))
            
            # Tri par score décroissant
            sentence_scores.sort(reverse=True)
            
            # Sélection des meilleures phrases
            summary_sentences = []
            current_length = 0
            
            for score, sentence in sentence_scores:
                if current_length + len(sentence) <= max_length:
                    summary_sentences.append(sentence)
                    current_length += len(sentence)
                else:
                    break
            
            return ' '.join(summary_sentences)
        except Exception:
            # Fallback simple: premières phrases
            summary = []
            current_length = 0
            
            for sentence in sentences:
                if current_length + len(sentence) <= max_length:
                    summary.append(sentence)
                    current_length += len(sentence)
                else:
                    break
            
            return ' '.join(summary)
    
    def _save_document(
        self,
        text: str,
        output_path: str,
        config: 'TransformationConfig'
    ) -> None:
        """Sauvegarde le document traité"""        
        # Création du répertoire si nécessaire
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        output_format = config.output_format or Path(output_path).suffix.lstrip('.').lower()
        
        if output_format == 'docx':
            self._save_as_docx(text, output_path)
        elif output_format == 'pdf':
            self._save_as_pdf(text, output_path)
        elif output_format in ['md', 'html', 'txt']:
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
        else:
            # Format texte par défaut
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
    
    def _save_as_docx(self, text: str, output_path: str) -> None:
        """Sauvegarde en format DOCX"""        
        try:
            doc = DocxDocument()
            
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)
            
            doc.save(output_path)
        except Exception as e:
            raise DocumentProcessingError(f"Échec sauvegarde DOCX: {str(e)}")
    
    def _save_as_pdf(self, text: str, output_path: str) -> None:
        """Sauvegarde en format PDF"""        
        try:
            # Utilisation de reportlab pour génération PDF
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.utils import simpleSplit
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Configuration du texte
            lines = text.split('\n')
            y_position = height - 50
            
            for line in lines:
                if y_position < 50:  # Nouvelle page
                    c.showPage()
                    y_position = height - 50
                
                # Division des lignes trop longues
                wrapped_lines = simpleSplit(line, 'Helvetica', 12, width - 100)
                
                for wrapped_line in wrapped_lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    
                    c.drawString(50, y_position, wrapped_line)
                    y_position -= 15
            
            c.save()
        except Exception as e:
            # Fallback: sauvegarde en texte
            self.logger.warning(f"Échec PDF, sauvegarde en texte: {e}")
            with open(output_path.replace('.pdf', '.txt'), 'w', encoding='utf-8') as file:
                file.write(text)
    
    def _generate_output_path(self, input_path: str, config: 'TransformationConfig') -> str:
        """Génère le chemin de sortie automatiquement"""        
        input_path_obj = Path(input_path)
        
        # Détermination de l'extension selon la transformation
        output_format = config.output_format
        
        # Application du template créateur si spécifié
        creator_type = config.parameters.get('creator_type')
        if creator_type and creator_type in self.creator_templates:
            template = self.creator_templates[creator_type]
            if 'format' in template:
                output_format = template['format']
        
        if not output_format:
            if config.type.value == 'document_convert':
                output_format = config.parameters.get('format', 'txt')
            else:
                output_format = input_path_obj.suffix.lstrip('.')
        
        # Nom de fichier avec suffixe de transformation
        transform_suffix = config.type.value.replace('document_', '')
        new_name = f"{input_path_obj.stem}_{transform_suffix}.{output_format}"
        
        return str(input_path_obj.parent / new_name)
    
    def _calculate_quality_metrics(
        self,
        original: DocumentMetadata,
        processed: DocumentMetadata
    ) -> DocumentQualityMetrics:
        """Calcule les métriques de qualité de la transformation"""        
        # Comparaison des métriques de contenu
        word_count_ratio = processed.word_count / original.word_count if original.word_count > 0 else 1.0
        readability_improvement = processed.readability_score - original.readability_score
        seo_improvement = processed.seo_score - original.seo_score
        
        # Score de préservation du contenu
        content_preservation = min(1.0, word_count_ratio) if word_count_ratio <= 1.0 else 1.0 / word_count_ratio
        
        # Score global
        overall_quality = (content_preservation * 0.4 + 
                          (processed.readability_score) * 0.3 + 
                          (processed.seo_score) * 0.3)
        
        return DocumentQualityMetrics(
            content_preservation_score=content_preservation,
            readability_improvement=max(0.0, readability_improvement),
            seo_improvement=max(0.0, seo_improvement),
            structure_quality=1.0,  # Estimation par défaut
            keyword_optimization=processed.seo_score,
            language_quality=processed.readability_score,
            overall_quality_score=max(0.0, min(1.0, overall_quality))
        )

class AsyncDocumentTransformer:
    """Version asynchrone du transformateur de document"""    
    def __init__(self):
        self.sync_transformer = DocumentTransformer()
        self.logger = logging.getLogger(__name__)
    
    async def transform_async(
        self,
        input_path: str,
        config: 'TransformationConfig',
        output_path: Optional[str] = None
    ) -> 'TransformationResult':
        """Transformation de document asynchrone"""        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.sync_transformer.transform,
            input_path,
            config,
            output_path
        )
    
    async def transform_batch_async(
        self,
        inputs: List[Tuple[str, 'TransformationConfig']],
        max_concurrent: int = 4
    ) -> List['TransformationResult']:
        """Transformation en lot asynchrone"""        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def transform_single(input_config_tuple):
            async with semaphore:
                input_path, config = input_config_tuple
                return await self.transform_async(input_path, config)
        
        tasks = [transform_single(item) for item in inputs]
        return await asyncio.gather(*tasks, return_exceptions=True)

# Export des classes
__all__ = [
    'DocumentTransformer',
    'AsyncDocumentTransformer',
    'DocumentAnalyzer',
    'ContentOptimizer',
    'DocumentFormat',
    'DocumentType',
    'ContentQuality',
    'TextAnalysis',
    'DocumentProcessingResult'
]
