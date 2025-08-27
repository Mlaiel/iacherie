"""
🏷️ Metadata Transformation Engine - IA Influencer Agent Platform Enterprise
=========================================================================
Module: backend/data_management/transformers/metadata_transformer.py
Author: Fahed Mlaiel (mlaiel@live.de)
=========================================================================

⚠️  PROPRIÉTÉ INTELLECTUELLE EXCLUSIVE - FAHED MLAIEL ⚠️
© 2025 Fahed Mlaiel. Tous droits réservés.
Contact: mlaiel@live.de

AVERTISSEMENT: Toute tentative de vol, copie ou utilisation non autorisée
de ce code ou de cette technologie est strictement interdite et sera
poursuivie selon les lois allemandes et internationales.
"""

import asyncio
import logging
import time
import json
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
from datetime import datetime, timezone
import uuid

from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import mutagen
from mutagen.id3 import ID3, TIT2, TPE1, TALB, TDRC, TCON, COMM, TXXX
import pypdf
from pypdf import PdfReader
import cv2
import numpy as np
from exif import Image as ExifImage
import xml.etree.ElementTree as ET
from pathvalidate import sanitize_filename

from ..models.metadata_models import MetadataStandard, MetadataRecord, EnrichmentResult
from ...core.exceptions import MetadataProcessingError, ValidationError
from ...core.config import get_settings
from ...utils.file_manager import FileManager
from ...utils.validation import validate_file_path

settings = get_settings()
logger = logging.getLogger(__name__)

class MetadataStandard(Enum):
    """Standards de métadonnées supportés"""
    DUBLIN_CORE = "dublin_core"
    ID3V2 = "id3v2"
    EXIF = "exif"
    XMP = "xmp"
    IPTC = "iptc"
    VORBIS_COMMENT = "vorbis_comment"
    PDF_INFO = "pdf_info"
    CUSTOM = "custom"

class MetadataField(Enum):
    """Champs de métadonnées standardisés"""
    # Identification
    TITLE = "title"
    CREATOR = "creator"
    AUTHOR = "author"
    ARTIST = "artist"
    ALBUM = "album"
    
    # Dates
    DATE_CREATED = "date_created"
    DATE_MODIFIED = "date_modified"
    DATE_PUBLISHED = "date_published"
    
    # Description
    DESCRIPTION = "description"
    SUBJECT = "subject"
    KEYWORDS = "keywords"
    TAGS = "tags"
    
    # Technique
    FORMAT = "format"
    FILE_SIZE = "file_size"
    DURATION = "duration"
    RESOLUTION = "resolution"
    
    # Droits
    COPYRIGHT = "copyright"
    LICENSE = "license"
    RIGHTS_HOLDER = "rights_holder"
    
    # Géolocalisation
    GPS_LATITUDE = "gps_latitude"
    GPS_LONGITUDE = "gps_longitude"
    LOCATION = "location"
    
    # Créateur spécifique
    CREATOR_TYPE = "creator_type"
    PLATFORM = "platform"
    SOCIAL_MEDIA_HANDLE = "social_media_handle"

class EnrichmentType(Enum):
    """Types d'enrichissement de métadonnées"""
    AI_TAGGING = "ai_tagging"
    GEO_ENRICHMENT = "geo_enrichment"
    CONTENT_ANALYSIS = "content_analysis"
    FACE_RECOGNITION = "face_recognition"
    OBJECT_DETECTION = "object_detection"
    TEXT_EXTRACTION = "text_extraction"
    MUSIC_ANALYSIS = "music_analysis"
    COPYRIGHT_DETECTION = "copyright_detection"

@dataclass
class MetadataProcessingResult:
    """Résultat du traitement de métadonnées"""
    success: bool
    input_file: str
    output_file: Optional[str]
    original_metadata: Dict[str, Any]
    processed_metadata: Dict[str, Any]
    enrichment_results: List[EnrichmentResult]
    processing_time: float
    operations_performed: List[str]
    warnings: List[str]
    errors: List[str]

class MetadataExtractor:
    """Extracteur de métadonnées multi-format pour créateurs de contenu"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.file_manager = FileManager()
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrait toutes les métadonnées disponibles d'un fichier"""
        
        try:
            file_path_obj = Path(file_path)
            file_extension = file_path_obj.suffix.lower()
            
            # Métadonnées de base du système de fichiers
            metadata = self._extract_filesystem_metadata(file_path)
            
            # Extraction spécialisée selon le type de fichier
            if file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                metadata.update(self._extract_image_metadata(file_path))
            elif file_extension in ['.mp3', '.wav', '.flac', '.ogg', '.m4a']:
                metadata.update(self._extract_audio_metadata(file_path))
            elif file_extension in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
                metadata.update(self._extract_video_metadata(file_path))
            elif file_extension == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_extension in ['.docx', '.doc']:
                metadata.update(self._extract_document_metadata(file_path))
            
            # Ajout d'un hash unique pour l'identification
            metadata['content_hash'] = self._calculate_content_hash(file_path)
            metadata['extraction_timestamp'] = datetime.now(timezone.utc).isoformat()
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Erreur extraction métadonnées {file_path}: {e}")
            raise MetadataProcessingError(f"Échec extraction métadonnées: {str(e)}")
    
    def _extract_filesystem_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées du système de fichiers"""
        
        file_stats = Path(file_path).stat()
        
        return {
            'filename': Path(file_path).name,
            'file_extension': Path(file_path).suffix.lower(),
            'file_size': file_stats.st_size,
            'creation_time': datetime.fromtimestamp(file_stats.st_ctime, timezone.utc).isoformat(),
            'modification_time': datetime.fromtimestamp(file_stats.st_mtime, timezone.utc).isoformat(),
            'access_time': datetime.fromtimestamp(file_stats.st_atime, timezone.utc).isoformat(),
            'file_mode': oct(file_stats.st_mode),
            'file_path': str(file_path)
        }
    
    def _extract_image_metadata(self, image_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées d'image (EXIF, IPTC, XMP)"""
        
        metadata = {}
        
        try:
            # Extraction EXIF avec PIL
            with Image.open(image_path) as img:
                metadata['image_width'] = img.width
                metadata['image_height'] = img.height
                metadata['image_mode'] = img.mode
                metadata['image_format'] = img.format
                
                # Métadonnées EXIF
                if hasattr(img, '_getexif') and img._getexif() is not None:
                    exif_data = {}
                    for tag_id, value in img._getexif().items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_data[tag] = str(value)
                    
                    metadata['exif'] = exif_data
                    
                    # Extraction GPS si disponible
                    gps_info = self._extract_gps_info(img._getexif())
                    if gps_info:
                        metadata['gps'] = gps_info
            
            # Extraction avec exif library pour plus de détails
            try:
                with open(image_path, 'rb') as f:
                    exif_img = ExifImage(f)
                    if exif_img.has_exif:
                        detailed_exif = {}
                        for attr in dir(exif_img):
                            if not attr.startswith('_') and hasattr(exif_img, attr):
                                try:
                                    value = getattr(exif_img, attr)
                                    if not callable(value):
                                        detailed_exif[attr] = str(value)
                                except:
                                    continue
                        metadata['detailed_exif'] = detailed_exif
            except Exception as e:
                self.logger.debug(f"Extraction EXIF détaillée échouée: {e}")
            
        except Exception as e:
            self.logger.warning(f"Erreur extraction métadonnées image: {e}")
        
        return metadata
    
    def _extract_gps_info(self, exif_data: Dict) -> Optional[Dict[str, float]]:
        """Extrait les informations GPS des données EXIF"""
        
        try:
            gps_info = exif_data.get(34853)  # GPS tag
            if not gps_info:
                return None
            
            def convert_to_degrees(value):
                """Convertit les coordonnées GPS en degrés décimaux"""
                d, m, s = value
                return d + (m / 60.0) + (s / 3600.0)
            
            gps_data = {}
            
            if 2 in gps_info and 4 in gps_info:  # Latitude et Longitude
                lat = convert_to_degrees(gps_info[2])
                if gps_info[1] == 'S':
                    lat = -lat
                
                lon = convert_to_degrees(gps_info[4])
                if gps_info[3] == 'W':
                    lon = -lon
                
                gps_data['latitude'] = lat
                gps_data['longitude'] = lon
            
            if 6 in gps_info:  # Altitude
                gps_data['altitude'] = float(gps_info[6])
            
            return gps_data if gps_data else None
            
        except Exception:
            return None
    
    def _extract_audio_metadata(self, audio_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées audio (ID3, Vorbis, etc.)"""
        
        metadata = {}
        
        try:
            # Utilisation de mutagen pour l'extraction
            audio_file = mutagen.File(audio_path)
            
            if audio_file is not None:
                # Informations techniques
                if hasattr(audio_file, 'info'):
                    info = audio_file.info
                    metadata['duration'] = getattr(info, 'length', 0)
                    metadata['bitrate'] = getattr(info, 'bitrate', 0)
                    metadata['sample_rate'] = getattr(info, 'sample_rate', 0)
                    metadata['channels'] = getattr(info, 'channels', 0)
                
                # Tags métadonnées
                if audio_file.tags:
                    tags = {}
                    for key, value in audio_file.tags.items():
                        if isinstance(value, list):
                            tags[key] = [str(v) for v in value]
                        else:
                            tags[key] = str(value)
                    metadata['tags'] = tags
                    
                    # Extraction des champs standardisés
                    self._extract_standard_audio_fields(metadata, audio_file.tags)
            
        except Exception as e:
            self.logger.warning(f"Erreur extraction métadonnées audio: {e}")
        
        return metadata
    
    def _extract_standard_audio_fields(self, metadata: Dict, tags: Any) -> None:
        """Extrait les champs audio standardisés"""
        
        # Mapping des tags communs
        field_mapping = {
            'title': ['TIT2', 'TITLE', '\xa9nam'],
            'artist': ['TPE1', 'ARTIST', '\xa9ART'],
            'album': ['TALB', 'ALBUM', '\xa9alb'],
            'date': ['TDRC', 'DATE', '\xa9day'],
            'genre': ['TCON', 'GENRE', '\xa9gen'],
            'track': ['TRCK', 'TRACKNUMBER', 'trkn'],
            'albumartist': ['TPE2', 'ALBUMARTIST', 'aART']
        }
        
        for field, possible_tags in field_mapping.items():
            for tag in possible_tags:
                if tag in tags:
                    value = tags[tag]
                    if isinstance(value, list):
                        metadata[field] = str(value[0]) if value else ''
                    else:
                        metadata[field] = str(value)
                    break
    
    def _extract_video_metadata(self, video_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées vidéo"""
        
        metadata = {}
        
        try:
            # Utilisation d'OpenCV pour les métadonnées de base
            cap = cv2.VideoCapture(video_path)
            
            if cap.isOpened():
                metadata['video_width'] = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                metadata['video_height'] = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                metadata['fps'] = cap.get(cv2.CAP_PROP_FPS)
                metadata['frame_count'] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                
                duration = metadata['frame_count'] / metadata['fps'] if metadata['fps'] > 0 else 0
                metadata['duration'] = duration
                
                cap.release()
            
            # Tentative d'extraction avec mutagen (pour conteneurs avec audio)
            try:
                video_file = mutagen.File(video_path)
                if video_file and video_file.tags:
                    metadata['video_tags'] = {k: str(v) for k, v in video_file.tags.items()}
            except:
                pass
            
        except Exception as e:
            self.logger.warning(f"Erreur extraction métadonnées vidéo: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées PDF"""
        
        metadata = {}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Informations de base
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Métadonnées du document
                if pdf_reader.metadata:
                    pdf_metadata = {}
                    for key, value in pdf_reader.metadata.items():
                        if key.startswith('/'):
                            clean_key = key[1:]  # Supprimer le '/' initial
                            pdf_metadata[clean_key] = str(value)
                    metadata['pdf_metadata'] = pdf_metadata
                
                # Extraction de texte des premières pages pour analyse
                if pdf_reader.pages:
                    sample_text = ""
                    for i, page in enumerate(pdf_reader.pages[:3]):  # 3 premières pages
                        sample_text += page.extract_text()
                        if len(sample_text) > 1000:  # Limitation
                            break
                    
                    if sample_text.strip():
                        metadata['sample_text'] = sample_text[:1000]
                        metadata['estimated_word_count'] = len(sample_text.split())
            
        except Exception as e:
            self.logger.warning(f"Erreur extraction métadonnées PDF: {e}")
        
        return metadata
    
    def _extract_document_metadata(self, doc_path: str) -> Dict[str, Any]:
        """Extrait les métadonnées de document Office"""
        
        metadata = {}
        
        try:
            file_extension = Path(doc_path).suffix.lower()
            
            if file_extension == '.docx':
                from docx import Document
                doc = Document(doc_path)
                
                # Propriétés du document
                core_props = doc.core_properties
                metadata['document_title'] = core_props.title or ''
                metadata['document_author'] = core_props.author or ''
                metadata['document_subject'] = core_props.subject or ''
                metadata['document_keywords'] = core_props.keywords or ''
                metadata['document_created'] = core_props.created.isoformat() if core_props.created else ''
                metadata['document_modified'] = core_props.modified.isoformat() if core_props.modified else ''
                
                # Statistiques du contenu
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                metadata['word_count'] = len(text_content.split())
                metadata['paragraph_count'] = len(doc.paragraphs)
                metadata['table_count'] = len(doc.tables)
            
        except Exception as e:
            self.logger.warning(f"Erreur extraction métadonnées document: {e}")
        
        return metadata
    
    def _calculate_content_hash(self, file_path: str) -> str:
        """Calcule un hash SHA-256 du contenu du fichier"""
        
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            self.logger.warning(f"Erreur calcul hash: {e}")
            return ""

class MetadataEnricher:
    """Enrichisseur de métadonnées IA pour créateurs de contenu"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Chargement des modèles IA si disponibles
        try:
            from ultralytics import YOLO
            self.yolo_model = YOLO('yolov8n.pt')
        except:
            self.yolo_model = None
            
        try:
            import face_recognition
            self.face_recognition_available = True
        except:
            self.face_recognition_available = False
    
    def enrich_metadata(
        self,
        file_path: str,
        metadata: Dict[str, Any],
        enrichment_types: List[EnrichmentType]
    ) -> Dict[str, Any]:
        """Enrichit les métadonnées avec des analyses IA"""
        
        enriched_metadata = metadata.copy()
        enrichment_results = []
        
        for enrichment_type in enrichment_types:
            try:
                if enrichment_type == EnrichmentType.AI_TAGGING:
                    result = self._ai_tag_content(file_path, metadata)
                elif enrichment_type == EnrichmentType.CONTENT_ANALYSIS:
                    result = self._analyze_content(file_path, metadata)
                elif enrichment_type == EnrichmentType.FACE_RECOGNITION:
                    result = self._detect_faces(file_path, metadata)
                elif enrichment_type == EnrichmentType.OBJECT_DETECTION:
                    result = self._detect_objects(file_path, metadata)
                elif enrichment_type == EnrichmentType.GEO_ENRICHMENT:
                    result = self._enrich_geolocation(metadata)
                elif enrichment_type == EnrichmentType.COPYRIGHT_DETECTION:
                    result = self._detect_copyright_info(file_path, metadata)
                else:
                    continue
                
                if result:
                    enriched_metadata.update(result.get('metadata', {}))
                    enrichment_results.append(result)
                    
            except Exception as e:
                self.logger.warning(f"Erreur enrichissement {enrichment_type}: {e}")
        
        enriched_metadata['enrichment_results'] = enrichment_results
        enriched_metadata['enrichment_timestamp'] = datetime.now(timezone.utc).isoformat()
        
        return enriched_metadata
    
    def _ai_tag_content(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Génère des tags automatiques basés sur le contenu"""
        
        file_extension = Path(file_path).suffix.lower()
        tags = []
        confidence_scores = {}
        
        try:
            if file_extension in ['.jpg', '.jpeg', '.png']:
                # Tags d'image
                if self.yolo_model:
                    from PIL import Image
                    img = Image.open(file_path)
                    results = self.yolo_model(np.array(img))
                    
                    for result in results:
                        if result.boxes is not None:
                            for box in result.boxes:
                                class_name = result.names[int(box.cls)]
                                confidence = float(box.conf)
                                
                                if confidence > 0.5:  # Seuil de confiance
                                    tags.append(class_name)
                                    confidence_scores[class_name] = confidence
                
                # Tags basés sur les couleurs dominantes
                color_tags = self._generate_color_tags(file_path)
                tags.extend(color_tags)
            
            elif file_extension in ['.mp3', '.wav', '.flac']:
                # Tags audio basés sur les métadonnées existantes
                if 'genre' in metadata:
                    tags.append(metadata['genre'])
                if 'tags' in metadata:
                    existing_tags = metadata['tags']
                    for key, value in existing_tags.items():
                        if 'genre' in key.lower() or 'style' in key.lower():
                            if isinstance(value, list):
                                tags.extend([str(v) for v in value])
                            else:
                                tags.append(str(value))
            
            # Nettoyage et déduplication des tags
            tags = list(set([tag.lower().strip() for tag in tags if tag and len(tag) > 2]))
            
            return {
                'type': EnrichmentType.AI_TAGGING,
                'metadata': {
                    'ai_generated_tags': tags,
                    'tag_confidence_scores': confidence_scores
                },
                'confidence': np.mean(list(confidence_scores.values())) if confidence_scores else 0.5,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur génération tags IA: {e}")
            return {}
    
    def _generate_color_tags(self, image_path: str) -> List[str]:
        """Génère des tags basés sur les couleurs dominantes"""
        
        try:
            from PIL import Image
            import colorsys
            
            img = Image.open(image_path)
            img = img.convert('RGB')
            img = img.resize((150, 150))  # Réduction pour performance
            
            # Extraction des couleurs dominantes
            colors = img.getcolors(maxcolors=256*256*256)
            if not colors:
                return []
            
            # Tri par fréquence
            colors.sort(key=lambda x: x[0], reverse=True)
            
            color_tags = []
            for count, (r, g, b) in colors[:5]:  # Top 5 couleurs
                # Conversion en HSV pour classification
                h, s, v = colorsys.rgb_to_hsv(r/255, g/255, b/255)
                
                # Classification par teinte
                if s < 0.1:  # Couleurs désaturées
                    if v > 0.8:
                        color_tags.append('white')
                    elif v < 0.2:
                        color_tags.append('black')
                    else:
                        color_tags.append('gray')
                else:  # Couleurs saturées
                    hue_deg = h * 360
                    if hue_deg < 15 or hue_deg > 345:
                        color_tags.append('red')
                    elif 15 <= hue_deg < 45:
                        color_tags.append('orange')
                    elif 45 <= hue_deg < 75:
                        color_tags.append('yellow')
                    elif 75 <= hue_deg < 150:
                        color_tags.append('green')
                    elif 150 <= hue_deg < 210:
                        color_tags.append('cyan')
                    elif 210 <= hue_deg < 270:
                        color_tags.append('blue')
                    elif 270 <= hue_deg < 330:
                        color_tags.append('purple')
                    else:
                        color_tags.append('pink')
            
            return list(set(color_tags))
            
        except Exception as e:
            self.logger.warning(f"Erreur génération tags couleur: {e}")
            return []
    
    def _analyze_content(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Analyse approfondie du contenu"""
        
        file_extension = Path(file_path).suffix.lower()
        analysis_results = {}
        
        try:
            if file_extension in ['.jpg', '.jpeg', '.png']:
                # Analyse d'image
                analysis_results.update(self._analyze_image_content(file_path))
            elif file_extension in ['.mp3', '.wav', '.flac']:
                # Analyse audio
                analysis_results.update(self._analyze_audio_content(file_path, metadata))
            elif file_extension in ['.mp4', '.avi', '.mov']:
                # Analyse vidéo
                analysis_results.update(self._analyze_video_content(file_path))
            
            return {
                'type': EnrichmentType.CONTENT_ANALYSIS,
                'metadata': {
                    'content_analysis': analysis_results
                },
                'confidence': 0.8,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur analyse contenu: {e}")
            return {}
    
    def _analyze_image_content(self, image_path: str) -> Dict[str, Any]:
        """Analyse spécialisée d'image"""
        
        analysis = {}
        
        try:
            from PIL import Image
            import cv2
            
            # Chargement de l'image
            img_pil = Image.open(image_path)
            img_cv = cv2.imread(image_path)
            
            # Analyse de la composition
            width, height = img_pil.size
            analysis['aspect_ratio'] = round(width / height, 2)
            analysis['resolution_category'] = self._categorize_resolution(width, height)
            
            # Analyse de la luminosité et du contraste
            img_gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
            analysis['brightness'] = float(np.mean(img_gray) / 255.0)
            analysis['contrast'] = float(np.std(img_gray) / 255.0)
            
            # Détection de netteté (variance du Laplacien)
            laplacian_var = cv2.Laplacian(img_gray, cv2.CV_64F).var()
            analysis['sharpness_score'] = float(min(1.0, laplacian_var / 1000.0))
            
            # Classification de la qualité
            if analysis['sharpness_score'] > 0.5 and analysis['contrast'] > 0.1:
                analysis['quality_assessment'] = 'high'
            elif analysis['sharpness_score'] > 0.2:
                analysis['quality_assessment'] = 'medium'
            else:
                analysis['quality_assessment'] = 'low'
            
        except Exception as e:
            self.logger.warning(f"Erreur analyse image: {e}")
        
        return analysis
    
    def _categorize_resolution(self, width: int, height: int) -> str:
        """Catégorise la résolution de l'image"""
        
        total_pixels = width * height
        
        if total_pixels >= 8000000:  # 8MP+
            return 'ultra_high'
        elif total_pixels >= 2000000:  # 2MP+
            return 'high'
        elif total_pixels >= 500000:   # 0.5MP+
            return 'medium'
        else:
            return 'low'
    
    def _analyze_audio_content(self, audio_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Analyse spécialisée audio"""
        
        analysis = {}
        
        try:
            # Analyse basée sur les métadonnées existantes
            if 'duration' in metadata:
                duration = metadata['duration']
                if duration < 60:
                    analysis['content_type'] = 'short_form'
                elif duration < 300:
                    analysis['content_type'] = 'medium_form'
                else:
                    analysis['content_type'] = 'long_form'
            
            if 'bitrate' in metadata:
                bitrate = metadata['bitrate']
                if bitrate >= 320:
                    analysis['quality_tier'] = 'high'
                elif bitrate >= 192:
                    analysis['quality_tier'] = 'medium'
                else:
                    analysis['quality_tier'] = 'low'
            
            # Analyse du genre si disponible
            if 'genre' in metadata:
                genre = metadata['genre'].lower()
                if any(term in genre for term in ['electronic', 'edm', 'techno', 'house']):
                    analysis['music_category'] = 'electronic'
                elif any(term in genre for term in ['rock', 'metal', 'punk']):
                    analysis['music_category'] = 'rock'
                elif any(term in genre for term in ['jazz', 'blues', 'soul']):
                    analysis['music_category'] = 'jazz_blues'
                elif any(term in genre for term in ['classical', 'orchestra', 'symphony']):
                    analysis['music_category'] = 'classical'
                else:
                    analysis['music_category'] = 'other'
            
        except Exception as e:
            self.logger.warning(f"Erreur analyse audio: {e}")
        
        return analysis
    
    def _analyze_video_content(self, video_path: str) -> Dict[str, Any]:
        """Analyse spécialisée vidéo"""
        
        analysis = {}
        
        try:
            import cv2
            
            cap = cv2.VideoCapture(video_path)
            
            if cap.isOpened():
                width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                fps = cap.get(cv2.CAP_PROP_FPS)
                frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                
                duration = frame_count / fps if fps > 0 else 0
                
                # Catégorisation par résolution
                if width >= 3840:  # 4K
                    analysis['resolution_category'] = '4k'
                elif width >= 1920:  # 1080p
                    analysis['resolution_category'] = '1080p'
                elif width >= 1280:  # 720p
                    analysis['resolution_category'] = '720p'
                else:
                    analysis['resolution_category'] = 'sd'
                
                # Catégorisation par durée
                if duration < 60:
                    analysis['duration_category'] = 'short'
                elif duration < 600:
                    analysis['duration_category'] = 'medium'
                else:
                    analysis['duration_category'] = 'long'
                
                # Format d'aspect ratio
                aspect_ratio = width / height
                if abs(aspect_ratio - 16/9) < 0.1:
                    analysis['aspect_format'] = 'widescreen'
                elif abs(aspect_ratio - 4/3) < 0.1:
                    analysis['aspect_format'] = 'standard'
                elif abs(aspect_ratio - 1) < 0.1:
                    analysis['aspect_format'] = 'square'
                elif aspect_ratio < 1:
                    analysis['aspect_format'] = 'vertical'
                else:
                    analysis['aspect_format'] = 'custom'
                
                cap.release()
            
        except Exception as e:
            self.logger.warning(f"Erreur analyse vidéo: {e}")
        
        return analysis
    
    def _detect_faces(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Détection et analyse des visages"""
        
        if not self.face_recognition_available:
            return {}
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in ['.jpg', '.jpeg', '.png']:
            return {}
        
        try:
            import face_recognition
            from PIL import Image
            
            # Chargement de l'image
            img = face_recognition.load_image_file(file_path)
            
            # Détection des visages
            face_locations = face_recognition.face_locations(img)
            face_encodings = face_recognition.face_encodings(img, face_locations)
            
            faces_data = []
            for i, (face_location, face_encoding) in enumerate(zip(face_locations, face_encodings)):
                top, right, bottom, left = face_location
                
                face_info = {
                    'face_id': i,
                    'location': {
                        'top': top,
                        'right': right,
                        'bottom': bottom,
                        'left': left
                    },
                    'size': (right - left) * (bottom - top),
                    'encoding_hash': hashlib.md5(face_encoding.tobytes()).hexdigest()[:16]
                }
                faces_data.append(face_info)
            
            return {
                'type': EnrichmentType.FACE_RECOGNITION,
                'metadata': {
                    'faces_detected': len(faces_data),
                    'face_data': faces_data
                },
                'confidence': 0.9,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur détection visages: {e}")
            return {}
    
    def _detect_objects(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Détection d'objets avec YOLO"""
        
        if not self.yolo_model:
            return {}
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in ['.jpg', '.jpeg', '.png']:
            return {}
        
        try:
            from PIL import Image
            
            img = Image.open(file_path)
            results = self.yolo_model(np.array(img))
            
            objects_data = []
            for result in results:
                if result.boxes is not None:
                    for i, box in enumerate(result.boxes):
                        bbox = box.xyxy.tolist()[0]
                        
                        object_info = {
                            'object_id': i,
                            'class': result.names[int(box.cls)],
                            'confidence': float(box.conf),
                            'bbox': {
                                'x1': bbox[0],
                                'y1': bbox[1],
                                'x2': bbox[2],
                                'y2': bbox[3]
                            },
                            'area': (bbox[2] - bbox[0]) * (bbox[3] - bbox[1])
                        }
                        objects_data.append(object_info)
            
            return {
                'type': EnrichmentType.OBJECT_DETECTION,
                'metadata': {
                    'objects_detected': len(objects_data),
                    'object_data': objects_data
                },
                'confidence': 0.85,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur détection objets: {e}")
            return {}
    
    def _enrich_geolocation(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Enrichissement géographique basé sur les coordonnées GPS"""
        
        gps_data = metadata.get('gps')
        if not gps_data or 'latitude' not in gps_data or 'longitude' not in gps_data:
            return {}
        
        try:
            lat = gps_data['latitude']
            lon = gps_data['longitude']
            
            # Ici on pourrait intégrer un service de géocodage inverse
            # Pour l'exemple, on génère des informations basiques
            
            geo_enrichment = {
                'coordinates': {
                    'latitude': lat,
                    'longitude': lon,
                    'decimal_degrees': f"{lat}, {lon}"
                },
                'location_precision': 'approximate',  # Dépendrait du service de géocodage
                'timezone_estimate': self._estimate_timezone(lat, lon)
            }
            
            return {
                'type': EnrichmentType.GEO_ENRICHMENT,
                'metadata': {
                    'geo_enrichment': geo_enrichment
                },
                'confidence': 0.7,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur enrichissement géo: {e}")
            return {}
    
    def _estimate_timezone(self, lat: float, lon: float) -> str:
        """Estimation approximative du fuseau horaire"""
        
        # Estimation basique basée sur la longitude
        # (15 degrés = 1 heure)
        utc_offset = round(lon / 15)
        
        if utc_offset >= 0:
            return f"UTC+{utc_offset}"
        else:
            return f"UTC{utc_offset}"
    
    def _detect_copyright_info(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Détection d'informations de copyright"""
        
        copyright_info = {}
        
        try:
            # Recherche dans les métadonnées existantes
            for key, value in metadata.items():
                if any(term in str(key).lower() for term in ['copyright', 'rights', 'author', 'creator']):
                    copyright_info[key] = str(value)
            
            # Recherche spécifique selon le type de fichier
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension in ['.jpg', '.jpeg', '.png']:
                # Watermark detection pourrait être ajouté ici
                pass
            elif file_extension in ['.mp3', '.wav', '.flac']:
                # Recherche dans les tags audio
                if 'tags' in metadata:
                    for tag_key, tag_value in metadata['tags'].items():
                        if 'copyright' in tag_key.lower() or 'rights' in tag_key.lower():
                            copyright_info[tag_key] = str(tag_value)
            
            # Score de protection basé sur la présence d'informations
            protection_score = min(1.0, len(copyright_info) * 0.25)
            
            return {
                'type': EnrichmentType.COPYRIGHT_DETECTION,
                'metadata': {
                    'copyright_information': copyright_info,
                    'protection_score': protection_score,
                    'has_copyright_notice': len(copyright_info) > 0
                },
                'confidence': 0.8,
                'timestamp': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Erreur détection copyright: {e}")
            return {}

class MetadataNormalizer:
    """Normalisateur de métadonnées selon standards internationaux"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Mapping vers Dublin Core
        self.dublin_core_mapping = {
            'title': ['title', 'filename', 'name'],
            'creator': ['creator', 'author', 'artist', 'photographer'],
            'subject': ['subject', 'keywords', 'tags', 'genre'],
            'description': ['description', 'comment', 'abstract'],
            'publisher': ['publisher', 'label', 'company'],
            'contributor': ['contributor', 'editor', 'producer'],
            'date': ['date', 'date_created', 'creation_time'],
            'type': ['type', 'format', 'content_type'],
            'format': ['format', 'file_extension', 'mime_type'],
            'identifier': ['identifier', 'id', 'content_hash'],
            'source': ['source', 'origin', 'file_path'],
            'language': ['language', 'lang'],
            'relation': ['relation', 'reference', 'link'],
            'coverage': ['coverage', 'location', 'gps'],
            'rights': ['rights', 'copyright', 'license']
        }
    
    def normalize_to_standard(
        self,
        metadata: Dict[str, Any],
        target_standard: MetadataStandard
    ) -> Dict[str, Any]:
        """Normalise les métadonnées vers un standard spécifique"""
        
        try:
            if target_standard == MetadataStandard.DUBLIN_CORE:
                return self._normalize_to_dublin_core(metadata)
            elif target_standard == MetadataStandard.ID3V2:
                return self._normalize_to_id3v2(metadata)
            elif target_standard == MetadataStandard.EXIF:
                return self._normalize_to_exif(metadata)
            else:
                return metadata
                
        except Exception as e:
            self.logger.error(f"Erreur normalisation métadonnées: {e}")
            return metadata
    
    def _normalize_to_dublin_core(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Normalise vers Dublin Core"""
        
        dublin_core = {}
        
        for dc_field, source_fields in self.dublin_core_mapping.items():
            value = None
            
            # Recherche dans les champs sources
            for source_field in source_fields:
                if source_field in metadata:
                    value = metadata[source_field]
                    break
                
                # Recherche partielle (contient le terme)
                for meta_key in metadata.keys():
                    if source_field.lower() in meta_key.lower():
                        value = metadata[meta_key]
                        break
                
                if value:
                    break
            
            if value is not None:
                # Nettoyage et formatage
                if isinstance(value, list):
                    dublin_core[dc_field] = '; '.join(str(v) for v in value)
                else:
                    dublin_core[dc_field] = str(value)
        
        # Ajout de métadonnées techniques spécifiques
        dublin_core['metadata_standard'] = 'dublin_core'
        dublin_core['normalization_timestamp'] = datetime.now(timezone.utc).isoformat()
        
        return dublin_core
    
    def _normalize_to_id3v2(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Normalise vers ID3v2 pour audio"""
        
        id3_mapping = {
            'TIT2': ['title', 'filename'],  # Title
            'TPE1': ['artist', 'creator'],  # Artist
            'TALB': ['album'],              # Album
            'TDRC': ['date', 'year'],       # Recording time
            'TCON': ['genre'],              # Content type
            'COMM': ['comment', 'description'],  # Comments
            'TXXX': ['tags', 'keywords']    # User defined text
        }
        
        id3_metadata = {}
        
        for id3_frame, source_fields in id3_mapping.items():
            for source_field in source_fields:
                if source_field in metadata:
                    value = metadata[source_field]
                    if isinstance(value, list):
                        id3_metadata[id3_frame] = '; '.join(str(v) for v in value)
                    else:
                        id3_metadata[id3_frame] = str(value)
                    break
        
        return id3_metadata
    
    def _normalize_to_exif(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Normalise vers EXIF pour images"""
        
        exif_mapping = {
            'DateTime': ['date_created', 'creation_time'],
            'Artist': ['artist', 'creator', 'photographer'],
            'Copyright': ['copyright', 'rights'],
            'ImageDescription': ['description', 'comment'],
            'Software': ['software', 'editor'],
            'Make': ['camera_make', 'device_make'],
            'Model': ['camera_model', 'device_model']
        }
        
        exif_metadata = {}
        
        for exif_tag, source_fields in exif_mapping.items():
            for source_field in source_fields:
                if source_field in metadata:
                    exif_metadata[exif_tag] = str(metadata[source_field])
                    break
        
        return exif_metadata

class MetadataTransformer:
    """Transformateur de métadonnées principal pour créateurs de contenu"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.file_manager = FileManager()
        self.extractor = MetadataExtractor()
        self.enricher = MetadataEnricher()
        self.normalizer = MetadataNormalizer()
    
    def transform(
        self,
        input_path: str,
        config: 'TransformationConfig',
        output_path: Optional[str] = None
    ) -> 'TransformationResult':
        """Transformation de métadonnées selon configuration"""
        
        start_time = time.time()
        operations = []
        warnings = []
        errors = []
        
        try:
            # Validation du fichier d'entrée
            if not validate_file_path(input_path):
                raise ValidationError(f"Fichier invalide: {input_path}")
            
            # Extraction des métadonnées originales
            original_metadata = self.extractor.extract_metadata(input_path)
            operations.append("Extraction métadonnées")
            
            # Préparation du chemin de sortie
            if not output_path:
                output_path = self._generate_output_path(input_path, config)
            
            # Application des transformations selon le type
            processed_metadata = original_metadata.copy()
            enrichment_results = []
            
            if config.type.value == 'metadata_normalize':
                processed_metadata = self._normalize_metadata(processed_metadata, config.parameters)
                operations.append("Normalisation")
                
            elif config.type.value == 'metadata_enrich':
                enrichment_result = self._enrich_metadata(input_path, processed_metadata, config.parameters)
                processed_metadata = enrichment_result['metadata']
                enrichment_results = enrichment_result['enrichment_results']
                operations.append("Enrichissement IA")
            
            # Sauvegarde des métadonnées traitées
            self._save_metadata(processed_metadata, output_path, config)
            operations.append("Sauvegarde")
            
            processing_time = time.time() - start_time
            
            from . import TransformationResult, TransformationType
            return TransformationResult(
                success=True,
                input_path=input_path,
                output_path=output_path,
                transformation_type=TransformationType(config.type.value),
                metadata={
                    'original': original_metadata,
                    'processed': processed_metadata,
                    'enrichment_results': enrichment_results
                },
                errors=errors,
                warnings=warnings,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Erreur transformation métadonnées {input_path}: {e}")
            processing_time = time.time() - start_time
            
            from . import TransformationResult, TransformationType
            return TransformationResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transformation_type=TransformationType(config.type.value),
                metadata={},
                errors=[str(e)],
                warnings=warnings,
                processing_time=processing_time
            )
    
    def _normalize_metadata(
        self,
        metadata: Dict[str, Any],
        params: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Normalise les métadonnées selon les standards spécifiés"""
        
        standards = params.get('standards', ['dublin_core'])
        normalized_metadata = metadata.copy()
        
        for standard_name in standards:
            try:
                standard = MetadataStandard(standard_name)
                normalized = self.normalizer.normalize_to_standard(metadata, standard)
                normalized_metadata[f'{standard_name}_normalized'] = normalized
            except ValueError:
                self.logger.warning(f"Standard non reconnu: {standard_name}")
        
        return normalized_metadata
    
    def _enrich_metadata(
        self,
        file_path: str,
        metadata: Dict[str, Any],
        params: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Enrichit les métadonnées avec des analyses IA"""
        
        enrichment_types_str = params.get('enrichment_types', ['ai_tagging'])
        enrichment_types = []
        
        for type_str in enrichment_types_str:
            try:
                enrichment_types.append(EnrichmentType(type_str))
            except ValueError:
                self.logger.warning(f"Type d'enrichissement non reconnu: {type_str}")
        
        enriched_metadata = self.enricher.enrich_metadata(
            file_path, metadata, enrichment_types
        )
        
        return {
            'metadata': enriched_metadata,
            'enrichment_results': enriched_metadata.get('enrichment_results', [])
        }
    
    def _save_metadata(
        self,
        metadata: Dict[str, Any],
        output_path: str,
        config: 'TransformationConfig'
    ) -> None:
        """Sauvegarde les métadonnées traitées"""
        
        # Création du répertoire si nécessaire
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        output_format = config.output_format or 'json'
        
        if output_format == 'json':
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False, default=str)
        
        elif output_format == 'xml':
            self._save_as_xml(metadata, output_path)
        
        elif output_format == 'csv':
            self._save_as_csv(metadata, output_path)
        
        else:
            # Format JSON par défaut
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False, default=str)
    
    def _save_as_xml(self, metadata: Dict[str, Any], output_path: str) -> None:
        """Sauvegarde en format XML"""
        
        def dict_to_xml(tag, d):
            elem = ET.Element(tag)
            for key, val in d.items():
                child = ET.Element(sanitize_filename(str(key)))
                if isinstance(val, dict):
                    child = dict_to_xml(sanitize_filename(str(key)), val)
                elif isinstance(val, list):
                    for item in val:
                        if isinstance(item, dict):
                            child.append(dict_to_xml('item', item))
                        else:
                            item_elem = ET.Element('item')
                            item_elem.text = str(item)
                            child.append(item_elem)
                else:
                    child.text = str(val)
                elem.append(child)
            return elem
        
        root = dict_to_xml('metadata', metadata)
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
    
    def _save_as_csv(self, metadata: Dict[str, Any], output_path: str) -> None:
        """Sauvegarde en format CSV (aplati)"""
        
        import csv
        
        def flatten_dict(d, parent_key='', sep='_'):
            items = []
            for k, v in d.items():
                new_key = f"{parent_key}{sep}{k}" if parent_key else k
                if isinstance(v, dict):
                    items.extend(flatten_dict(v, new_key, sep=sep).items())
                elif isinstance(v, list):
                    items.append((new_key, '; '.join(str(item) for item in v)))
                else:
                    items.append((new_key, str(v)))
            return dict(items)
        
        flattened = flatten_dict(metadata)
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Field', 'Value'])
            for key, value in flattened.items():
                writer.writerow([key, value])
    
    def _generate_output_path(self, input_path: str, config: 'TransformationConfig') -> str:
        """Génère le chemin de sortie automatiquement"""
        
        input_path_obj = Path(input_path)
        output_format = config.output_format or 'json'
        
        # Nom de fichier avec suffixe de transformation
        transform_suffix = config.type.value.replace('metadata_', '')
        new_name = f"{input_path_obj.stem}_{transform_suffix}_metadata.{output_format}"
        
        return str(input_path_obj.parent / new_name)

class AsyncMetadataTransformer:
    """Version asynchrone du transformateur de métadonnées"""
    
    def __init__(self):
        self.sync_transformer = MetadataTransformer()
        self.logger = logging.getLogger(__name__)
    
    async def transform_async(
        self,
        input_path: str,
        config: 'TransformationConfig',
        output_path: Optional[str] = None
    ) -> 'TransformationResult':
        """Transformation de métadonnées asynchrone"""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.sync_transformer.transform,
            input_path,
            config,
            output_path
        )
    
    async def transform_batch_async(
        self,
        inputs: List[Tuple[str, 'TransformationConfig']],
        max_concurrent: int = 8  # Plus de concurrence pour métadonnées
    ) -> List['TransformationResult']:
        """Transformation en lot asynchrone"""
        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def transform_single(input_config_tuple):
            async with semaphore:
                input_path, config = input_config_tuple
                return await self.transform_async(input_path, config)
        
        tasks = [transform_single(item) for item in inputs]
        return await asyncio.gather(*tasks, return_exceptions=True)

# Export des classes
__all__ = [
    'MetadataTransformer',
    'AsyncMetadataTransformer',
    'MetadataExtractor',
    'MetadataEnricher',
    'MetadataNormalizer',
    'MetadataStandard',
    'MetadataField',
    'EnrichmentType',
    'MetadataProcessingResult'
]
