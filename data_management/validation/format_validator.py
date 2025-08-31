"""🚀 Format Validation System - IA Influencer Agent Platform Enterprise
==================================================================
Module: backend/data_management/validation/format_validator.py
Author: Fahed Mlaiel (mlaiel@live.de)
==================================================================

⚠️  PROPRIÉTÉ INTELLECTUELLE EXCLUSIVE - FAHED MLAIEL ⚠️
© 2025 Fahed Mlaiel. Tous droits réservés.
Contact: mlaiel@live.de

🎯 SYSTÈME DE VALIDATION DE FORMATS MULTIMÉDIA
Validation avancée des formats pour tous types de créateurs
- Support multi-format (audio, vidéo, image, document)
- Détection de corruption et intégrité
- Validation headers et métadonnées
- Support formats professionnels
"""
from typing import Dict, List, Optional, Any, Union, Tuple, Set
import asyncio
import logging
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
import os
import magic
import mimetypes
import struct
import hashlib
import json

# File analysis libraries
import filetype
from mutagen import File as MutagenFile
from mutagen.id3 import ID3NoHeaderError

# Image format validation
from PIL import Image, UnidentifiedImageError
import piexif

# Video format validation
import cv2
from moviepy.editor import VideoFileClip

# Audio format validation
import soundfile as sf
import librosa

# Document format validation
import docx
import PyPDF2
from markdownify import markdownify

logger = logging.getLogger(__name__)

@dataclass
class FormatValidationResult:
    """Résultat de validation de format"""    is_valid: bool
    format_detected: str
    mime_type: str
    file_extension: str
    file_size: int
    errors: List[str]
    warnings: List[str]
    metadata: Dict[str, Any]
    integrity_score: float  # 0.0 - 1.0

class FormatDetector:
    """Détecteur avancé de formats de fichiers"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.FormatDetector")
        
        # Signatures de fichiers pour validation
        self.file_signatures = {
            # Audio formats
            'mp3': [b'\xFF\xFB', b'\xFF\xF3', b'\xFF\xF2', b'ID3'],
            'wav': [b'RIFF', b'WAVE'],
            'flac': [b'fLaC'],
            'ogg': [b'OggS'],
            'm4a': [b'ftypM4A'],
            'aiff': [b'FORM', b'AIFF'],
            
            # Video formats
            'mp4': [b'ftyp', b'moov'],
            'avi': [b'RIFF', b'AVI '],
            'mov': [b'ftypqt', b'moov'],
            'mkv': [b'\x1A\x45\xDF\xA3'],
            'webm': [b'\x1A\x45\xDF\xA3'],
            
            # Image formats
            'jpg': [b'\xFF\xD8\xFF'],
            'jpeg': [b'\xFF\xD8\xFF'],
            'png': [b'\x89PNG\r\n\x1a\n'],
            'gif': [b'GIF87a', b'GIF89a'],
            'bmp': [b'BM'],
            'tiff': [b'II*\x00', b'MM\x00*'],
            'webp': [b'RIFF', b'WEBP'],
            
            # Document formats
            'pdf': [b'%PDF'],
            'docx': [b'PK\x03\x04'],
            'rtf': [b'{\\rtf1'],
            'txt': [],  # Text files have no specific signature
            'md': [],   # Markdown files are text-based
        }
    
    def detect_format(self, file_path: str) -> Tuple[str, str, Dict[str, Any]]:
        """Détecte le format et le type MIME d'un fichier"""        metadata = {}
        
        try:
            # 1. Détection par extension
            extension = Path(file_path).suffix.lower().lstrip('.')
            
            # 2. Détection par libmagic
            try:
                mime_type = magic.from_file(file_path, mime=True)
                file_type_desc = magic.from_file(file_path)
                metadata['magic_description'] = file_type_desc
            except:
                mime_type = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
            
            # 3. Détection par filetype (basée sur les signatures)
            try:
                kind = filetype.guess(file_path)
                if kind:
                    detected_extension = kind.extension
                    detected_mime = kind.mime
                    metadata['filetype_detection'] = {
                        'extension': detected_extension,
                        'mime': detected_mime
                    }
                else:
                    detected_extension = extension
                    detected_mime = mime_type
            except:
                detected_extension = extension
                detected_mime = mime_type
            
            # 4. Validation par signature binaire
            signature_valid = self._validate_file_signature(file_path, detected_extension)
            metadata['signature_valid'] = signature_valid
            
            # 5. Cohérence entre détections
            if extension != detected_extension and detected_extension:
                metadata['extension_mismatch'] = {
                    'declared': extension,
                    'detected': detected_extension
                }
            
            return detected_extension or extension, detected_mime, metadata
            
        except Exception as e:
            self.logger.error(f"Erreur détection format {file_path}: {e}")
            return 'unknown', 'application/octet-stream', {'error': str(e)}
    
    def _validate_file_signature(self, file_path: str, format_name: str) -> bool:
        """Valide la signature binaire du fichier"""        try:
            if format_name not in self.file_signatures:
                return True  # Pas de signature à vérifier
            
            signatures = self.file_signatures[format_name]
            if not signatures:  # Formats sans signature spécifique
                return True
            
            with open(file_path, 'rb') as f:
                header = f.read(32)  # Lire les premiers 32 bytes
                
                for signature in signatures:
                    if header.startswith(signature) or signature in header:
                        return True
            
            return False
            
        except Exception as e:
            self.logger.error(f"Erreur validation signature {file_path}: {e}")
            return False

class AudioFormatValidator:
    """Validateur spécialisé pour formats audio"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.AudioFormatValidator")
        
        self.supported_formats = {
            'mp3': {'container': 'MPEG', 'codecs': ['MP3']},
            'wav': {'container': 'WAV', 'codecs': ['PCM']},
            'flac': {'container': 'FLAC', 'codecs': ['FLAC']},
            'ogg': {'container': 'OGG', 'codecs': ['Vorbis', 'FLAC', 'Opus']},
            'm4a': {'container': 'MP4', 'codecs': ['AAC', 'ALAC']},
            'aiff': {'container': 'AIFF', 'codecs': ['PCM']},
            'wma': {'container': 'ASF', 'codecs': ['WMA']}
        }
    
    def validate_audio_format(self, file_path: str, expected_format: str) -> FormatValidationResult:
        """Valide un format audio spécifique"""        errors = []
        warnings = []
        metadata = {}
        
        try:
            file_size = os.path.getsize(file_path)
            metadata['file_size'] = file_size
            
            # Validation avec soundfile
            try:
                with sf.SoundFile(file_path) as f:
                    metadata.update({
                        'channels': f.channels,
                        'samplerate': f.samplerate,
                        'subtype': f.subtype,
                        'format': f.format,
                        'frames': f.frames,
                        'duration': f.frames / f.samplerate
                    })
                    
                    # Validation paramètres audio
                    if f.samplerate < 8000:
                        warnings.append(f"Taux d'échantillonnage très bas: {f.samplerate}Hz")
                    elif f.samplerate > 192000:
                        warnings.append(f"Taux d'échantillonnage très élevé: {f.samplerate}Hz")
                    
                    if f.channels > 8:
                        warnings.append(f"Nombre de canaux élevé: {f.channels}")
                    
            except Exception as e:
                # Fallback avec librosa
                try:
                    y, sr = librosa.load(file_path, sr=None)
                    duration = librosa.get_duration(y=y, sr=sr)
                    
                    metadata.update({
                        'samplerate': sr,
                        'duration': duration,
                        'samples': len(y)
                    })
                except Exception as e2:
                    errors.append(f"Impossible de lire le fichier audio: {str(e2)}")
            
            # Validation des métadonnées avec Mutagen
            try:
                audio_file = MutagenFile(file_path)
                if audio_file:
                    mutagen_info = {}
                    if hasattr(audio_file, 'info'):
                        info = audio_file.info
                        mutagen_info.update({
                            'bitrate': getattr(info, 'bitrate', None),
                            'length': getattr(info, 'length', None),
                            'channels': getattr(info, 'channels', None),
                            'sample_rate': getattr(info, 'sample_rate', None)
                        })
                    
                    # Tags/métadonnées
                    if audio_file.tags:
                        tag_info = {}
                        for key, value in audio_file.tags.items():
                            tag_info[str(key)] = str(value[0]) if isinstance(value, list) else str(value)
                        mutagen_info['tags'] = tag_info
                    
                    metadata['mutagen'] = mutagen_info
                    
            except ID3NoHeaderError:
                warnings.append("Pas de header ID3 trouvé")
            except Exception as e:
                warnings.append(f"Erreur lecture métadonnées: {str(e)}")
            
            # Validation intégrité
            integrity_score = self._calculate_audio_integrity(file_path, metadata)
            
            # Détection du format réel
            format_detected = self._detect_audio_format(file_path, metadata)
            mime_type = f"audio/{format_detected}" if format_detected != 'unknown' else 'application/octet-stream'
            
            # Validation cohérence format
            if expected_format and format_detected != expected_format:
                if format_detected != 'unknown':
                    warnings.append(f"Format détecté ({format_detected}) différent de l'attendu ({expected_format})")
            
            return FormatValidationResult(
                is_valid=len(errors) == 0,
                format_detected=format_detected,
                mime_type=mime_type,
                file_extension=Path(file_path).suffix.lower(),
                file_size=file_size,
                errors=errors,
                warnings=warnings,
                metadata=metadata,
                integrity_score=integrity_score
            )
            
        except Exception as e:
            self.logger.error(f"Erreur validation audio {file_path}: {e}")
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension=Path(file_path).suffix.lower(),
                file_size=0,
                errors=[f"Erreur système: {str(e)}"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
    
    def _calculate_audio_integrity(self, file_path: str, metadata: Dict) -> float:
        """Calcule le score d'intégrité du fichier audio"""        scores = []
        
        # 1. Cohérence des métadonnées
        if 'duration' in metadata and metadata['duration'] > 0:
            scores.append(1.0)
        else:
            scores.append(0.0)
        
        # 2. Validité du taux d'échantillonnage
        if 'samplerate' in metadata:
            sr = metadata['samplerate']
            if 8000 <= sr <= 192000:
                scores.append(1.0)
            else:
                scores.append(0.5)
        
        # 3. Lecture possible du fichier
        try:
            y, sr = librosa.load(file_path, duration=1.0)  # Test 1 seconde
            if len(y) > 0:
                scores.append(1.0)
            else:
                scores.append(0.0)
        except:
            scores.append(0.0)
        
        return sum(scores) / len(scores) if scores else 0.0
    
    def _detect_audio_format(self, file_path: str, metadata: Dict) -> str:
        """Détecte le format audio réel"""        extension = Path(file_path).suffix.lower().lstrip('.')
        
        # Validation basée sur les informations Mutagen
        if 'mutagen' in metadata:
            mutagen_info = metadata['mutagen']
            
            # Implement more precise detection based on mutagen metadata
            try:
                # Check for specific audio format indicators
                if 'mime' in mutagen_info:
                    mime_type = mutagen_info['mime']
                    mime_to_format = {
                        'audio/mpeg': 'mp3',
                        'audio/mp4': 'm4a',
                        'audio/flac': 'flac',
                        'audio/ogg': 'ogg',
                        'audio/wav': 'wav',
                        'audio/x-wav': 'wav',
                        'audio/aiff': 'aiff',
                        'audio/x-aiff': 'aiff'
                    }
                    if mime_type in mime_to_format:
                        detected_format = mime_to_format[mime_type]
                        if detected_format in self.supported_formats:
                            return detected_format
                
                # Check for codec information
                if 'codec' in mutagen_info:
                    codec = mutagen_info['codec'].lower()
                    if 'mp3' in codec or 'mpeg' in codec:
                        return 'mp3'
                    elif 'flac' in codec:
                        return 'flac'
                    elif 'aac' in codec:
                        return 'm4a'
                    elif 'vorbis' in codec:
                        return 'ogg'
                
                # Check for bitrate and other quality indicators
                if 'bitrate' in mutagen_info:
                    bitrate = mutagen_info['bitrate']
                    # High bitrate often indicates lossless formats
                    if bitrate > 1000000 and extension in ['flac', 'wav', 'aiff']:
                        return extension
                
                # Check for file format specific fields
                if 'length' in mutagen_info and 'filesize' in mutagen_info:
                    # Calculate approximate bitrate to help identify format
                    length = mutagen_info['length']
                    filesize = mutagen_info['filesize']
                    if length > 0:
                        calculated_bitrate = (filesize * 8) / length
                        
                        # Typical bitrates for different formats
                        if calculated_bitrate > 1000000:  # > 1 Mbps, likely lossless
                            if extension in ['flac', 'wav', 'aiff']:
                                return extension
                        elif calculated_bitrate > 200000:  # > 200 kbps, likely high quality compressed
                            if extension in ['mp3', 'm4a', 'ogg']:
                                return extension
                
            except Exception as e:
                self.logger.warning(f"Error in precise audio format detection: {e}")
        
        
        return extension if extension in self.supported_formats else 'unknown'

class VideoFormatValidator:
    """Validateur spécialisé pour formats vidéo"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.VideoFormatValidator")
        
        self.supported_formats = {
            'mp4': {'container': 'MP4', 'video_codecs': ['H.264', 'H.265', 'AV1'], 'audio_codecs': ['AAC', 'MP3']},
            'avi': {'container': 'AVI', 'video_codecs': ['H.264', 'MJPEG', 'DivX'], 'audio_codecs': ['MP3', 'PCM']},
            'mov': {'container': 'QuickTime', 'video_codecs': ['H.264', 'H.265', 'ProRes'], 'audio_codecs': ['AAC', 'PCM']},
            'mkv': {'container': 'Matroska', 'video_codecs': ['H.264', 'H.265', 'VP9'], 'audio_codecs': ['AAC', 'FLAC', 'Opus']},
            'webm': {'container': 'WebM', 'video_codecs': ['VP8', 'VP9', 'AV1'], 'audio_codecs': ['Vorbis', 'Opus']},
            'flv': {'container': 'FLV', 'video_codecs': ['H.264', 'VP6'], 'audio_codecs': ['AAC', 'MP3']}
        }
    
    def validate_video_format(self, file_path: str, expected_format: str) -> FormatValidationResult:
        """Valide un format vidéo spécifique"""        errors = []
        warnings = []
        metadata = {}
        
        try:
            file_size = os.path.getsize(file_path)
            metadata['file_size'] = file_size
            
            # Validation avec OpenCV
            cap = cv2.VideoCapture(file_path)
            
            if not cap.isOpened():
                errors.append("Impossible d'ouvrir le fichier vidéo avec OpenCV")
            else:
                # Extraction métadonnées de base
                fps = cap.get(cv2.CAP_PROP_FPS)
                frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                fourcc = int(cap.get(cv2.CAP_PROP_FOURCC))
                
                metadata.update({
                    'fps': fps,
                    'frame_count': frame_count,
                    'width': width,
                    'height': height,
                    'fourcc': fourcc,
                    'duration': frame_count / fps if fps > 0 else 0
                })
                
                # Validation paramètres vidéo
                if fps <= 0:
                    errors.append("FPS invalide ou non détecté")
                elif fps < 1:
                    warnings.append(f"FPS très bas: {fps}")
                elif fps > 240:
                    warnings.append(f"FPS très élevé: {fps}")
                
                if width <= 0 or height <= 0:
                    errors.append("Résolution invalide")
                elif width < 160 or height < 120:
                    warnings.append(f"Résolution très basse: {width}x{height}")
                
                cap.release()
            
            # Validation avec MoviePy pour métadonnées avancées
            try:
                with VideoFileClip(file_path) as clip:
                    moviepy_metadata = {
                        'duration': clip.duration,
                        'fps': clip.fps,
                        'size': clip.size,
                        'aspect_ratio': clip.w / clip.h if clip.h > 0 else 0
                    }
                    
                    # Audio track info
                    if clip.audio:
                        moviepy_metadata['has_audio'] = True
                        if hasattr(clip.audio, 'fps'):
                            moviepy_metadata['audio_fps'] = clip.audio.fps
                    else:
                        moviepy_metadata['has_audio'] = False
                    
                    metadata['moviepy'] = moviepy_metadata
                    
            except Exception as e:
                warnings.append(f"Erreur MoviePy: {str(e)}")
            
            # Validation intégrité
            integrity_score = self._calculate_video_integrity(file_path, metadata)
            
            # Détection du format réel
            format_detected = self._detect_video_format(file_path, metadata)
            mime_type = f"video/{format_detected}" if format_detected != 'unknown' else 'application/octet-stream'
            
            # Validation cohérence format
            if expected_format and format_detected != expected_format:
                if format_detected != 'unknown':
                    warnings.append(f"Format détecté ({format_detected}) différent de l'attendu ({expected_format})")
            
            return FormatValidationResult(
                is_valid=len(errors) == 0,
                format_detected=format_detected,
                mime_type=mime_type,
                file_extension=Path(file_path).suffix.lower(),
                file_size=file_size,
                errors=errors,
                warnings=warnings,
                metadata=metadata,
                integrity_score=integrity_score
            )
            
        except Exception as e:
            self.logger.error(f"Erreur validation vidéo {file_path}: {e}")
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension=Path(file_path).suffix.lower(),
                file_size=0,
                errors=[f"Erreur système: {str(e)}"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
    
    def _calculate_video_integrity(self, file_path: str, metadata: Dict) -> float:
        """Calcule le score d'intégrité du fichier vidéo"""        scores = []
        
        # 1. Métadonnées valides
        if 'duration' in metadata and metadata['duration'] > 0:
            scores.append(1.0)
        else:
            scores.append(0.0)
        
        # 2. FPS valide
        if 'fps' in metadata and metadata['fps'] > 0:
            scores.append(1.0)
        else:
            scores.append(0.0)
        
        # 3. Résolution valide
        if 'width' in metadata and 'height' in metadata:
            if metadata['width'] > 0 and metadata['height'] > 0:
                scores.append(1.0)
            else:
                scores.append(0.0)
        
        # 4. Test de lecture premier frame
        try:
            cap = cv2.VideoCapture(file_path)
            ret, frame = cap.read()
            cap.release()
            
            if ret and frame is not None:
                scores.append(1.0)
            else:
                scores.append(0.0)
        except:
            scores.append(0.0)
        
        return sum(scores) / len(scores) if scores else 0.0
    
    def _detect_video_format(self, file_path: str, metadata: Dict) -> str:
        """Détecte le format vidéo réel"""        extension = Path(file_path).suffix.lower().lstrip('.')
        
        # Validation basée sur fourcc et autres métadonnées
        if 'fourcc' in metadata:
            fourcc = metadata['fourcc']
            
            # Implement mapping fourcc to format
            try:
                # Convert fourcc code to string if it's an integer
                if isinstance(fourcc, int):
                    # Convert fourcc integer to 4-character string
                    fourcc_str = ''.join([chr((fourcc >> 8*i) & 0xFF) for i in range(4)])
                else:
                    fourcc_str = str(fourcc)
                
                # FourCC to format mapping
                fourcc_to_format = {
                    # H.264 variants
                    'avc1': 'mp4',  # H.264 in MP4
                    'AVC1': 'mp4',
                    'H264': 'mp4',
                    'x264': 'mp4',
                    
                    # H.265/HEVC variants
                    'hvc1': 'mp4',  # H.265 in MP4
                    'hev1': 'mp4',
                    'H265': 'mp4',
                    'HEVC': 'mp4',
                    
                    # VP8/VP9 (WebM)
                    'VP80': 'webm',
                    'VP90': 'webm',
                    'VP8 ': 'webm',
                    'VP9 ': 'webm',
                    
                    # AV1
                    'AV01': 'mp4',  # AV1 can be in MP4 or WebM
                    'av01': 'mp4',
                    
                    # MPEG-4 variants
                    'mp4v': 'mp4',
                    'MP4V': 'mp4',
                    'DIVX': 'avi',
                    'DX50': 'avi',
                    'XVID': 'avi',
                    
                    # Legacy codecs
                    'MJPG': 'avi',  # Motion JPEG
                    'mjpg': 'avi',
                    'YUY2': 'avi',  # Raw YUV
                    'UYVY': 'avi',
                    
                    # QuickTime
                    'qt  ': 'mov',
                    'QT  ': 'mov',
                    
                    # Flash Video
                    'FLV1': 'flv',
                    'flv1': 'flv',
                    
                    # ProRes (typically in MOV)
                    'apch': 'mov',  # ProRes 422 HQ
                    'apcn': 'mov',  # ProRes 422
                    'apcs': 'mov',  # ProRes 422 LT
                    'apco': 'mov',  # ProRes 422 Proxy
                    'ap4h': 'mov',  # ProRes 4444
                }
                
                # Try exact match first
                if fourcc_str in fourcc_to_format:
                    detected_format = fourcc_to_format[fourcc_str]
                    if detected_format in self.supported_formats:
                        self.logger.info(f"Detected format {detected_format} from fourcc {fourcc_str}")
                        return detected_format
                
                # Try case-insensitive match
                fourcc_lower = fourcc_str.lower()
                for known_fourcc, format_name in fourcc_to_format.items():
                    if known_fourcc.lower() == fourcc_lower:
                        if format_name in self.supported_formats:
                            self.logger.info(f"Detected format {format_name} from fourcc {fourcc_str} (case-insensitive)")
                            return format_name
                
                # Log unknown fourcc for debugging
                self.logger.debug(f"Unknown fourcc code: {fourcc_str} (0x{fourcc:08x} if int)")
                
            except Exception as e:
                self.logger.warning(f"Error processing fourcc code: {e}")
        
        # Check other metadata for format hints
        if 'codec_name' in metadata:
            codec_name = metadata['codec_name'].lower()
            if 'h264' in codec_name or 'avc' in codec_name:
                return 'mp4'
            elif 'h265' in codec_name or 'hevc' in codec_name:
                return 'mp4'
            elif 'vp8' in codec_name or 'vp9' in codec_name:
                return 'webm'
            elif 'av1' in codec_name:
                return 'mp4'
        
        
        return extension if extension in self.supported_formats else 'unknown'

class ImageFormatValidator:
    """Validateur spécialisé pour formats image"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.ImageFormatValidator")
        
        self.supported_formats = {
            'jpg': {'compression': 'JPEG', 'color_modes': ['RGB', 'CMYK', 'L']},
            'jpeg': {'compression': 'JPEG', 'color_modes': ['RGB', 'CMYK', 'L']},
            'png': {'compression': 'PNG', 'color_modes': ['RGB', 'RGBA', 'L', 'LA', 'P']},
            'gif': {'compression': 'GIF', 'color_modes': ['P', 'L']},
            'bmp': {'compression': 'None', 'color_modes': ['RGB', 'L']},
            'tiff': {'compression': 'Various', 'color_modes': ['RGB', 'CMYK', 'L']},
            'webp': {'compression': 'WebP', 'color_modes': ['RGB', 'RGBA']},
            'raw': {'compression': 'None', 'color_modes': ['RGB']},
            'dng': {'compression': 'Various', 'color_modes': ['RGB']}
        }
    
    def validate_image_format(self, file_path: str, expected_format: str) -> FormatValidationResult:
        """Valide un format image spécifique"""        errors = []
        warnings = []
        metadata = {}
        
        try:
            file_size = os.path.getsize(file_path)
            metadata['file_size'] = file_size
            
            # Validation avec PIL
            try:
                with Image.open(file_path) as img:
                    metadata.update({
                        'width': img.width,
                        'height': img.height,
                        'mode': img.mode,
                        'format': img.format,
                        'size': img.size
                    })
                    
                    # Validation paramètres image
                    if img.width <= 0 or img.height <= 0:
                        errors.append("Dimensions invalides")
                    elif img.width < 32 or img.height < 32:
                        warnings.append(f"Image très petite: {img.width}x{img.height}")
                    elif img.width > 16384 or img.height > 16384:
                        warnings.append(f"Image très grande: {img.width}x{img.height}")
                    
                    # Validation mode couleur
                    if img.mode not in ['RGB', 'RGBA', 'L', 'LA', 'P', 'CMYK']:
                        warnings.append(f"Mode couleur inhabituel: {img.mode}")
                    
                    # Extraction métadonnées supplémentaires
                    if hasattr(img, 'info') and img.info:
                        metadata['pil_info'] = dict(img.info)
                    
                    # Validation format spécifique
                    if img.format:
                        pil_format = img.format.lower()
                        metadata['pil_format'] = pil_format
                        
                        if expected_format and pil_format != expected_format:
                            warnings.append(f"Format PIL ({pil_format}) différent de l'attendu ({expected_format})")
                    
            except UnidentifiedImageError:
                errors.append("Format d'image non reconnu par PIL")
            except Exception as e:
                errors.append(f"Erreur PIL: {str(e)}")
            
            # Extraction EXIF pour JPEG/TIFF
            if Path(file_path).suffix.lower() in ['.jpg', '.jpeg', '.tiff', '.tif']:
                try:
                    exif_data = piexif.load(file_path)
                    if exif_data:
                        # Conversion des données EXIF en format lisible
                        exif_readable = {}
                        for ifd_name, ifd in exif_data.items():
                            if isinstance(ifd, dict):
                                exif_readable[ifd_name] = {}
                                for tag, value in ifd.items():
                                    try:
                                        exif_readable[ifd_name][str(tag)] = str(value)
                                    except:
                                        pass
                        
                        metadata['exif'] = exif_readable
                        
                except Exception as e:
                    warnings.append(f"Erreur extraction EXIF: {str(e)}")
            
            # Validation intégrité
            integrity_score = self._calculate_image_integrity(file_path, metadata)
            
            # Détection du format réel
            format_detected = self._detect_image_format(file_path, metadata)
            mime_type = f"image/{format_detected}" if format_detected != 'unknown' else 'application/octet-stream'
            
            return FormatValidationResult(
                is_valid=len(errors) == 0,
                format_detected=format_detected,
                mime_type=mime_type,
                file_extension=Path(file_path).suffix.lower(),
                file_size=file_size,
                errors=errors,
                warnings=warnings,
                metadata=metadata,
                integrity_score=integrity_score
            )
            
        except Exception as e:
            self.logger.error(f"Erreur validation image {file_path}: {e}")
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension=Path(file_path).suffix.lower(),
                file_size=0,
                errors=[f"Erreur système: {str(e)}"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
    
    def _calculate_image_integrity(self, file_path: str, metadata: Dict) -> float:
        """Calcule le score d'intégrité du fichier image"""        scores = []
        
        # 1. Dimensions valides
        if 'width' in metadata and 'height' in metadata:
            if metadata['width'] > 0 and metadata['height'] > 0:
                scores.append(1.0)
            else:
                scores.append(0.0)
        
        # 2. Mode couleur valide
        if 'mode' in metadata:
            valid_modes = ['RGB', 'RGBA', 'L', 'LA', 'P', 'CMYK']
            if metadata['mode'] in valid_modes:
                scores.append(1.0)
            else:
                scores.append(0.5)
        
        # 3. Test de chargement complet
        try:
            with Image.open(file_path) as img:
                img.load()  # Force le chargement complet
                scores.append(1.0)
        except:
            scores.append(0.0)
        
        # 4. Cohérence taille fichier / résolution
        if 'width' in metadata and 'height' in metadata and 'file_size' in metadata:
            expected_min_size = metadata['width'] * metadata['height'] * 0.1  # Estimation très conservative
            if metadata['file_size'] >= expected_min_size:
                scores.append(1.0)
            else:
                scores.append(0.5)
        
        return sum(scores) / len(scores) if scores else 0.0
    
    def _detect_image_format(self, file_path: str, metadata: Dict) -> str:
        """Détecte le format image réel"""        if 'pil_format' in metadata:
            return metadata['pil_format']
        
        extension = Path(file_path).suffix.lower().lstrip('.')
        return extension if extension in self.supported_formats else 'unknown'

class DocumentFormatValidator:
    """Validateur spécialisé pour formats de documents"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentFormatValidator")
        
        self.supported_formats = {
            'txt': {'encoding': ['utf-8', 'ascii', 'latin-1'], 'structured': False},
            'md': {'encoding': ['utf-8'], 'structured': True},
            'html': {'encoding': ['utf-8'], 'structured': True},
            'pdf': {'encoding': ['binary'], 'structured': True},
            'docx': {'encoding': ['binary'], 'structured': True},
            'rtf': {'encoding': ['ascii', 'utf-8'], 'structured': True},
            'json': {'encoding': ['utf-8'], 'structured': True},
            'xml': {'encoding': ['utf-8'], 'structured': True}
        }
    
    def validate_document_format(self, file_path: str, expected_format: str) -> FormatValidationResult:
        """Valide un format de document spécifique"""        errors = []
        warnings = []
        metadata = {}
        
        try:
            file_size = os.path.getsize(file_path)
            metadata['file_size'] = file_size
            
            # Détection d'encodage pour fichiers texte
            if expected_format in ['txt', 'md', 'html', 'json', 'xml', 'rtf']:
                try:
                    import chardet
                    with open(file_path, 'rb') as f:
                        raw_data = f.read(10000)  # Échantillon pour détection
                        encoding_info = chardet.detect(raw_data)
                        
                        metadata.update({
                            'detected_encoding': encoding_info.get('encoding'),
                            'encoding_confidence': encoding_info.get('confidence', 0)
                        })
                        
                        if encoding_info['confidence'] < 0.8:
                            warnings.append(f"Détection d'encodage incertaine: {encoding_info['confidence']:.2f}")
                            
                except Exception as e:
                    warnings.append(f"Erreur détection encodage: {str(e)}")
            
            # Validation spécifique par format
            if expected_format == 'pdf':
                self._validate_pdf(file_path, metadata, errors, warnings)
            elif expected_format == 'docx':
                self._validate_docx(file_path, metadata, errors, warnings)
            elif expected_format in ['txt', 'md']:
                self._validate_text(file_path, metadata, errors, warnings)
            elif expected_format == 'json':
                self._validate_json(file_path, metadata, errors, warnings)
            elif expected_format == 'xml':
                self._validate_xml(file_path, metadata, errors, warnings)
            elif expected_format == 'html':
                self._validate_html(file_path, metadata, errors, warnings)
            
            # Validation intégrité
            integrity_score = self._calculate_document_integrity(file_path, metadata, expected_format)
            
            # Format détecté
            format_detected = expected_format if len(errors) == 0 else 'unknown'
            mime_type = self._get_document_mime_type(expected_format)
            
            return FormatValidationResult(
                is_valid=len(errors) == 0,
                format_detected=format_detected,
                mime_type=mime_type,
                file_extension=Path(file_path).suffix.lower(),
                file_size=file_size,
                errors=errors,
                warnings=warnings,
                metadata=metadata,
                integrity_score=integrity_score
            )
            
        except Exception as e:
            self.logger.error(f"Erreur validation document {file_path}: {e}")
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension=Path(file_path).suffix.lower(),
                file_size=0,
                errors=[f"Erreur système: {str(e)}"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
    
    def _validate_pdf(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique PDF"""        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                metadata.update({
                    'num_pages': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted
                })
                
                # Métadonnées du document
                if pdf_reader.metadata:
                    pdf_metadata = {}
                    for key, value in pdf_reader.metadata.items():
                        pdf_metadata[key] = str(value)
                    metadata['pdf_metadata'] = pdf_metadata
                
                # Validation structure
                if len(pdf_reader.pages) == 0:
                    errors.append("PDF sans pages")
                
                # Test lecture première page
                try:
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()
                    metadata['has_text'] = len(text.strip()) > 0
                except:
                    warnings.append("Impossible d'extraire le texte de la première page")
                    
        except Exception as e:
            errors.append(f"Erreur lecture PDF: {str(e)}")
    
    def _validate_docx(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique DOCX"""        try:
            doc = docx.Document(file_path)
            
            metadata.update({
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables)
            })
            
            # Propriétés du document
            if doc.core_properties:
                core_props = {}
                for prop in ['author', 'created', 'modified', 'title', 'subject']:
                    value = getattr(doc.core_properties, prop, None)
                    if value:
                        core_props[prop] = str(value)
                metadata['core_properties'] = core_props
            
            # Test contenu
            total_text = ""
            for paragraph in doc.paragraphs:
                total_text += paragraph.text
            
            metadata.update({
                'has_content': len(total_text.strip()) > 0,
                'content_length': len(total_text)
            })
            
        except Exception as e:
            errors.append(f"Erreur lecture DOCX: {str(e)}")
    
    def _validate_text(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique fichiers texte"""        try:
            encoding = metadata.get('detected_encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            metadata.update({
                'line_count': content.count('\n') + 1,
                'character_count': len(content),
                'word_count': len(content.split()),
                'empty': len(content.strip()) == 0
            })
            
            if len(content.strip()) == 0:
                warnings.append("Fichier texte vide")
                
        except Exception as e:
            errors.append(f"Erreur lecture fichier texte: {str(e)}")
    
    def _validate_json(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique JSON"""        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            metadata.update({
                'json_type': type(data).__name__,
                'valid_json': True
            })
            
            if isinstance(data, dict):
                metadata['json_keys'] = list(data.keys())
            elif isinstance(data, list):
                metadata['json_length'] = len(data)
                
        except json.JSONDecodeError as e:
            errors.append(f"JSON invalide: {str(e)}")
        except Exception as e:
            errors.append(f"Erreur lecture JSON: {str(e)}")
    
    def _validate_xml(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique XML"""        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            metadata.update({
                'xml_root_tag': root.tag,
                'xml_namespace': root.tag.split('}')[0][1:] if '}' in root.tag else None,
                'xml_children_count': len(root)
            })
            
        except ET.ParseError as e:
            errors.append(f"XML invalide: {str(e)}")
        except Exception as e:
            errors.append(f"Erreur lecture XML: {str(e)}")
    
    def _validate_html(self, file_path: str, metadata: Dict, errors: List[str], warnings: List[str]):
        """Validation spécifique HTML"""        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            metadata.update({
                'has_html_tag': soup.html is not None,
                'has_head_tag': soup.head is not None,
                'has_body_tag': soup.body is not None,
                'title': soup.title.string if soup.title else None
            })
            
            if not soup.html:
                warnings.append("Pas de balise <html> trouvée")
                
        except Exception as e:
            warnings.append(f"Erreur parsing HTML: {str(e)}")
    
    def _calculate_document_integrity(self, file_path: str, metadata: Dict, format_name: str) -> float:
        """Calcule le score d'intégrité du document"""        scores = []
        
        # 1. Fichier lisible
        try:
            with open(file_path, 'rb') as f:
                f.read(1024)  # Test lecture
            scores.append(1.0)
        except:
            scores.append(0.0)
        
        # 2. Taille cohérente
        file_size = metadata.get('file_size', 0)
        if file_size > 0:
            scores.append(1.0)
        else:
            scores.append(0.0)
        
        # 3. Structure valide selon le format
        if format_name == 'pdf':
            if metadata.get('num_pages', 0) > 0:
                scores.append(1.0)
            else:
                scores.append(0.0)
        elif format_name == 'docx':
            if not metadata.get('empty', True):
                scores.append(1.0)
            else:
                scores.append(0.0)
        elif format_name == 'json':
            if metadata.get('valid_json', False):
                scores.append(1.0)
            else:
                scores.append(0.0)
        else:
            scores.append(0.8)  # Score par défaut
        
        return sum(scores) / len(scores) if scores else 0.0
    
    def _get_document_mime_type(self, format_name: str) -> str:
        """Retourne le type MIME pour un format de document"""        mime_types = {
            'txt': 'text/plain',
            'md': 'text/markdown',
            'html': 'text/html',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'rtf': 'application/rtf',
            'json': 'application/json',
            'xml': 'application/xml'
        }
        
        return mime_types.get(format_name, 'application/octet-stream')

class FormatValidator:
    """Validateur principal de formats"""    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.FormatValidator")
        
        # Initialisation des validateurs spécialisés
        self.format_detector = FormatDetector()
        self.audio_validator = AudioFormatValidator()
        self.video_validator = VideoFormatValidator()
        self.image_validator = ImageFormatValidator()
        self.document_validator = DocumentFormatValidator()
        
        # Mapping des types de contenu
        self.content_type_mapping = {
            'audio': ['mp3', 'wav', 'flac', 'ogg', 'm4a', 'aiff', 'wma'],
            'video': ['mp4', 'avi', 'mov', 'mkv', 'webm', 'flv', 'wmv'],
            'image': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp', 'raw', 'dng'],
            'document': ['txt', 'md', 'html', 'pdf', 'docx', 'rtf', 'json', 'xml']
        }
    
    def validate_format(self, file_path: str, expected_content_type: Optional[str] = None) -> FormatValidationResult:
        """Valide le format d'un fichier"""        
        if not os.path.exists(file_path):
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension='',
                file_size=0,
                errors=["Fichier introuvable"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
        
        try:
            # Détection automatique du format
            detected_format, detected_mime, detection_metadata = self.format_detector.detect_format(file_path)
            
            # Détermination du type de contenu
            content_type = self._determine_content_type(detected_format, expected_content_type)
            
            # Validation spécialisée selon le type
            if content_type == 'audio':
                result = self.audio_validator.validate_audio_format(file_path, detected_format)
            elif content_type == 'video':
                result = self.video_validator.validate_video_format(file_path, detected_format)
            elif content_type == 'image':
                result = self.image_validator.validate_image_format(file_path, detected_format)
            elif content_type == 'document':
                result = self.document_validator.validate_document_format(file_path, detected_format)
            else:
                # Format non supporté
                return FormatValidationResult(
                    is_valid=False,
                    format_detected=detected_format,
                    mime_type=detected_mime,
                    file_extension=Path(file_path).suffix.lower(),
                    file_size=os.path.getsize(file_path),
                    errors=[f"Type de contenu non supporté: {content_type}"],
                    warnings=[],
                    metadata=detection_metadata,
                    integrity_score=0.0
                )
            
            # Enrichissement avec métadonnées de détection
            result.metadata.update(detection_metadata)
            
            return result
            
        except Exception as e:
            self.logger.error(f"Erreur validation format {file_path}: {e}")
            return FormatValidationResult(
                is_valid=False,
                format_detected='unknown',
                mime_type='application/octet-stream',
                file_extension=Path(file_path).suffix.lower(),
                file_size=0,
                errors=[f"Erreur système: {str(e)}"],
                warnings=[],
                metadata={},
                integrity_score=0.0
            )
    
    def _determine_content_type(self, format_name: str, expected_type: Optional[str]) -> str:
        """Détermine le type de contenu basé sur le format détecté"""        
        # Si un type est attendu, vérifier la cohérence
        if expected_type:
            expected_formats = self.content_type_mapping.get(expected_type, [])
            if format_name in expected_formats:
                return expected_type
        
        # Détection automatique
        for content_type, formats in self.content_type_mapping.items():
            if format_name in formats:
                return content_type
        
        return 'unknown'

class AsyncFormatValidator:
    """Version asynchrone du validateur de formats"""    
    def __init__(self):
        self.sync_validator = FormatValidator()
        self.logger = logging.getLogger(f"{__name__}.AsyncFormatValidator")
    
    async def validate_format(self, file_path: str, expected_content_type: Optional[str] = None) -> FormatValidationResult:
        """Valide le format de manière asynchrone"""        loop = asyncio.get_event_loop()
        
        result = await loop.run_in_executor(
            None,
            self.sync_validator.validate_format,
            file_path,
            expected_content_type
        )
        
        return result
    
    async def validate_batch(self, file_paths: List[str], expected_types: Optional[List[str]] = None) -> Dict[str, FormatValidationResult]:
        """Valide un lot de fichiers de manière asynchrone"""        if expected_types is None:
            expected_types = [None] * len(file_paths)
        
        # Création des tâches asynchrones
        tasks = []
        for i, file_path in enumerate(file_paths):
            expected_type = expected_types[i] if i < len(expected_types) else None
            task = self.validate_format(file_path, expected_type)
            tasks.append(task)
        
        # Exécution en parallèle
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Formatage des résultats
        validation_results = {}
        for i, result in enumerate(results):
            file_path = file_paths[i]
            
            if isinstance(result, Exception):
                validation_results[file_path] = FormatValidationResult(
                    is_valid=False,
                    format_detected='unknown',
                    mime_type='application/octet-stream',
                    file_extension='',
                    file_size=0,
                    errors=[f"Erreur validation: {str(result)}"],
                    warnings=[],
                    metadata={},
                    integrity_score=0.0
                )
            else:
                validation_results[file_path] = result
        
        return validation_results

# Export des classes principales
__all__ = [
    'FormatValidator',
    'AsyncFormatValidator',
    'FormatValidationResult',
    'FormatDetector',
    'AudioFormatValidator',
    'VideoFormatValidator',
    'ImageFormatValidator',
    'DocumentFormatValidator'
]
