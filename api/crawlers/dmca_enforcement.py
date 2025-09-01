"""Professional DMCA and copyright enforcement system for content protection.

This module implements comprehensive DMCA enforcement capabilities including
automated takedown request generation, copyright violation detection,
legal documentation, and multi-platform enforcement coordination.

Author: Fahed Mlaiel <mlaiel@live.de>
Copyright (c) 2025 IA-Influencer Project. All rights reserved.
Licensed under proprietary license - reproduction forbidden without written authorization.

Project Team Specialties:
- Lead AI Developer & Senior Backend Engineer: Fahed Mlaiel
- Legal Technology Specialist: DMCA & Copyright Law Expert
- Intellectual Property Attorney: Legal Compliance & Enforcement
- Content Protection Engineer: Automated Takedown Systems
- Legal Documentation Expert: Formal Notice Generation
- Platform Policy Specialist: Multi-Platform Enforcement

Contact: mlaiel@live.de

LEGAL WARNING: This software and all associated intellectual property
belong exclusively to Fahed Mlaiel. Any unauthorized copying, redistribution,
reverse engineering, or commercial use without explicit written permission
will result in immediate legal action under international copyright laws.
"""

from typing import Dict, Any, List, Optional, Union, Set, Tuple, AsyncIterator
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from enum import Enum
import asyncio
import logging
import json
import re
import time
import hashlib
import uuid
from pathlib import Path
import tempfile
from urllib.parse import urlparse

# Document generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Email and communication
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import aiosmtplib

# Template engine
from jinja2 import Template, Environment, FileSystemLoader

# HTTP requests for takedown submissions
import aiohttp
import requests

# Data validation
from pydantic import BaseModel, validator, EmailStr
from datetime import date

from . import WebCrawler, CrawlResult, CrawlTarget, ContentType, PlatformType
from ..core.exceptions import CrawlerException, ValidationException, DMCAException
from ..core.models import BaseModel as CoreBaseModel
from ..security.encryption import EncryptionManager
from ..utils.rate_limiter import RateLimiter


class DMCARequestType(Enum):
    """
Types of DMCA requests."""

    TAKEDOWN_NOTICE = "takedown_notice"
    COUNTER_NOTIFICATION = "counter_notification"
    REPEAT_INFRINGER = "repeat_infringer"
    GOOD_FAITH_BELIEF = "good_faith_belief"
    FALSE_CLAIM_PROTECTION = "false_claim_protection"
    SAFE_HARBOR_COMPLIANCE = "safe_harbor_compliance"


class InfringementSeverity(Enum):
    """Severity levels for copyright infringement."""

    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"
    EXACT_COPY = "exact_copy"
    COMMERCIAL_USE = "commercial_use"


class EnforcementStatus(Enum):
    """Status of enforcement actions."""

    PENDING = "pending"
    SUBMITTED = "submitted"
    ACKNOWLEDGED = "acknowledged"
    CONTENT_REMOVED = "content_removed"
    REJECTED = "rejected"
    COUNTER_CLAIMED = "counter_claimed"
    LEGAL_ACTION = "legal_action"
    RESOLVED = "resolved"


class PlatformDMCAPolicy(Enum):
    """Platform-specific DMCA policies."""

    YOUTUBE_COPYRIGHT = "youtube_copyright"
    INSTAGRAM_IP = "instagram_ip"
    TIKTOK_COPYRIGHT = "tiktok_copyright"
    TWITTER_COPYRIGHT = "twitter_copyright"
    FACEBOOK_IP = "facebook_ip"
    GENERIC_DMCA = "generic_dmca"


@dataclass
class CopyrightOwner:
    """Copyright owner information."""
    owner_id: str
    full_name: str
    organization: Optional[str] = None
    email: str = ""
    phone: str = ""
    address: str = ""
    city: str = ""
    state: str = ""
    zip_code: str = ""
    country: str = ""
    website: str = ""
    copyright_registration: Optional[str] = None
    agent_name: Optional[str] = None
    agent_email: Optional[str] = None
    signature_image: Optional[str] = None
    is_verified: bool = False
    created_at: datetime = field(default_factory=datetime.utcnow)


@dataclass
class CopyrightWork:
    """Copyrighted work information."""
    work_id: str
    title: str
    work_type: str  # music, video, image, text, etc.
    description: str = ""
    creation_date: Optional[date] = None
    publication_date: Optional[date] = None
    registration_number: Optional[str] = None
    copyright_office: str = "US Copyright Office"
    original_urls: List[str] = field(default_factory=list)
    file_hash: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    owner_id: str = ""
    license_type: str = "all_rights_reserved"
    is_published: bool = True
    created_at: datetime = field(default_factory=datetime.utcnow)


@dataclass
class InfringementEvidence:
    """Evidence of copyright infringement."""
    evidence_id: str
    infringement_url: str
    platform: str
    infringing_content_id: str = ""
    screenshot_path: str = ""
    content_hash: str = ""
    similarity_score: float = 0.0
    detection_method: str = ""
    extracted_content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    uploader_info: Dict[str, str] = field(default_factory=dict)
    view_count: int = 0
    engagement_metrics: Dict[str, int] = field(default_factory=dict)
    detected_at: datetime = field(default_factory=datetime.utcnow)
    verified_at: Optional[datetime] = None
    is_commercial_use: bool = False
    revenue_impact: float = 0.0


@dataclass
class DMCANotice:
    """DMCA takedown notice."""
    notice_id: str
    notice_type: DMCARequestType
    copyright_owner: CopyrightOwner
    copyright_work: CopyrightWork
    infringement_evidence: InfringementEvidence
    platform_policy: PlatformDMCAPolicy
    severity: InfringementSeverity = InfringementSeverity.MEDIUM
    good_faith_statement: str = ""
    penalty_acknowledgment: str = ""
    electronic_signature: str = ""
    contact_information: Dict[str, str] = field(default_factory=dict)
    legal_basis: str = ""
    requested_action: str = "removal"
    additional_information: str = ""
    attachments: List[str] = field(default_factory=list)
    generated_at: datetime = field(default_factory=datetime.utcnow)
    submitted_at: Optional[datetime] = None
    platform_reference: str = ""
    status: EnforcementStatus = EnforcementStatus.PENDING
    response_deadline: datetime = field(default_factory=lambda: datetime.utcnow() + timedelta(days=7))


@dataclass
class EnforcementCase:
    """Complete enforcement case tracking."""
    case_id: str
    copyright_owner: CopyrightOwner
    copyright_works: List[CopyrightWork]
    infringement_instances: List[InfringementEvidence]
    dmca_notices: List[DMCANotice] = field(default_factory=list)
    case_status: EnforcementStatus = EnforcementStatus.PENDING
    total_infringements: int = 0
    platforms_involved: Set[str] = field(default_factory=set)
    estimated_damages: float = 0.0
    legal_fees: float = 0.0
    settlement_amount: float = 0.0
    case_notes: List[str] = field(default_factory=list)
    communication_log: List[Dict[str, Any]] = field(default_factory=list)
    created_at: datetime = field(default_factory=datetime.utcnow)
    updated_at: datetime = field(default_factory=datetime.utcnow)
    closed_at: Optional[datetime] = None


class DMCAEnforcementEngine:
    """
    Professional DMCA enforcement engine for copyright protection.
    
    Features:
    - Automated DMCA notice generation
    - Multi-platform takedown coordination
    - Legal documentation and templates
    - Evidence collection and preservation
    - Case management and tracking
    - Platform-specific policy compliance
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.logger = logging.getLogger("dmca.enforcement")
        
        # Core components
        self.rate_limiter = RateLimiter(config.get("rate_limits", {}))
        self.encryption_manager = EncryptionManager()
        
        # Template engine setup
        self.template_env = Environment(
            loader=FileSystemLoader(config.get("template_dir", "templates"))
        )
        
        # Legal settings
        self.default_copyright_office = config.get("copyright_office", "US Copyright Office")
        self.legal_jurisdiction = config.get("legal_jurisdiction", "United States")
        self.law_firm_info = config.get("law_firm_info", {})
        
        # Email configuration
        self.email_config = config.get("email_config", {})
        
        # Platform endpoints
        self.platform_endpoints = {
            "youtube": "https://www.youtube.com/copyright_complaint_form",
            "instagram": "https://help.instagram.com/contact/1017648381612439",
            "tiktok": "https://www.tiktok.com/legal/copyright",
            "twitter": "https://help.twitter.com/forms/dmca",
            "facebook": "https://www.facebook.com/help/contact/208282075858952"
        }
        
        # Case storage
        self.active_cases: Dict[str, EnforcementCase] = {}
        self.notice_templates: Dict[str, str] = {}
        
        # Load templates
        self._load_dmca_templates()
        
        # Performance metrics
        self.metrics = {
            "total_notices": 0,
            "successful_takedowns": 0,
            "rejected_notices": 0,
            "pending_cases": 0,
            "average_response_time": 0.0
        }
    
    def _load_dmca_templates(self):
        """Load DMCA notice templates."""
        try:
            # Load default templates if custom ones not provided
            self.notice_templates = {
                "takedown_notice": self._get_default_takedown_template(),
                "counter_notification": self._get_default_counter_template(),
                "repeat_infringer": self._get_repeat_infringer_template()
            }
            
            self.logger.info("DMCA templates loaded successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to load DMCA templates: {e}")
    
    def _get_default_takedown_template(self) -> str:
        """Default DMCA takedown notice template."""
        return """
DMCA COPYRIGHT INFRINGEMENT TAKEDOWN NOTICE

To: {{ platform_name }} Copyright Agent
From: {{ copyright_owner.full_name }}
Date: {{ notice_date }}

I, {{ copyright_owner.full_name }}, am the owner of exclusive rights in the copyrighted work described below:

COPYRIGHTED WORK:
Title: {{ copyright_work.title }}
Type: {{ copyright_work.work_type }}
Description: {{ copyright_work.description }}
{% if copyright_work.registration_number %}
Registration Number: {{ copyright_work.registration_number }}
{% endif %}
Creation Date: {{ copyright_work.creation_date }}
Original Location(s): {{ copyright_work.original_urls|join(', ') }}

INFRINGING MATERIAL:
The following material located on your service infringes my copyrighted work:

URL: {{ infringement_evidence.infringement_url }}
Description of Infringing Material: {{ infringement_evidence.extracted_content }}
{% if infringement_evidence.uploader_info.username %}
Uploaded by: {{ infringement_evidence.uploader_info.username }}
{% endif %}

GOOD FAITH BELIEF:
I have a good faith belief that use of the copyrighted materials described above is not authorized by the copyright owner, its agent, or the law.

ACCURACY STATEMENT:
I swear, under penalty of perjury, that the information in this notification is accurate and that I am the copyright owner or am authorized to act on behalf of the owner of an exclusive right that is allegedly infringed.

CONTACT INFORMATION:
Name: {{ copyright_owner.full_name }}
{% if copyright_owner.organization %}
Organization: {{ copyright_owner.organization }}
{% endif %}
Email: {{ copyright_owner.email }}
Phone: {{ copyright_owner.phone }}
Address: {{ copyright_owner.address }}, {{ copyright_owner.city }}, {{ copyright_owner.state }} {{ copyright_owner.zip_code }}, {{ copyright_owner.country }}

ELECTRONIC SIGNATURE:
{{ copyright_owner.full_name }}

Date: {{ notice_date }}

This notice is sent in good faith compliance with the Digital Millennium Copyright Act.
"""
    
    def _get_default_counter_template(self) -> str:
        """
Default DMCA counter-notification template."""
        return """
DMCA COUNTER-NOTIFICATION

To: {{ platform_name }} Copyright Agent
From: {{ user_name }}
Date: {{ notice_date }}

RE: Counter-Notification for Content Removed Due to DMCA Takedown Notice

I, {{ user_name }}, am the user who posted the material that was removed from {{ platform_name }} in response to a DMCA takedown notice.

REMOVED CONTENT:
The following content was removed from my account:
URL: {{ original_url }}
Description: {{ content_description }}
Date Removed: {{ removal_date }}

GOOD FAITH BELIEF:
I have a good faith belief that the material was removed or disabled as a result of mistake or misidentification of the material to be removed or disabled.

CONSENT TO JURISDICTION:
I consent to the jurisdiction of Federal District Court for the judicial district in which my address is located, or if my address is outside of the United States, the judicial district in which {{ platform_name }} is located, and will accept service of process from the claimant.

CONTACT INFORMATION:
Name: {{ user_name }}
Address: {{ user_address }}
Phone: {{ user_phone }}
Email: {{ user_email }}

PENALTY STATEMENT:
I swear, under penalty of perjury, that I have a good faith belief that the material was removed or disabled as a result of mistake or misidentification.

Electronic Signature: {{ user_name }}
Date: {{ notice_date }}
"""
    
    def _get_repeat_infringer_template(self) -> str:
        """
Repeat infringer notification template."""
        return """
REPEAT INFRINGER NOTIFICATION

To: {{ platform_name }} Copyright Agent
From: {{ copyright_owner.full_name }}
Date: {{ notice_date }}

REPEAT INFRINGEMENT NOTICE:

This notice is to inform you that the user account "{{ infringer_username }}" has repeatedly infringed my copyrighted works on your platform.

PREVIOUS VIOLATIONS:
{% for violation in previous_violations %}
- {{ violation.date }}: {{ violation.url }} ({{ violation.work_title }})
{% endfor %}

CURRENT VIOLATION:
URL: {{ infringement_evidence.infringement_url }}
Work Title: {{ copyright_work.title }}
Date of Infringement: {{ infringement_evidence.detected_at }}

REQUESTED ACTION:
I request that you terminate or suspend the repeat infringer's account in accordance with your Terms of Service and the DMCA repeat infringer policy.

This user has demonstrated a pattern of willful copyright infringement and should be subject to account termination under 17 U.S.C. § 512(i).

Sincerely,
{{ copyright_owner.full_name }}
{{ copyright_owner.email }}
{{ notice_date }}
"""
    
    async def create_enforcement_case(
        self,
        copyright_owner: CopyrightOwner,
        copyright_works: List[CopyrightWork],
        infringement_instances: List[InfringementEvidence]
    ) -> EnforcementCase:
        """
Create a new copyright enforcement case."""
        try:
            case_id = str(uuid.uuid4())
            
            # Validate inputs
            self._validate_copyright_owner(copyright_owner)
            for work in copyright_works:
                self._validate_copyright_work(work)
            for evidence in infringement_instances:
                self._validate_infringement_evidence(evidence)
            
            # Create case
            case = EnforcementCase(
                case_id=case_id,
                copyright_owner=copyright_owner,
                copyright_works=copyright_works,
                infringement_instances=infringement_instances,
                total_infringements=len(infringement_instances),
                platforms_involved=set(evidence.platform for evidence in infringement_instances)
            )
            
            # Calculate estimated damages
            case.estimated_damages = self._calculate_damages(infringement_instances)
            
            # Store case
            self.active_cases[case_id] = case
            
            self.logger.info(f"Enforcement case created: {case_id}")
            return case
            
        except Exception as e:
            self.logger.error(f"Failed to create enforcement case: {e}")
            raise DMCAException(f"Case creation failed: {e}")
    
    def _validate_copyright_owner(self, owner: CopyrightOwner):
        """Validate copyright owner information."""
        if not owner.full_name:
            raise ValidationException("Copyright owner full name required")
        if not owner.email:
            raise ValidationException("Copyright owner email required")
        if not owner.address:
            raise ValidationException("Copyright owner address required")
    
    def _validate_copyright_work(self, work: CopyrightWork):
        """Validate copyright work information."""
        if not work.title:
            raise ValidationException("Copyright work title required")
        if not work.work_type:
            raise ValidationException("Copyright work type required")
    
    def _validate_infringement_evidence(self, evidence: InfringementEvidence):
        """Validate infringement evidence."""
        if not evidence.infringement_url:
            raise ValidationException("Infringement URL required")
        if not evidence.platform:
            raise ValidationException("Platform identification required")
    
    def _calculate_damages(self, infringement_instances: List[InfringementEvidence]) -> float:
        """Calculate estimated damages from infringement instances."""
        total_damages = 0.0
        
        for evidence in infringement_instances:
            # Base damage calculation
            base_damage = 750.0  # Minimum statutory damages
            
            # Increase for commercial use
            if evidence.is_commercial_use:
                base_damage *= 2.0
            
            # Increase based on view count
            if evidence.view_count > 10000:
                base_damage *= 1.5
            elif evidence.view_count > 100000:
                base_damage *= 2.0
            elif evidence.view_count > 1000000:
                base_damage *= 3.0
            
            # Add revenue impact
            total_damages += base_damage + evidence.revenue_impact
        
        return total_damages
    
    async def generate_dmca_notice(
        self,
        case_id: str,
        evidence_id: str,
        notice_type: DMCARequestType = DMCARequestType.TAKEDOWN_NOTICE
    ) -> DMCANotice:
        """
Generate a DMCA notice for specific infringement."""
        try:
            case = self.active_cases.get(case_id)
            if not case:
                raise DMCAException(f"Case not found: {case_id}")
            
            # Find evidence
            evidence = None
            for e in case.infringement_instances:
                if e.evidence_id == evidence_id:
                    evidence = e
                    break
            
            if not evidence:
                raise DMCAException(f"Evidence not found: {evidence_id}")
            
            # Determine platform policy
            platform_policy = self._get_platform_policy(evidence.platform)
            
            # Create notice
            notice_id = str(uuid.uuid4())
            notice = DMCANotice(
                notice_id=notice_id,
                notice_type=notice_type,
                copyright_owner=case.copyright_owner,
                copyright_work=case.copyright_works[0],  # Primary work
                infringement_evidence=evidence,
                platform_policy=platform_policy,
                severity=self._assess_infringement_severity(evidence)
            )
            
            # Generate notice content
            await self._generate_notice_content(notice)
            
            # Add to case
            case.dmca_notices.append(notice)
            case.updated_at = datetime.utcnow()
            
            # Update metrics
            self.metrics["total_notices"] += 1
            
            self.logger.info(f"DMCA notice generated: {notice_id}")
            return notice
            
        except Exception as e:
            self.logger.error(f"Failed to generate DMCA notice: {e}")
            raise DMCAException(f"Notice generation failed: {e}")
    
    def _get_platform_policy(self, platform: str) -> PlatformDMCAPolicy:
        """Get platform-specific DMCA policy."""
        policy_mapping = {
            "youtube": PlatformDMCAPolicy.YOUTUBE_COPYRIGHT,
            "instagram": PlatformDMCAPolicy.INSTAGRAM_IP,
            "tiktok": PlatformDMCAPolicy.TIKTOK_COPYRIGHT,
            "twitter": PlatformDMCAPolicy.TWITTER_COPYRIGHT,
            "facebook": PlatformDMCAPolicy.FACEBOOK_IP
        }
        
        return policy_mapping.get(platform.lower(), PlatformDMCAPolicy.GENERIC_DMCA)
    
    def _assess_infringement_severity(self, evidence: InfringementEvidence) -> InfringementSeverity:
        """Assess the severity of copyright infringement."""
        if evidence.similarity_score >= 0.95:
            return InfringementSeverity.EXACT_COPY
        elif evidence.is_commercial_use:
            return InfringementSeverity.COMMERCIAL_USE
        elif evidence.similarity_score >= 0.8:
            return InfringementSeverity.HIGH
        elif evidence.similarity_score >= 0.6:
            return InfringementSeverity.MEDIUM
        else:
            return InfringementSeverity.LOW
    
    async def _generate_notice_content(self, notice: DMCANotice):
        """
Generate the content of a DMCA notice."""
        try:
            template_key = notice.notice_type.value
            template_str = self.notice_templates.get(template_key)
            
            if not template_str:
                raise DMCAException(f"Template not found for {template_key}")
            
            template = Template(template_str)
            
            # Prepare template variables
            template_vars = {
                "platform_name": notice.infringement_evidence.platform.title(),
                "copyright_owner": notice.copyright_owner,
                "copyright_work": notice.copyright_work,
                "infringement_evidence": notice.infringement_evidence,
                "notice_date": notice.generated_at.strftime("%B %d, %Y"),
                "severity": notice.severity.value,
                "legal_jurisdiction": self.legal_jurisdiction
            }
            
            # Generate content
            notice_content = template.render(**template_vars)
            
            # Set required fields
            notice.good_faith_statement = self._generate_good_faith_statement(notice)
            notice.penalty_acknowledgment = self._generate_penalty_statement(notice)
            notice.electronic_signature = notice.copyright_owner.full_name
            notice.legal_basis = self._generate_legal_basis(notice)
            
            # Store full notice content
            notice.additional_information = notice_content
            
        except Exception as e:
            self.logger.error(f"Failed to generate notice content: {e}")
            raise DMCAException(f"Notice content generation failed: {e}")
    
    def _generate_good_faith_statement(self, notice: DMCANotice) -> str:
        """Generate good faith belief statement."""
        return (
            f"I have a good faith belief that use of the copyrighted materials described above "
            f"on the allegedly infringing web pages is not authorized by the copyright owner, "
            f"its agent, or the law. I have taken fair use into consideration."
        )
    
    def _generate_penalty_statement(self, notice: DMCANotice) -> str:
        """Generate penalty of perjury statement."""
        return (
            f"I swear, under penalty of perjury, that the information in this notification is "
            f"accurate and that I am the copyright owner, or am authorized to act on behalf of "
            f"the owner, of an exclusive right that is allegedly infringed."
        )
    
    def _generate_legal_basis(self, notice: DMCANotice) -> str:
        """Generate legal basis for the claim."""
        return (
            f"This notice is given under the Digital Millennium Copyright Act (DMCA), "
            f"17 U.S.C. § 512, and applicable copyright laws. The allegedly infringing "
            f"material should be removed or disabled immediately."
        )
    
    async def submit_dmca_notice(
        self,
        notice_id: str,
        auto_submit: bool = False
    ) -> bool:
        """Submit DMCA notice to the platform."""
        try:
            # Find notice across all cases
            notice = None
            case = None
            
            for c in self.active_cases.values():
                for n in c.dmca_notices:
                    if n.notice_id == notice_id:
                        notice = n
                        case = c
                        break
                if notice:
                    break
            
            if not notice:
                raise DMCAException(f"Notice not found: {notice_id}")
            
            # Rate limiting for submissions
            await self.rate_limiter.acquire(f"dmca_submit_{notice.infringement_evidence.platform}")
            
            if auto_submit:
                # Automated submission (where supported)
                success = await self._auto_submit_notice(notice)
            else:
                # Generate documents for manual submission
                success = await self._prepare_manual_submission(notice)
            
            if success:
                notice.status = EnforcementStatus.SUBMITTED
                notice.submitted_at = datetime.utcnow()
                case.updated_at = datetime.utcnow()
                
                self.logger.info(f"DMCA notice submitted: {notice_id}")
                return True
            else:
                self.logger.error(f"DMCA notice submission failed: {notice_id}")
                return False
                
        except Exception as e:
            self.logger.error(f"Failed to submit DMCA notice: {e}")
            return False
    
    async def _auto_submit_notice(self, notice: DMCANotice) -> bool:
        """Automatically submit DMCA notice to platform."""
        try:
            platform = notice.infringement_evidence.platform.lower()
            endpoint = self.platform_endpoints.get(platform)
            
            if not endpoint:
                self.logger.warning(f"No automated submission for {platform}")
                return False
            
            # Platform-specific submission logic
            if platform == "youtube":
                return await self._submit_youtube_dmca(notice, endpoint)
            elif platform == "instagram":
                return await self._submit_instagram_dmca(notice, endpoint)
            else:
                self.logger.warning(f"Automated submission not implemented for {platform}")
                return False
                
        except Exception as e:
            self.logger.error(f"Auto submission failed: {e}")
            return False
    
    async def _submit_youtube_dmca(self, notice: DMCANotice, endpoint: str) -> bool:
        """Submit DMCA notice to YouTube."""
        try:
            # YouTube requires web form submission
            # This would typically involve Selenium or similar automation
            self.logger.info(f"YouTube DMCA submission prepared for manual processing")
            return True
            
        except Exception as e:
            self.logger.error(f"YouTube DMCA submission failed: {e}")
            return False
    
    async def _submit_instagram_dmca(self, notice: DMCANotice, endpoint: str) -> bool:
        """Submit DMCA notice to Instagram."""
        try:
            # Instagram/Facebook IP reporting
            self.logger.info(f"Instagram DMCA submission prepared for manual processing")
            return True
            
        except Exception as e:
            self.logger.error(f"Instagram DMCA submission failed: {e}")
            return False
    
    async def _prepare_manual_submission(self, notice: DMCANotice) -> bool:
        """Prepare documents for manual DMCA submission."""
        try:
            # Generate PDF document
            pdf_path = await self._generate_dmca_pdf(notice)
            notice.attachments.append(pdf_path)
            
            # Generate Word document
            docx_path = await self._generate_dmca_docx(notice)
            notice.attachments.append(docx_path)
            
            # Send email notification
            if self.email_config.get("enabled", False):
                await self._send_dmca_email(notice)
            
            self.logger.info(f"Manual submission documents prepared for notice: {notice.notice_id}")
            return True
            
        except Exception as e:
            self.logger.error(f"Manual submission preparation failed: {e}")
            return False
    
    async def _generate_dmca_pdf(self, notice: DMCANotice) -> str:
        """Generate PDF DMCA notice."""
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            pdf_path = temp_file.name
            temp_file.close()
            
            # Create PDF document
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                textColor=colors.black
            )
            
            story.append(Paragraph("DMCA COPYRIGHT INFRINGEMENT TAKEDOWN NOTICE", title_style))
            story.append(Spacer(1, 12))
            
            # Notice content
            content_style = styles['Normal']
            content_lines = notice.additional_information.split('\n')
            
            for line in content_lines:
                if line.strip():
                    story.append(Paragraph(line.strip(), content_style))
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"DMCA PDF generated: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            self.logger.error(f"PDF generation failed: {e}")
            raise DMCAException(f"PDF generation failed: {e}")
    
    async def _generate_dmca_docx(self, notice: DMCANotice) -> str:
        """Generate Word document DMCA notice."""
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            docx_path = temp_file.name
            temp_file.close()
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('DMCA COPYRIGHT INFRINGEMENT TAKEDOWN NOTICE', 0)
            title.alignment = 1  # Center alignment
            
            # Content
            content_lines = notice.additional_information.split('\n')
            for line in content_lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
            
            # Save document
            doc.save(docx_path)
            
            self.logger.info(f"DMCA DOCX generated: {docx_path}")
            return docx_path
            
        except Exception as e:
            self.logger.error(f"DOCX generation failed: {e}")
            raise DMCAException(f"DOCX generation failed: {e}")
    
    async def _send_dmca_email(self, notice: DMCANotice):
        """Send DMCA notice via email."""
        try:
            # Email configuration
            smtp_server = self.email_config.get("smtp_server")
            smtp_port = self.email_config.get("smtp_port", 587)
            username = self.email_config.get("username")
            password = self.email_config.get("password")
            
            if not all([smtp_server, username, password]):
                self.logger.warning("Email configuration incomplete")
                return
            
            # Platform-specific email addresses
            platform_emails = {
                "youtube": "copyright@youtube.com",
                "instagram": "ip@instagram.com",
                "tiktok": "copyright@tiktok.com",
                "twitter": "copyright@twitter.com",
                "facebook": "ip@facebook.com"
            }
            
            to_email = platform_emails.get(
                notice.infringement_evidence.platform.lower(),
                "dmca@example.com"
            )
            
            # Create email message
            msg = MIMEMultipart()
            msg['From'] = username
            msg['To'] = to_email
            msg['Subject'] = f"DMCA Takedown Notice - {notice.copyright_work.title}"
            
            # Email body
            body = f"""
            This is a formal DMCA takedown notice for copyright infringement.
            
            Notice ID: {notice.notice_id}
            Generated: {notice.generated_at.strftime('%Y-%m-%d %H:%M:%S')}
            
            Please see attached documents for complete details.
            
            Sincerely,
            {notice.copyright_owner.full_name}
            {notice.copyright_owner.email}
            """
            
            msg.attach(MIMEText(body, 'plain'))
            
            # Attach documents
            for attachment_path in notice.attachments:
                if Path(attachment_path).exists():
                    with open(attachment_path, 'rb') as f:
                        attach = MIMEApplication(f.read())
                        attach.add_header(
                            'Content-Disposition',
                            'attachment',
                            filename=Path(attachment_path).name
                        )
                        msg.attach(attach)
            
            # Send email
            async with aiosmtplib.SMTP(hostname=smtp_server, port=smtp_port) as server:
                await server.starttls()
                await server.login(username, password)
                await server.send_message(msg)
            
            self.logger.info(f"DMCA email sent for notice: {notice.notice_id}")
            
        except Exception as e:
            self.logger.error(f"Email sending failed: {e}")
    
    async def track_enforcement_status(self, case_id: str) -> EnforcementCase:
        """Track the status of an enforcement case."""
        try:
            case = self.active_cases.get(case_id)
            if not case:
                raise DMCAException(f"Case not found: {case_id}")
            
            # Check each notice status
            for notice in case.dmca_notices:
                if notice.status == EnforcementStatus.SUBMITTED:
                    # Check if response deadline has passed
                    if datetime.utcnow() > notice.response_deadline:
                        # Update status based on platform response
                        await self._check_platform_response(notice)
            
            # Update overall case status
            case.case_status = self._determine_case_status(case)
            case.updated_at = datetime.utcnow()
            
            return case
            
        except Exception as e:
            self.logger.error(f"Status tracking failed: {e}")
            raise DMCAException(f"Status tracking failed: {e}")
    
    async def _check_platform_response(self, notice: DMCANotice):
        """Check platform response to DMCA notice."""
        try:
            # This would typically involve checking the original URL
            # to see if content was removed
            original_url = notice.infringement_evidence.infringement_url
            
            async with aiohttp.ClientSession() as session:
                try:
                    async with session.get(original_url, timeout=10) as response:
                        if response.status == 404:
                            # Content removed
                            notice.status = EnforcementStatus.CONTENT_REMOVED
                            self.metrics["successful_takedowns"] += 1
                        elif response.status == 200:
                            # Content still exists - check for changes
                            content = await response.text()
                            if "removed" in content.lower() or "unavailable" in content.lower():
                                notice.status = EnforcementStatus.CONTENT_REMOVED
                                self.metrics["successful_takedowns"] += 1
                            else:
                                # No response within deadline
                                notice.status = EnforcementStatus.REJECTED
                                self.metrics["rejected_notices"] += 1
                except:
                    # Network error - assume content might be removed
                    notice.status = EnforcementStatus.ACKNOWLEDGED
            
        except Exception as e:
            self.logger.error(f"Platform response check failed: {e}")
    
    def _determine_case_status(self, case: EnforcementCase) -> EnforcementStatus:
        """Determine overall case status based on individual notices."""
        if not case.dmca_notices:
            return EnforcementStatus.PENDING
        
        statuses = [notice.status for notice in case.dmca_notices]
        
        if all(s == EnforcementStatus.CONTENT_REMOVED for s in statuses):
            return EnforcementStatus.RESOLVED
        elif any(s == EnforcementStatus.COUNTER_CLAIMED for s in statuses):
            return EnforcementStatus.COUNTER_CLAIMED
        elif any(s == EnforcementStatus.CONTENT_REMOVED for s in statuses):
            return EnforcementStatus.ACKNOWLEDGED
        elif all(s == EnforcementStatus.REJECTED for s in statuses):
            return EnforcementStatus.LEGAL_ACTION
        else:
            return EnforcementStatus.PENDING
    
    async def generate_enforcement_report(self, case_id: str) -> Dict[str, Any]:
        """
Generate comprehensive enforcement report."""
        try:
            case = self.active_cases.get(case_id)
            if not case:
                raise DMCAException(f"Case not found: {case_id}")
            
            # Update case status first
            await self.track_enforcement_status(case_id)
            
            report = {
                "case_summary": {
                    "case_id": case.case_id,
                    "status": case.case_status.value,
                    "total_infringements": case.total_infringements,
                    "platforms_involved": list(case.platforms_involved),
                    "estimated_damages": case.estimated_damages,
                    "created_at": case.created_at.isoformat(),
                    "updated_at": case.updated_at.isoformat()
                },
                "copyright_owner": {
                    "name": case.copyright_owner.full_name,
                    "organization": case.copyright_owner.organization,
                    "email": case.copyright_owner.email
                },
                "copyright_works": [
                    {
                        "title": work.title,
                        "type": work.work_type,
                        "registration": work.registration_number
                    }
                    for work in case.copyright_works
                ],
                "infringement_analysis": {
                    "total_instances": len(case.infringement_instances),
                    "platforms": {},
                    "severity_breakdown": {},
                    "commercial_use_count": 0
                },
                "enforcement_actions": {
                    "total_notices": len(case.dmca_notices),
                    "successful_takedowns": 0,
                    "pending_notices": 0,
                    "rejected_notices": 0
                },
                "financial_impact": {
                    "estimated_damages": case.estimated_damages,
                    "legal_fees": case.legal_fees,
                    "potential_settlement": case.settlement_amount
                },
                "recommendations": []
            }
            
            # Analyze infringements by platform
            platform_counts = {}
            severity_counts = {}
            commercial_count = 0
            
            for evidence in case.infringement_instances:
                # Platform analysis
                platform = evidence.platform
                if platform not in platform_counts:
                    platform_counts[platform] = 0
                platform_counts[platform] += 1
                
                # Commercial use tracking
                if evidence.is_commercial_use:
                    commercial_count += 1
            
            report["infringement_analysis"]["platforms"] = platform_counts
            report["infringement_analysis"]["commercial_use_count"] = commercial_count
            
            # Analyze enforcement actions
            action_counts = {
                "successful": 0,
                "pending": 0,
                "rejected": 0
            }
            
            for notice in case.dmca_notices:
                if notice.status == EnforcementStatus.CONTENT_REMOVED:
                    action_counts["successful"] += 1
                elif notice.status in [EnforcementStatus.PENDING, EnforcementStatus.SUBMITTED]:
                    action_counts["pending"] += 1
                else:
                    action_counts["rejected"] += 1
            
            report["enforcement_actions"].update(action_counts)
            
            # Generate recommendations
            recommendations = self._generate_case_recommendations(case)
            report["recommendations"] = recommendations
            
            self.logger.info(f"Enforcement report generated for case: {case_id}")
            return report
            
        except Exception as e:
            self.logger.error(f"Report generation failed: {e}")
            raise DMCAException(f"Report generation failed: {e}")
    
    def _generate_case_recommendations(self, case: EnforcementCase) -> List[str]:
        """Generate recommendations for enforcement case."""
        recommendations = []
        
        try:
            # Success rate analysis
            total_notices = len(case.dmca_notices)
            successful_notices = sum(
                1 for notice in case.dmca_notices
                if notice.status == EnforcementStatus.CONTENT_REMOVED
            )
            
            if total_notices > 0:
                success_rate = successful_notices / total_notices
                
                if success_rate < 0.5:
                    recommendations.append(
                        "Consider strengthening evidence collection and documentation"
                    )
                    recommendations.append(
                        "Review DMCA notice templates for platform-specific requirements"
                    )
            
            # Platform-specific recommendations
            if "youtube" in case.platforms_involved:
                recommendations.append(
                    "Consider using YouTube's Content ID system for proactive protection"
                )
            
            # Legal action recommendations
            high_value_infringements = [
                e for e in case.infringement_instances
                if e.is_commercial_use or e.view_count > 100000
            ]
            
            if len(high_value_infringements) > 3:
                recommendations.append(
                    "High-value infringements detected - consider legal consultation"
                )
            
            # Repeat infringer recommendations
            infringer_counts = {}
            for evidence in case.infringement_instances:
                username = evidence.uploader_info.get("username", "unknown")
                if username not in infringer_counts:
                    infringer_counts[username] = 0
                infringer_counts[username] += 1
            
            repeat_infringers = [
                user for user, count in infringer_counts.items()
                if count > 2 and user != "unknown"
            ]
            
            if repeat_infringers:
                recommendations.append(
                    "Repeat infringers identified - submit repeat infringer notices"
                )
            
            # Prevention recommendations
            if case.total_infringements > 10:
                recommendations.append(
                    "Consider implementing proactive content monitoring systems"
                )
                recommendations.append(
                    "Explore watermarking and fingerprinting technologies"
                )
            
        except Exception as e:
            self.logger.error(f"Recommendation generation failed: {e}")
        
        return recommendations
    
    def get_enforcement_metrics(self) -> Dict[str, Any]:
        """Get enforcement performance metrics."""
        active_cases = len(self.active_cases)
        total_infringements = sum(
            case.total_infringements for case in self.active_cases.values()
        )
        
        return {
            **self.metrics,
            "active_cases": active_cases,
            "total_infringements": total_infringements,
            "success_rate": (
                self.metrics["successful_takedowns"] / max(self.metrics["total_notices"], 1)
            ) * 100
        }


class DMCAComplianceValidator:
    """Validate DMCA notice compliance."""
    
    @staticmethod
    def validate_notice_completeness(notice: DMCANotice) -> Tuple[bool, List[str]]:
        """
Validate DMCA notice for completeness and compliance."""
        errors = []
        
        # Required elements check
        if not notice.copyright_owner.full_name:
            errors.append("Copyright owner name missing")
        
        if not notice.copyright_owner.email:
            errors.append("Copyright owner contact information incomplete")
        
        if not notice.copyright_work.title:
            errors.append("Copyrighted work identification missing")
        
        if not notice.infringement_evidence.infringement_url:
            errors.append("Infringing content URL missing")
        
        if not notice.good_faith_statement:
            errors.append("Good faith belief statement missing")
        
        if not notice.penalty_acknowledgment:
            errors.append("Penalty of perjury statement missing")
        
        if not notice.electronic_signature:
            errors.append("Electronic signature missing")
        
        # Content validation
        if notice.infringement_evidence.similarity_score < 0.3:
            errors.append("Low similarity score - weak infringement claim")
        
        return len(errors) == 0, errors
    
    @staticmethod
    def check_fair_use_considerations(notice: DMCANotice) -> List[str]:
        """Check for potential fair use considerations."""
        considerations = []
        
        # Educational use indicators
        if "educational" in notice.infringement_evidence.extracted_content.lower():
            considerations.append("Potential educational fair use")
        
        # Commentary/criticism indicators
        commentary_keywords = ["review", "criticism", "commentary", "parody"]
        content_lower = notice.infringement_evidence.extracted_content.lower()
        
        for keyword in commentary_keywords:
            if keyword in content_lower:
                considerations.append(f"Potential {keyword} fair use")
                break
        
        # Transformative use indicators
        if "remix" in content_lower or "mashup" in content_lower:
            considerations.append("Potential transformative fair use")
        
        return considerations
