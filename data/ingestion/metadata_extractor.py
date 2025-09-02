"""Metadata Extraction Engine
=========================

Professional metadata extraction system for multi-format content analysis.
Provides comprehensive metadata extraction, enhancement, and standardization
for audio, video, image, and text content with AI-powered analysis.

Author: Fahed Mlaiel (mlaiel@live.de)
Copyright: (c) 2025 Fahed Mlaiel - All Rights Reserved

⚠️  INTELLECTUAL PROPERTY WARNING ⚠️
This code is proprietary and confidential. Any unauthorized copying, distribution,
or use without explicit written permission from Fahed Mlaiel (mlaiel@live.de) is
strictly prohibited and will result in legal action.
"""

import asyncio
import logging
import mimetypes
import tempfile
import os
from datetime import datetime
from typing import Dict, List, Optional, Any, Union, BinaryIO
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import json
import hashlib

# Media processing libraries
import librosa
import cv2
import numpy as np
from PIL import Image
from PIL.ExifTags import TAGS
import ffmpeg
import mutagen
from mutagen.id3 import ID3NoHeaderError
import fitz  # PyMuPDF
import docx
import magic

# AI and NLP libraries  
from transformers import pipeline
import spacy
from langdetect import detect
import textstat

from ...core.exceptions import MetadataExtractionError, ValidationError


class MetadataType(Enum):
    """
Types of metadata that can be extracted"""

    TECHNICAL = "technical"
    DESCRIPTIVE = "descriptive"
    ADMINISTRATIVE = "administrative"
    STRUCTURAL = "structural"
    AI_GENERATED = "ai_generated"
    PRESERVATION = "preservation"


class ContentFormat(Enum):
    """Supported content formats for metadata extraction"""

    AUDIO = "audio"
    VIDEO = "video"
    IMAGE = "image"
    TEXT = "text"
    DOCUMENT = "document"


@dataclass
class MetadataField:
    """Individual metadata field with validation and type information"""
    name: str
    value: Any
    type: str
    confidence: float = 1.0
    source: str = "system"
    extracted_at: datetime = field(default_factory=datetime.utcnow)
    validation_status: str = "valid"


@dataclass
class MetadataCollection:
    """Collection of metadata organized by type and source"""
    content_id: str
    content_type: ContentFormat
    technical_metadata: Dict[str, Any] = field(default_factory=dict)
    descriptive_metadata: Dict[str, Any] = field(default_factory=dict)
    administrative_metadata: Dict[str, Any] = field(default_factory=dict)
    structural_metadata: Dict[str, Any] = field(default_factory=dict)
    ai_metadata: Dict[str, Any] = field(default_factory=dict)
    preservation_metadata: Dict[str, Any] = field(default_factory=dict)
    custom_metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_summary: Dict[str, Any] = field(default_factory=dict)
    quality_score: float = 0.0
    completeness_score: float = 0.0


class MetadataExtractor:
    """
    Professional metadata extraction engine for IA Influencer Agent platform.
    
    Provides comprehensive metadata extraction capabilities including:
    - Technical metadata (format, dimensions, encoding, etc.)
    - Descriptive metadata (title, description, tags, etc.)
    - AI-generated metadata (content analysis, features, etc.)
    - EXIF/XMP data for images
    - ID3/Vorbis tags for audio
    - Document properties and text analysis
    - Video stream information and analysis
    """
    
    def __init__(self):
        """
Initialize MetadataExtractor with AI models and processors."""
        self.logger = logging.getLogger(__name__)
        
        # Initialize AI models for content analysis
        self._init_ai_models()
        
        # Supported formats configuration
        self.supported_formats = {
            ContentFormat.AUDIO: ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a', '.wma', '.mp4'],
            ContentFormat.VIDEO: ['.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.m4v'],
            ContentFormat.IMAGE: ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg'],
            ContentFormat.TEXT: ['.txt', '.md', '.html', '.json', '.xml', '.csv'],
            ContentFormat.DOCUMENT: ['.pdf', '.docx', '.doc', '.rtf', '.odt']
        }
        
        # Metadata extraction rules
        self.extraction_rules = {
            'required_fields': ['format', 'size', 'created_at', 'mime_type'],
            'optional_fields': ['title', 'description', 'author', 'tags', 'duration'],
            'ai_fields': ['content_description', 'sentiment', 'topics', 'quality_score']
        }
    
    def _init_ai_models(self):
        """
Initialize AI models for content analysis"""
        try:
            # Image analysis models
            self.image_classifier = pipeline(
                "image-classification",
                model="microsoft/resnet-50",
                device=-1
            )
            
            # Text analysis models
            self.text_classifier = pipeline(
                "text-classification",
                model="distilbert-base-uncased-finetuned-sst-2-english"
            )
            
            self.ner_model = pipeline(
                "ner",
                model="dbmdz/bert-large-cased-finetuned-conll03-english",
                aggregation_strategy="simple"
            )
            
            # Load spaCy model for advanced NLP
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                self.logger.warning("spaCy English model not found, some NLP features disabled")
                self.nlp = None
            
            self.logger.info("AI models initialized successfully")
            
        except Exception as e:
            self.logger.warning(f"AI model initialization failed: {str(e)}")
            # Set fallback None values
            self.image_classifier = None
            self.text_classifier = None
            self.ner_model = None
            self.nlp = None
    
    async def extract_metadata(self, file_data: Union[bytes, BinaryIO], 
                             filename: str, content_type: ContentFormat = None,
                             include_ai_analysis: bool = True) -> MetadataCollection:
        """
        Extract comprehensive metadata from content.
        
        Args:
            file_data: Content file data
            filename: Original filename
            content_type: Content format (auto-detected if None)
            include_ai_analysis: Whether to include AI-powered analysis
            
        Returns:
            Complete metadata collection
        """
        content_id = hashlib.sha256(
            (filename + str(datetime.utcnow())).encode()
        ).hexdigest()[:16]
        
        try:
            self.logger.info(f"Starting metadata extraction: {content_id}")
            
            # Convert file data to bytes if needed
            if hasattr(file_data, 'read'):
                file_bytes = file_data.read()
                file_data.seek(0)
            else:
                file_bytes = file_data
            
            # Auto-detect content type if not provided
            if content_type is None:
                content_type = await self._detect_content_format(file_bytes, filename)
            
            # Initialize metadata collection
            metadata_collection = MetadataCollection(
                content_id=content_id,
                content_type=content_type
            )
            
            # Extract basic file metadata
            await self._extract_basic_metadata(metadata_collection, file_bytes, filename)
            
            # Extract format-specific metadata
            if content_type == ContentFormat.AUDIO:
                await self._extract_audio_metadata(metadata_collection, file_bytes, filename)
            elif content_type == ContentFormat.VIDEO:
                await self._extract_video_metadata(metadata_collection, file_bytes, filename)
            elif content_type == ContentFormat.IMAGE:
                await self._extract_image_metadata(metadata_collection, file_bytes, filename)
            elif content_type in [ContentFormat.TEXT, ContentFormat.DOCUMENT]:
                await self._extract_text_metadata(metadata_collection, file_bytes, filename)
            
            # AI-powered metadata enhancement
            if include_ai_analysis:
                await self._extract_ai_metadata(metadata_collection, file_bytes, content_type)
            
            # Calculate quality and completeness scores
            await self._calculate_metadata_scores(metadata_collection)
            
            # Generate extraction summary
            await self._generate_extraction_summary(metadata_collection)
            
            self.logger.info(f"Metadata extraction completed: {content_id}")
            return metadata_collection
            
        except Exception as e:
            self.logger.error(f"Metadata extraction failed: {content_id} - {str(e)}")
            raise MetadataExtractionError(f"Metadata extraction failed: {str(e)}")
    
    async def batch_extract_metadata(self, content_items: List[tuple],
                                   include_ai_analysis: bool = True) -> List[MetadataCollection]:
        """
        Extract metadata from multiple content items in batch.
        
        Args:
            content_items: List of (file_data, filename, content_type) tuples
            include_ai_analysis: Whether to include AI analysis
            
        Returns:
            List of metadata collections
        """
        try:
            self.logger.info(f"Starting batch metadata extraction: {len(content_items)} items")
            
            # Process items concurrently with limited concurrency
            semaphore = asyncio.Semaphore(3)  # Limit concurrent extractions
            
            async def extract_single(item):
        try:
                    # AI model processing
                    if not hasattr(self, 'model') or self.model is None:
                        raise RuntimeError("AI model not initialized")
            
                    # Preprocess input
                    processed_input = await self._preprocess_extract_single_input(data)
            
                    # Run inference
                    result = await self.model.predict(processed_input)
            
                    # Postprocess result
                    final_result = await self._postprocess_extract_single_result(result)
            
                    logger.info(f"AI processing extract_single completed")
                    return final_result
            
                except Exception as e:
                    logger.error(f"AI processing extract_single failed: {e}")
                    raise
                    file_data, filename = item[:2]
                    content_type = item[2] if len(item) > 2 else None
                    return await self.extract_metadata(
                        file_data, filename, content_type, include_ai_analysis
                    )
            
            tasks = [extract_single(item) for item in content_items]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Handle exceptions
            metadata_collections = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    self.logger.error(f"Batch extraction error for item {i}: {str(result)}")
                    # Create empty metadata collection for failed items
                    failed_collection = MetadataCollection(
                        content_id=f"failed_{i}",
                        content_type=ContentFormat.TEXT
                    )
                    failed_collection.extraction_summary = {
                        'status': 'failed',
                        'error': str(result)
                    }
                    metadata_collections.append(failed_collection)
                else:
                    metadata_collections.append(result)
            
            self.logger.info(f"Batch metadata extraction completed: {len(metadata_collections)} results")
            return metadata_collections
            
        except Exception as e:
            self.logger.error(f"Batch metadata extraction failed: {str(e)}")
            raise
    
    async def enrich_metadata(self, metadata_collection: MetadataCollection,
                            additional_sources: Dict[str, Any] = None) -> MetadataCollection:
        """
        Enrich existing metadata with additional sources and AI analysis.
        
        Args:
            metadata_collection: Existing metadata collection
            additional_sources: Additional metadata sources (APIs, databases, etc.)
            
        Returns:
            Enriched metadata collection
        """
        try:
            self.logger.info(f"Enriching metadata: {metadata_collection.content_id}")
            
            # Add additional sources if provided
            if additional_sources:
                for source, data in additional_sources.items():
                    metadata_collection.custom_metadata[source] = data
            
            # Perform additional AI analysis if not already done
            if not metadata_collection.ai_metadata:
                # This would require re-extracting from the original content
                # For now, add placeholder enrichment
                metadata_collection.ai_metadata['enrichment'] = {
                    'enriched_at': datetime.utcnow().isoformat(),
                    'enrichment_sources': list(additional_sources.keys()) if additional_sources else []
                }
            
            # Recalculate scores
            await self._calculate_metadata_scores(metadata_collection)
            
            # Update extraction summary
            metadata_collection.extraction_summary['enriched'] = True
            metadata_collection.extraction_summary['enriched_at'] = datetime.utcnow().isoformat()
            
            return metadata_collection
            
        except Exception as e:
            self.logger.error(f"Metadata enrichment failed: {str(e)}")
            raise
    
    async def validate_metadata(self, metadata_collection: MetadataCollection) -> Dict[str, Any]:
        """
        Validate metadata completeness and consistency.
        
        Args:
            metadata_collection: Metadata collection to validate
            
        Returns:
            Validation results with issues and recommendations
        """
        try:
            validation_results = {
                'is_valid': True,
                'completeness_score': 0.0,
                'issues': [],
                'warnings': [],
                'recommendations': []
            }
            
            # Check required fields
            required_fields = self.extraction_rules['required_fields']
            missing_fields = []
            
            for field in required_fields:
                found = False
                for metadata_dict in [
                    metadata_collection.technical_metadata,
                    metadata_collection.descriptive_metadata,
                    metadata_collection.administrative_metadata
                ]:
                    if field in metadata_dict:
                        found = True
                        break
                
                if not found:
                    missing_fields.append(field)
            
            if missing_fields:
                validation_results['issues'].append(f"Missing required fields: {missing_fields}")
                validation_results['is_valid'] = False
            
            # Calculate completeness score
            total_possible_fields = len(required_fields) + len(self.extraction_rules['optional_fields'])
            present_fields = total_possible_fields - len(missing_fields)
            validation_results['completeness_score'] = (present_fields / total_possible_fields) * 100
            
            # Content-specific validation
            content_type = metadata_collection.content_type
            
            if content_type == ContentFormat.AUDIO:
                if 'duration' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Audio duration not extracted")
                if 'sample_rate' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Audio sample rate not extracted")
            
            elif content_type == ContentFormat.VIDEO:
                if 'resolution' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Video resolution not extracted")
                if 'fps' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Video frame rate not extracted")
            
            elif content_type == ContentFormat.IMAGE:
                if 'dimensions' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Image dimensions not extracted")
                if 'color_mode' not in metadata_collection.technical_metadata:
                    validation_results['warnings'].append("Image color mode not extracted")
            
            # Generate recommendations
            if validation_results['completeness_score'] < 80:
                validation_results['recommendations'].append(
                    "Consider enabling AI analysis for more complete metadata"
                )
            
            if not metadata_collection.ai_metadata:
                validation_results['recommendations'].append(
                    "AI-generated metadata would provide additional insights"
                )
            
            return validation_results
            
        except Exception as e:
            self.logger.error(f"Metadata validation failed: {str(e)}")
            return {
                'is_valid': False,
                'completeness_score': 0.0,
                'issues': [f"Validation error: {str(e)}"],
                'warnings': [],
                'recommendations': []
            }
    
    async def export_metadata(self, metadata_collection: MetadataCollection,
                            format: str = "json") -> Union[str, bytes]:
        """
        Export metadata in various formats.
        
        Args:
            metadata_collection: Metadata to export
            format: Export format (json, xml, dublin_core, etc.)
            
        Returns:
            Exported metadata in requested format
        """
        try:
            if format.lower() == "json":
                return await self._export_json(metadata_collection)
            elif format.lower() == "xml":
                return await self._export_xml(metadata_collection)
            elif format.lower() == "dublin_core":
                return await self._export_dublin_core(metadata_collection)
            else:
                raise ValueError(f"Unsupported export format: {format}")
                
        except Exception as e:
            self.logger.error(f"Metadata export failed: {str(e)}")
            raise
    
    # Private extraction methods
    
    async def _detect_content_format(self, file_data: bytes, filename: str) -> ContentFormat:
        """Auto-detect content format"""
        try:
            # Use python-magic for MIME type detection
            mime_type = magic.from_buffer(file_data, mime=True)
            
            # Map MIME types to content formats
            if mime_type.startswith('audio/'):
                return ContentFormat.AUDIO
            elif mime_type.startswith('video/'):
                return ContentFormat.VIDEO
            elif mime_type.startswith('image/'):
                return ContentFormat.IMAGE
            elif mime_type == 'application/pdf' or filename.endswith('.pdf'):
                return ContentFormat.DOCUMENT
            elif mime_type.startswith('text/') or filename.endswith(('.docx', '.doc', '.rtf')):
                if filename.endswith(('.docx', '.doc', '.rtf', '.odt')):
                    return ContentFormat.DOCUMENT
                else:
                    return ContentFormat.TEXT
            
            # Fallback to file extension
            file_ext = Path(filename).suffix.lower()
            for content_format, extensions in self.supported_formats.items():
                if file_ext in extensions:
                    return content_format
            
            return ContentFormat.TEXT  # Default fallback
            
        except Exception as e:
            self.logger.warning(f"Content format detection failed: {str(e)}")
            return ContentFormat.TEXT
    
    async def _extract_basic_metadata(self, metadata_collection: MetadataCollection,
                                    file_data: bytes, filename: str):
        """Extract basic file metadata"""
        try:
            # Basic file information
            metadata_collection.administrative_metadata.update({
                'original_filename': filename,
                'file_size': len(file_data),
                'file_extension': Path(filename).suffix.lower(),
                'mime_type': magic.from_buffer(file_data, mime=True),
                'file_hash_sha256': hashlib.sha256(file_data).hexdigest(),
                'file_hash_md5': hashlib.md5(file_data).hexdigest(),
                'extraction_timestamp': datetime.utcnow().isoformat(),
                'extractor_version': '1.0.0'
            })
            
            # Technical metadata
            metadata_collection.technical_metadata.update({
                'format': Path(filename).suffix.lower().strip('.'),
                'size_bytes': len(file_data),
                'created_at': datetime.utcnow().isoformat()
            })
            
        except Exception as e:
            self.logger.warning(f"Basic metadata extraction failed: {str(e)}")
    
    async def _extract_audio_metadata(self, metadata_collection: MetadataCollection,
                                    file_data: bytes, filename: str):
        """Extract audio-specific metadata"""
        try:
            # Save to temporary file for processing
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name
            
            try:
                # Use librosa for audio analysis
                audio_data, sample_rate = librosa.load(temp_path, sr=None)
                duration = len(audio_data) / sample_rate
                
                # Technical metadata
                metadata_collection.technical_metadata.update({
                    'duration': float(duration),
                    'sample_rate': int(sample_rate),
                    'channels': 1 if len(audio_data.shape) == 1 else audio_data.shape[1],
                    'bit_depth': 16,  # Default, could be extracted more precisely
                    'audio_codec': 'unknown'  # Would need more detailed analysis
                })
                
                # Use mutagen for ID3/metadata tags
                try:
                    audio_file = mutagen.File(temp_path)
                    if audio_file is not None:
                        # Extract common tags
                        tags = {}
                        
                        # Handle different tag formats
                        if hasattr(audio_file, 'tags') and audio_file.tags:
                            for key, value in audio_file.tags.items():
                                if isinstance(value, list):
                                    tags[key] = value[0] if value else ""
                                else:
                                    tags[key] = str(value)
                        
                        # Standardize common fields
                        descriptive_metadata = {}
                        if 'TIT2' in tags or 'TITLE' in tags:  # Title
                            descriptive_metadata['title'] = tags.get('TIT2', tags.get('TITLE', ''))
                        if 'TPE1' in tags or 'ARTIST' in tags:  # Artist
                            descriptive_metadata['artist'] = tags.get('TPE1', tags.get('ARTIST', ''))
                        if 'TALB' in tags or 'ALBUM' in tags:  # Album
                            descriptive_metadata['album'] = tags.get('TALB', tags.get('ALBUM', ''))
                        if 'TDRC' in tags or 'DATE' in tags:  # Year
                            descriptive_metadata['year'] = tags.get('TDRC', tags.get('DATE', ''))
                        if 'TCON' in tags or 'GENRE' in tags:  # Genre
                            descriptive_metadata['genre'] = tags.get('TCON', tags.get('GENRE', ''))
                        
                        metadata_collection.descriptive_metadata.update(descriptive_metadata)
                        
                        # Store all raw tags
                        metadata_collection.custom_metadata['audio_tags'] = tags
                        
                except Exception as e:
                    self.logger.warning(f"Audio tag extraction failed: {str(e)}")
                
                # Audio feature analysis
                try:
                    # Spectral features
                    spectral_centroids = librosa.feature.spectral_centroid(y=audio_data, sr=sample_rate)
                    mfccs = librosa.feature.mfcc(y=audio_data, sr=sample_rate, n_mfcc=13)
                    
                    audio_features = {
                        'spectral_centroid_mean': float(np.mean(spectral_centroids)),
                        'mfcc_mean': [float(np.mean(mfcc)) for mfcc in mfccs],
                        'tempo': float(librosa.beat.tempo(y=audio_data, sr=sample_rate)[0]),
                        'zero_crossing_rate': float(np.mean(librosa.feature.zero_crossing_rate(audio_data)))
                    }
                    
                    metadata_collection.ai_metadata['audio_features'] = audio_features
                    
                except Exception as e:
                    self.logger.warning(f"Audio feature analysis failed: {str(e)}")
                    
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            self.logger.warning(f"Audio metadata extraction failed: {str(e)}")
    
    async def _extract_video_metadata(self, metadata_collection: MetadataCollection,
                                    file_data: bytes, filename: str):
        """Extract video-specific metadata"""
        try:
            # Save to temporary file for processing
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name
            
            try:
                # Use ffmpeg-python to probe video
                probe = ffmpeg.probe(temp_path)
                
                # Extract format information
                format_info = probe.get('format', {})
                
                # Find video and audio streams
                video_stream = None
                audio_stream = None
                
                for stream in probe.get('streams', []):
                    if stream.get('codec_type') == 'video' and video_stream is None:
                        video_stream = stream
                    elif stream.get('codec_type') == 'audio' and audio_stream is None:
                        audio_stream = stream
                
                # Technical metadata
                technical_metadata = {
                    'duration': float(format_info.get('duration', 0)),
                    'bitrate': int(format_info.get('bit_rate', 0)),
                    'format_name': format_info.get('format_name', ''),
                    'container': Path(filename).suffix.lower().strip('.')
                }
                
                if video_stream:
                    technical_metadata.update({
                        'width': int(video_stream.get('width', 0)),
                        'height': int(video_stream.get('height', 0)),
                        'video_codec': video_stream.get('codec_name', ''),
                        'fps': eval(video_stream.get('r_frame_rate', '0/1')),
                        'pixel_format': video_stream.get('pix_fmt', ''),
                        'video_bitrate': int(video_stream.get('bit_rate', 0))
                    })
                    
                    # Calculate resolution category
                    width = technical_metadata.get('width', 0)
                    height = technical_metadata.get('height', 0)
                    if width >= 3840 and height >= 2160:
                        technical_metadata['resolution_category'] = '4K'
                    elif width >= 1920 and height >= 1080:
                        technical_metadata['resolution_category'] = 'HD'
                    elif width >= 1280 and height >= 720:
                        technical_metadata['resolution_category'] = 'HD Ready'
                    else:
                        technical_metadata['resolution_category'] = 'SD'
                
                if audio_stream:
                    technical_metadata.update({
                        'audio_codec': audio_stream.get('codec_name', ''),
                        'audio_sample_rate': int(audio_stream.get('sample_rate', 0)),
                        'audio_channels': int(audio_stream.get('channels', 0)),
                        'audio_bitrate': int(audio_stream.get('bit_rate', 0))
                    })
                
                metadata_collection.technical_metadata.update(technical_metadata)
                
                # Extract format tags/metadata
                format_tags = format_info.get('tags', {})
                if format_tags:
                    descriptive_metadata = {}
                    
                    # Common video metadata fields
                    title_fields = ['title', 'Title', 'TITLE']
                    for field in title_fields:
                        if field in format_tags:
                            descriptive_metadata['title'] = format_tags[field]
                            break
                    
                    comment_fields = ['comment', 'Comment', 'COMMENT', 'description']
                    for field in comment_fields:
                        if field in format_tags:
                            descriptive_metadata['description'] = format_tags[field]
                            break
                    
                    if descriptive_metadata:
                        metadata_collection.descriptive_metadata.update(descriptive_metadata)
                    
                    # Store all raw tags
                    metadata_collection.custom_metadata['video_tags'] = format_tags
                
                # Video analysis using OpenCV (basic frame analysis)
                try:
                    cap = cv2.VideoCapture(temp_path)
                    if cap.isOpened():
                        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                        
                        # Sample middle frame for analysis
                        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count // 2)
                        ret, frame = cap.read()
                        
                        if ret:
                            # Basic frame analysis
                            frame_analysis = {
                                'average_brightness': float(np.mean(cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY))),
                                'color_variance': float(np.var(frame)),
                                'frame_count': frame_count
                            }
                            
                            metadata_collection.ai_metadata['video_analysis'] = frame_analysis
                        
                        cap.release()
                        
                except Exception as e:
                    self.logger.warning(f"Video frame analysis failed: {str(e)}")
                    
            finally:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            self.logger.warning(f"Video metadata extraction failed: {str(e)}")
    
    async def _extract_image_metadata(self, metadata_collection: MetadataCollection,
                                    file_data: bytes, filename: str):
        """Extract image-specific metadata"""
        try:
            # Load image with PIL
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_data)
                temp_file.flush()
                
                image = Image.open(temp_file.name)
                
                # Basic image metadata
                technical_metadata = {
                    'width': image.width,
                    'height': image.height,
                    'mode': image.mode,
                    'format': image.format or Path(filename).suffix.upper().strip('.'),
                    'has_transparency': image.mode in ['RGBA', 'LA', 'P'],
                    'color_channels': len(image.getbands()) if hasattr(image, 'getbands') else 1
                }
                
                # Calculate additional properties
                technical_metadata['aspect_ratio'] = round(image.width / image.height, 2)
                technical_metadata['total_pixels'] = image.width * image.height
                
                # Image quality indicators
                if technical_metadata['total_pixels'] >= 2000000:  # 2MP+
                    technical_metadata['quality_category'] = 'high'
                elif technical_metadata['total_pixels'] >= 500000:  # 0.5MP+
                    technical_metadata['quality_category'] = 'medium'
                else:
                    technical_metadata['quality_category'] = 'low'
                
                metadata_collection.technical_metadata.update(technical_metadata)
                
                # Extract EXIF data
                try:
                    exif_data = image._getexif()
                    if exif_data:
                        exif_metadata = {}
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            # Convert value to string if it's not JSON serializable
                            if isinstance(value, (bytes, tuple)):
                                value = str(value)
                            exif_metadata[tag] = value
                        
                        # Extract common EXIF fields to descriptive metadata
                        descriptive_metadata = {}
                        if 'ImageDescription' in exif_metadata:
                            descriptive_metadata['description'] = exif_metadata['ImageDescription']
                        if 'Artist' in exif_metadata:
                            descriptive_metadata['artist'] = exif_metadata['Artist']
                        if 'Copyright' in exif_metadata:
                            descriptive_metadata['copyright'] = exif_metadata['Copyright']
                        if 'DateTime' in exif_metadata:
                            descriptive_metadata['date_taken'] = exif_metadata['DateTime']
                        
                        if descriptive_metadata:
                            metadata_collection.descriptive_metadata.update(descriptive_metadata)
                        
                        # Store all EXIF data
                        metadata_collection.custom_metadata['exif_data'] = exif_metadata
                        
                except Exception as e:
                    self.logger.warning(f"EXIF extraction failed: {str(e)}")
                
                # Image analysis
                try:
                    # Convert to numpy array for analysis
                    img_array = np.array(image)
                    
                    if len(img_array.shape) == 3:  # Color image
                        # Color analysis
                        avg_color = np.mean(img_array, axis=(0, 1))
                        color_variance = np.var(img_array, axis=(0, 1))
                        
                        image_analysis = {
                            'average_color_rgb': [float(c) for c in avg_color],
                            'color_variance_rgb': [float(c) for c in color_variance],
                            'brightness': float(np.mean(img_array)),
                            'contrast': float(np.std(img_array))
                        }
                    else:  # Grayscale
                        image_analysis = {
                            'brightness': float(np.mean(img_array)),
                            'contrast': float(np.std(img_array))
                        }
                    
                    # Sharpness estimation using Laplacian variance
                    if image.mode != 'L':
                        gray_image = image.convert('L')
                    else:
                        gray_image = image
                    
                    gray_array = np.array(gray_image)
                    laplacian_var = cv2.Laplacian(gray_array, cv2.CV_64F).var()
                    image_analysis['sharpness_score'] = float(laplacian_var)
                    
                    metadata_collection.ai_metadata['image_analysis'] = image_analysis
                    
                except Exception as e:
                    self.logger.warning(f"Image analysis failed: {str(e)}")
                    
        except Exception as e:
            self.logger.warning(f"Image metadata extraction failed: {str(e)}")
    
    async def _extract_text_metadata(self, metadata_collection: MetadataCollection,
                                   file_data: bytes, filename: str):
        """Extract text/document-specific metadata"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            # Extract text content based on format
            if file_ext == '.pdf':
                text_content, doc_metadata = await self._extract_pdf_metadata(file_data)
            elif file_ext in ['.docx', '.doc']:
                text_content, doc_metadata = await self._extract_docx_metadata(file_data)
            else:
                # Plain text
                try:
                    text_content = file_data.decode('utf-8')
                    doc_metadata = {}
                except UnicodeDecodeError:
                    try:
                        text_content = file_data.decode('latin-1')
                        doc_metadata = {'encoding_detected': 'latin-1'}
                    except UnicodeDecodeError:
                        text_content = file_data.decode('utf-8', errors='ignore')
                        doc_metadata = {'encoding_issues': True}
            
            # Basic text statistics
            technical_metadata = {
                'character_count': len(text_content),
                'word_count': len(text_content.split()),
                'line_count': len(text_content.split('\n')),
                'paragraph_count': len([p for p in text_content.split('\n\n') if p.strip()]),
                'encoding': 'utf-8'
            }
            
            # Add document-specific metadata
            technical_metadata.update(doc_metadata)
            
            metadata_collection.technical_metadata.update(technical_metadata)
            
            # Language detection
            try:
                if len(text_content.strip()) > 10:
                    detected_language = detect(text_content[:1000])  # Use first 1000 chars
                    metadata_collection.descriptive_metadata['language'] = detected_language
            except Exception as e:
                self.logger.warning(f"Language detection failed: {str(e)}")
            
            # Text analysis
            if len(text_content.strip()) > 0:
                try:
                    # Readability analysis
                    readability_scores = {
                        'flesch_reading_ease': textstat.flesch_reading_ease(text_content),
                        'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text_content),
                        'gunning_fog': textstat.gunning_fog(text_content),
                        'automated_readability_index': textstat.automated_readability_index(text_content)
                    }
                    
                    # Text quality metrics
                    avg_words_per_sentence = technical_metadata['word_count'] / max(
                        text_content.count('.') + text_content.count('!') + text_content.count('?'), 1
                    )
                    
                    text_analysis = {
                        'readability_scores': readability_scores,
                        'avg_words_per_sentence': avg_words_per_sentence,
                        'avg_chars_per_word': technical_metadata['character_count'] / max(technical_metadata['word_count'], 1)
                    }
                    
                    metadata_collection.ai_metadata['text_analysis'] = text_analysis
                    
                except Exception as e:
                    self.logger.warning(f"Text analysis failed: {str(e)}")
                
                # NLP analysis with spaCy (if available)
                if self.nlp and len(text_content) < 10000:  # Limit for performance
                    try:
                        doc = self.nlp(text_content[:1000])  # Analyze first 1000 chars
                        
                        # Extract entities
                        entities = []
                        for ent in doc.ents:
                            entities.append({
                                'text': ent.text,
                                'label': ent.label_,
                                'description': spacy.explain(ent.label_)
                            })
                        
                        # Extract key phrases (noun phrases)
                        noun_phrases = [chunk.text for chunk in doc.noun_chunks]
                        
                        nlp_analysis = {
                            'entities': entities[:10],  # Limit to top 10
                            'noun_phrases': noun_phrases[:20],  # Limit to top 20
                            'sentence_count': len(list(doc.sents))
                        }
                        
                        metadata_collection.ai_metadata['nlp_analysis'] = nlp_analysis
                        
                    except Exception as e:
                        self.logger.warning(f"NLP analysis failed: {str(e)}")
                        
        except Exception as e:
            self.logger.warning(f"Text metadata extraction failed: {str(e)}")
    
    async def _extract_pdf_metadata(self, pdf_data: bytes) -> tuple:
        """Extract metadata from PDF"""
        try:
            doc = fitz.open(stream=pdf_data, filetype="pdf")
            
            # Extract text content
            text_content = ""
            for page in doc:
                text_content += page.get_text()
            
            # Extract PDF metadata
            pdf_metadata = doc.metadata
            doc_metadata = {
                'page_count': doc.page_count,
                'pdf_title': pdf_metadata.get('title', ''),
                'pdf_author': pdf_metadata.get('author', ''),
                'pdf_subject': pdf_metadata.get('subject', ''),
                'pdf_creator': pdf_metadata.get('creator', ''),
                'pdf_producer': pdf_metadata.get('producer', ''),
                'pdf_creation_date': pdf_metadata.get('creationDate', ''),
                'pdf_modification_date': pdf_metadata.get('modDate', '')
            }
            
            doc.close()
            return text_content, doc_metadata
            
        except Exception as e:
            self.logger.warning(f"PDF metadata extraction failed: {str(e)}")
            return "", {}
    
    async def _extract_docx_metadata(self, docx_data: bytes) -> tuple:
        """Extract metadata from DOCX"""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(docx_data)
                temp_file.flush()
                
                doc = docx.Document(temp_file.name)
                
                # Extract text content
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Extract document properties
                core_props = doc.core_properties
                doc_metadata = {
                    'docx_title': core_props.title or '',
                    'docx_author': core_props.author or '',
                    'docx_subject': core_props.subject or '',
                    'docx_keywords': core_props.keywords or '',
                    'docx_category': core_props.category or '',
                    'docx_comments': core_props.comments or '',
                    'docx_created': core_props.created.isoformat() if core_props.created else '',
                    'docx_modified': core_props.modified.isoformat() if core_props.modified else '',
                    'docx_last_modified_by': core_props.last_modified_by or '',
                    'paragraph_count': len(doc.paragraphs),
                    'section_count': len(doc.sections)
                }
                
                return text_content, doc_metadata
                
        except Exception as e:
            self.logger.warning(f"DOCX metadata extraction failed: {str(e)}")
            return "", {}
    
    async def _extract_ai_metadata(self, metadata_collection: MetadataCollection,
                                 file_data: bytes, content_type: ContentFormat):
        """Extract AI-powered metadata analysis"""
        try:
            ai_metadata = {}
            
            # Content-specific AI analysis
            if content_type == ContentFormat.IMAGE and self.image_classifier:
                try:
                    # Save image temporarily for AI analysis
                    with tempfile.NamedTemporaryFile(suffix='.jpg') as temp_file:
                        temp_file.write(file_data)
                        temp_file.flush()
                        
                        # Load and classify image
                        image = Image.open(temp_file.name)
                        
                        # Classify image content
                        classifications = self.image_classifier(image)
                        
                        # Extract top predictions
                        top_predictions = classifications[:3]  # Top 3 predictions
                        ai_metadata['image_classification'] = [
                            {
                                'label': pred['label'],
                                'confidence': pred['score']
                            }
                            for pred in top_predictions
                        ]
                        
                except Exception as e:
                    self.logger.warning(f"AI image analysis failed: {str(e)}")
            
            elif content_type in [ContentFormat.TEXT, ContentFormat.DOCUMENT]:
                # Text-based AI analysis
                text_content = ""
                
                # Extract text for analysis
                if content_type == ContentFormat.TEXT:
                    try:
                        text_content = file_data.decode('utf-8')
                    except UnicodeDecodeError:
                        text_content = file_data.decode('utf-8', errors='ignore')
                
                if text_content and len(text_content.strip()) > 10:
                    try:
                        # Sentiment analysis
                        if self.text_classifier and len(text_content) < 512:
                            sentiment_result = self.text_classifier(text_content[:512])
                            ai_metadata['sentiment_analysis'] = {
                                'label': sentiment_result[0]['label'],
                                'confidence': sentiment_result[0]['score']
                            }
                        
                        # Named Entity Recognition
                        if self.ner_model and len(text_content) < 512:
                            entities = self.ner_model(text_content[:512])
                            ai_metadata['named_entities'] = [
                                {
                                    'entity': ent['word'],
                                    'label': ent['entity_group'],
                                    'confidence': ent['score']
                                }
                                for ent in entities[:10]  # Top 10 entities
                            ]
                        
                    except Exception as e:
                        self.logger.warning(f"AI text analysis failed: {str(e)}")
            
            # Add AI metadata timestamp
            if ai_metadata:
                ai_metadata['ai_analysis_timestamp'] = datetime.utcnow().isoformat()
                ai_metadata['ai_models_used'] = []
                
                if self.image_classifier:
                    ai_metadata['ai_models_used'].append('image_classifier')
                if self.text_classifier:
                    ai_metadata['ai_models_used'].append('text_classifier')
                if self.ner_model:
                    ai_metadata['ai_models_used'].append('ner_model')
                if self.nlp:
                    ai_metadata['ai_models_used'].append('spacy_nlp')
            
            metadata_collection.ai_metadata.update(ai_metadata)
            
        except Exception as e:
            self.logger.warning(f"AI metadata extraction failed: {str(e)}")
    
    async def _calculate_metadata_scores(self, metadata_collection: MetadataCollection):
        """Calculate quality and completeness scores for metadata"""
        try:
            # Count extracted fields
            total_fields = 0
            filled_fields = 0
            
            for metadata_dict in [
                metadata_collection.technical_metadata,
                metadata_collection.descriptive_metadata,
                metadata_collection.administrative_metadata,
                metadata_collection.structural_metadata,
                metadata_collection.ai_metadata
            ]:
                for key, value in metadata_dict.items():
                    total_fields += 1
                    if value is not None and value != '' and value != {}:
                        filled_fields += 1
            
            # Calculate completeness score
            if total_fields > 0:
                completeness_score = (filled_fields / total_fields) * 100
            else:
                completeness_score = 0
            
            # Quality score based on presence of key fields
            quality_score = 0
            quality_weights = {
                'required_fields': 40,
                'descriptive_fields': 25,
                'ai_analysis': 20,
                'technical_details': 15
            }
            
            # Check required fields
            required_fields = ['format', 'size_bytes', 'mime_type']
            required_present = sum(
                1 for field in required_fields
                if any(field in md for md in [
                    metadata_collection.technical_metadata,
                    metadata_collection.administrative_metadata
                ])
            )
            quality_score += (required_present / len(required_fields)) * quality_weights['required_fields']
            
            # Check descriptive fields
            descriptive_fields = ['title', 'description', 'author', 'language']
            descriptive_present = sum(
                1 for field in descriptive_fields
                if field in metadata_collection.descriptive_metadata
            )
            if descriptive_fields:
                quality_score += (descriptive_present / len(descriptive_fields)) * quality_weights['descriptive_fields']
            
            # Check AI analysis
            if metadata_collection.ai_metadata:
                quality_score += quality_weights['ai_analysis']
            
            # Check technical details (format-specific)
            content_type = metadata_collection.content_type
            tech_score = 0
            
            if content_type == ContentFormat.AUDIO:
                audio_fields = ['duration', 'sample_rate', 'channels']
                audio_present = sum(
                    1 for field in audio_fields
                    if field in metadata_collection.technical_metadata
                )
                tech_score = (audio_present / len(audio_fields)) * quality_weights['technical_details']
                
            elif content_type == ContentFormat.VIDEO:
                video_fields = ['duration', 'width', 'height', 'fps']
                video_present = sum(
                    1 for field in video_fields
                    if field in metadata_collection.technical_metadata
                )
                tech_score = (video_present / len(video_fields)) * quality_weights['technical_details']
                
            elif content_type == ContentFormat.IMAGE:
                image_fields = ['width', 'height', 'mode', 'format']
                image_present = sum(
                    1 for field in image_fields
                    if field in metadata_collection.technical_metadata
                )
                tech_score = (image_present / len(image_fields)) * quality_weights['technical_details']
                
            elif content_type in [ContentFormat.TEXT, ContentFormat.DOCUMENT]:
                text_fields = ['character_count', 'word_count', 'language']
                text_present = sum(
                    1 for field in text_fields
                    if field in metadata_collection.technical_metadata or
                       field in metadata_collection.descriptive_metadata
                )
                tech_score = (text_present / len(text_fields)) * quality_weights['technical_details']
            
            quality_score += tech_score
            
            # Set scores
            metadata_collection.completeness_score = round(completeness_score, 2)
            metadata_collection.quality_score = round(quality_score, 2)
            
        except Exception as e:
            self.logger.warning(f"Score calculation failed: {str(e)}")
            metadata_collection.completeness_score = 0.0
            metadata_collection.quality_score = 0.0
    
    async def _generate_extraction_summary(self, metadata_collection: MetadataCollection):
        """Generate extraction summary with statistics and insights"""
        try:
            summary = {
                'extraction_status': 'completed',
                'content_type': metadata_collection.content_type.value,
                'total_fields_extracted': 0,
                'metadata_categories': [],
                'ai_analysis_performed': bool(metadata_collection.ai_metadata),
                'quality_assessment': {
                    'completeness_score': metadata_collection.completeness_score,
                    'quality_score': metadata_collection.quality_score
                },
                'extraction_timestamp': datetime.utcnow().isoformat()
            }
            
            # Count fields by category
            categories = {
                'technical': metadata_collection.technical_metadata,
                'descriptive': metadata_collection.descriptive_metadata,
                'administrative': metadata_collection.administrative_metadata,
                'structural': metadata_collection.structural_metadata,
                'ai_generated': metadata_collection.ai_metadata,
                'custom': metadata_collection.custom_metadata
            }
            
            for category, metadata_dict in categories.items():
                if metadata_dict:
                    field_count = len(metadata_dict)
                    summary['total_fields_extracted'] += field_count
                    summary['metadata_categories'].append({
                        'category': category,
                        'field_count': field_count,
                        'fields': list(metadata_dict.keys())
                    })
            
            # Add insights based on content type
            insights = []
            
            if metadata_collection.content_type == ContentFormat.AUDIO:
                if 'duration' in metadata_collection.technical_metadata:
                    duration = metadata_collection.technical_metadata['duration']
                    if duration > 300:  # 5 minutes
                        insights.append("Long-form audio content detected")
                    elif duration < 30:
                        insights.append("Short audio clip detected")
                
                if 'audio_features' in metadata_collection.ai_metadata:
                    insights.append("Audio feature analysis completed")
            
            elif metadata_collection.content_type == ContentFormat.VIDEO:
                tech_meta = metadata_collection.technical_metadata
                if 'width' in tech_meta and 'height' in tech_meta:
                    if tech_meta['width'] >= 1920 and tech_meta['height'] >= 1080:
                        insights.append("High-definition video content")
                    if tech_meta.get('fps', 0) >= 60:
                        insights.append("High frame rate video")
            
            elif metadata_collection.content_type == ContentFormat.IMAGE:
                if 'exif_data' in metadata_collection.custom_metadata:
                    insights.append("Camera metadata available")
                if metadata_collection.technical_metadata.get('total_pixels', 0) > 2000000:
                    insights.append("High-resolution image")
            
            elif metadata_collection.content_type in [ContentFormat.TEXT, ContentFormat.DOCUMENT]:
                if 'language' in metadata_collection.descriptive_metadata:
                    lang = metadata_collection.descriptive_metadata['language']
                    insights.append(f"Content language: {lang}")
                
                word_count = metadata_collection.technical_metadata.get('word_count', 0)
                if word_count > 1000:
                    insights.append("Long-form text content")
                elif word_count < 100:
                    insights.append("Short text content")
            
            summary['insights'] = insights
            
            metadata_collection.extraction_summary = summary
            
        except Exception as e:
            self.logger.warning(f"Summary generation failed: {str(e)}")
            metadata_collection.extraction_summary = {
                'extraction_status': 'completed_with_errors',
                'error': str(e),
                'extraction_timestamp': datetime.utcnow().isoformat()
            }
    
    # Export methods
    
    async def _export_json(self, metadata_collection: MetadataCollection) -> str:
        """Export metadata as JSON"""
        try:
            export_data = {
                'content_id': metadata_collection.content_id,
                'content_type': metadata_collection.content_type.value,
                'technical_metadata': metadata_collection.technical_metadata,
                'descriptive_metadata': metadata_collection.descriptive_metadata,
                'administrative_metadata': metadata_collection.administrative_metadata,
                'structural_metadata': metadata_collection.structural_metadata,
                'ai_metadata': metadata_collection.ai_metadata,
                'preservation_metadata': metadata_collection.preservation_metadata,
                'custom_metadata': metadata_collection.custom_metadata,
                'extraction_summary': metadata_collection.extraction_summary,
                'quality_scores': {
                    'quality_score': metadata_collection.quality_score,
                    'completeness_score': metadata_collection.completeness_score
                },
                'export_timestamp': datetime.utcnow().isoformat(),
                'export_format': 'json'
            }
            
            return json.dumps(export_data, indent=2, default=str)
            
        except Exception as e:
            self.logger.error(f"JSON export failed: {str(e)}")
            raise
    
    async def _export_xml(self, metadata_collection: MetadataCollection) -> str:
        """Export metadata as XML"""
        try:
            import xml.etree.ElementTree as ET
            
            root = ET.Element("metadata")
            root.set("content_id", metadata_collection.content_id)
            root.set("content_type", metadata_collection.content_type.value)
            
            # Add metadata sections
            sections = {
                'technical': metadata_collection.technical_metadata,
                'descriptive': metadata_collection.descriptive_metadata,
                'administrative': metadata_collection.administrative_metadata,
                'structural': metadata_collection.structural_metadata,
                'ai_generated': metadata_collection.ai_metadata,
                'custom': metadata_collection.custom_metadata
            }
            
            for section_name, metadata_dict in sections.items():
                if metadata_dict:
                    section_elem = ET.SubElement(root, section_name)
                    for key, value in metadata_dict.items():
                        field_elem = ET.SubElement(section_elem, "field")
                        field_elem.set("name", key)
                        field_elem.text = str(value)
            
            # Add quality scores
            quality_elem = ET.SubElement(root, "quality_scores")
            quality_elem.set("quality_score", str(metadata_collection.quality_score))
            quality_elem.set("completeness_score", str(metadata_collection.completeness_score))
            
            return ET.tostring(root, encoding='unicode')
            
        except Exception as e:
            self.logger.error(f"XML export failed: {str(e)}")
            raise
    
    async def _export_dublin_core(self, metadata_collection: MetadataCollection) -> str:
        """Export metadata in Dublin Core format"""
        try:
            dublin_core_mapping = {
                'title': metadata_collection.descriptive_metadata.get('title', ''),
                'creator': metadata_collection.descriptive_metadata.get('author', ''),
                'subject': ', '.join(metadata_collection.descriptive_metadata.get('tags', [])),
                'description': metadata_collection.descriptive_metadata.get('description', ''),
                'publisher': metadata_collection.descriptive_metadata.get('publisher', ''),
                'contributor': metadata_collection.descriptive_metadata.get('contributor', ''),
                'date': metadata_collection.administrative_metadata.get('created_at', ''),
                'type': metadata_collection.content_type.value,
                'format': metadata_collection.technical_metadata.get('mime_type', ''),
                'identifier': metadata_collection.content_id,
                'source': metadata_collection.administrative_metadata.get('original_filename', ''),
                'language': metadata_collection.descriptive_metadata.get('language', ''),
                'relation': '',
                'coverage': '',
                'rights': metadata_collection.descriptive_metadata.get('copyright', '')
            }
            
            # Build Dublin Core XML
            import xml.etree.ElementTree as ET
            
            root = ET.Element("metadata")
            root.set("xmlns:dc", "http://purl.org/dc/elements/1.1/")
            
            for dc_field, value in dublin_core_mapping.items():
                if value:
                    elem = ET.SubElement(root, f"dc:{dc_field}")
                    elem.text = str(value)
            
            return ET.tostring(root, encoding='unicode')
            
        except Exception as e:
            self.logger.error(f"Dublin Core export failed: {str(e)}")
            raise
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get supported formats by content type"""
        return {fmt.value: exts for fmt, exts in self.supported_formats.items()}
    
    def get_extraction_capabilities(self) -> Dict[str, Any]:
        """
Get extractor capabilities and configuration"""
        return {
            'supported_formats': self.get_supported_formats(),
            'ai_models_available': {
                'image_classifier': self.image_classifier is not None,
                'text_classifier': self.text_classifier is not None,
                'ner_model': self.ner_model is not None,
                'nlp_model': self.nlp is not None
            },
            'metadata_types': [mt.value for mt in MetadataType],
            'export_formats': ['json', 'xml', 'dublin_core'],
            'extraction_rules': self.extraction_rules
        }
